from __future__ import annotations

import csv
import importlib.util
import json
import re
import wave
from array import array
from copy import copy
from pathlib import Path
from typing import Any


def render_output_file(
    service: Any,
    *,
    output_path: Path,
    output_format: str,
    title: str,
    content: str,
    table_rows: list[list[str]],
    formatting: dict[str, Any],
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if output_format in {"txt", "md", "html", "json"}:
        output_path.write_text(content, encoding="utf-8")
        return
    if output_format == "csv":
        rows = table_rows or service._extract_table_rows_from_markdown(content)
        if not rows:
            rows = [["内容"], *[[line] for line in content.splitlines() if line.strip()]]
        with output_path.open("w", encoding="utf-8-sig", newline="") as stream:
            writer = csv.writer(stream)
            writer.writerows(rows)
        return
    if output_format == "xlsx":
        write_xlsx(
            service,
            output_path=output_path,
            title=title,
            content=content,
            table_rows=table_rows,
            formatting=formatting,
        )
        return
    if output_format == "docx":
        write_docx(service, output_path=output_path, title=title, content=content, formatting=formatting)
        return
    if output_format == "pdf":
        write_pdf(service, output_path=output_path, title=title, content=content)
        return
    raise ValueError(f"不支持的输出格式：{output_format}")


def style_existing_xlsx(service: Any, *, source_path: Path, output_path: Path, formatting: dict[str, Any]) -> None:
    if importlib.util.find_spec("openpyxl") is None:
        raise RuntimeError("缺少 openpyxl，无法加工 Excel 文档。")
    from openpyxl import load_workbook  # type: ignore

    output_path.parent.mkdir(parents=True, exist_ok=True)
    workbook = load_workbook(str(source_path))
    sheet = workbook.active
    apply_xlsx_formatting(service, sheet=sheet, rows=[], formatting=formatting)
    workbook.save(str(output_path))


def style_existing_docx(service: Any, *, source_path: Path, output_path: Path, formatting: dict[str, Any]) -> None:
    if importlib.util.find_spec("docx") is None:
        raise RuntimeError("缺少 python-docx，无法加工 Word 文档。")
    from docx import Document  # type: ignore

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document(str(source_path))
    for paragraph in document.paragraphs:
        apply_docx_paragraph_rules(service, paragraph, formatting)
    for table in document.tables:
        rows = docx_table_rows(table)
        apply_docx_table_formatting(service, table, rows, formatting)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    apply_docx_paragraph_rules(service, paragraph, formatting)
    document.save(str(output_path))


def docx_table_rows(table: Any) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in getattr(table, "rows", []) or []:
        cells = [str(getattr(cell, "text", "") or "") for cell in getattr(row, "cells", []) or []]
        if cells:
            rows.append(cells)
    return rows


def write_docx(service: Any, *, output_path: Path, title: str, content: str, formatting: dict[str, Any]) -> None:
    if importlib.util.find_spec("docx") is None:
        raise RuntimeError("缺少 python-docx，无法生成 Word 文档。")
    from docx import Document  # type: ignore

    document = Document()
    if title:
        document.add_heading(title, level=1)
    for block in service._split_markdown_blocks(content):
        if not block:
            continue
        if block.startswith("#"):
            heading = block.lstrip("#").strip()
            level = min(4, max(1, len(block) - len(block.lstrip("#"))))
            paragraph = document.add_heading(heading or "标题", level=level)
            apply_docx_paragraph_rules(service, paragraph, formatting)
            continue
        table_rows = service._extract_table_rows_from_markdown(block)
        if table_rows and len(table_rows) >= 2:
            table = document.add_table(rows=len(table_rows), cols=max(len(row) for row in table_rows))
            table.style = "Table Grid"
            for row_index, row in enumerate(table_rows):
                for col_index, cell_text in enumerate(row):
                    table.cell(row_index, col_index).text = str(cell_text)
            apply_docx_table_formatting(service, table, table_rows, formatting)
            continue
        if block.startswith(("- ", "* ")):
            for line in block.splitlines():
                text = re.sub(r"^[-*]\s+", "", line).strip()
                if text:
                    paragraph = document.add_paragraph(text, style="List Bullet")
                    apply_docx_paragraph_rules(service, paragraph, formatting)
            continue
        paragraph = document.add_paragraph(block)
        apply_docx_paragraph_rules(service, paragraph, formatting)
    document.save(str(output_path))


def write_xlsx(
    service: Any,
    *,
    output_path: Path,
    title: str,
    content: str,
    table_rows: list[list[str]],
    formatting: dict[str, Any],
) -> None:
    if importlib.util.find_spec("openpyxl") is None:
        raise RuntimeError("缺少 openpyxl，无法生成 Excel 文档。")
    from openpyxl import Workbook  # type: ignore

    rows = table_rows or service._extract_table_rows_from_markdown(content)
    if not rows:
        rows = [["内容"], *[[line] for line in content.splitlines() if line.strip()]]
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = service._safe_sheet_name(title or "生成内容")
    for row in rows:
        sheet.append([str(cell) for cell in row])
    apply_xlsx_formatting(service, sheet=sheet, rows=rows, formatting=formatting)
    workbook.save(str(output_path))


def apply_xlsx_formatting(service: Any, *, sheet: Any, rows: list[list[str]], formatting: dict[str, Any]) -> None:
    if not formatting:
        return
    max_row = int(getattr(sheet, "max_row", 0) or 0)
    max_col = int(getattr(sheet, "max_column", 0) or 0)
    if max_row <= 0 or max_col <= 0:
        return
    headers = rows[0] if rows else [str(sheet.cell(row=1, column=col).value or "") for col in range(1, max_col + 1)]

    header_style = formatting.get("header") if isinstance(formatting.get("header"), dict) else {}
    if header_style:
        for col in range(1, max_col + 1):
            apply_xlsx_cell_style(sheet.cell(row=1, column=col), header_style)

    for rule in formatting.get("columns") or []:
        if not isinstance(rule, dict):
            continue
        col = service._resolve_column_index(rule, headers)
        if not col or col > max_col:
            continue
        for row_index in range(1, max_row + 1):
            apply_xlsx_cell_style(sheet.cell(row=row_index, column=col), rule)

    for rule in formatting.get("rows") or []:
        if not isinstance(rule, dict):
            continue
        start, end = service._resolve_row_range(rule, max_row=max_row)
        if not start:
            continue
        for row_index in range(start, end + 1):
            for col in range(1, max_col + 1):
                apply_xlsx_cell_style(sheet.cell(row=row_index, column=col), rule)

    for rule in formatting.get("row_rules") or []:
        if not isinstance(rule, dict):
            continue
        for row_index in range(1, max_row + 1):
            if xlsx_row_matches(service, sheet=sheet, row_index=row_index, headers=headers, rule=rule):
                for col in range(1, max_col + 1):
                    apply_xlsx_cell_style(sheet.cell(row=row_index, column=col), rule)

    for rule in formatting.get("cells") or []:
        if not isinstance(rule, dict):
            continue
        row_index = service._coerce_int(rule.get("row") or rule.get("row_index") or rule.get("index"))
        col = service._resolve_column_index(rule, headers)
        if not row_index or not col or row_index > max_row or col > max_col:
            continue
        apply_xlsx_cell_style(sheet.cell(row=row_index, column=col), rule)

    for rule in formatting.get("highlights") or []:
        if not isinstance(rule, dict):
            continue
        needle = str(rule.get("text") or rule.get("contains") or "").strip()
        if not needle:
            continue
        for row_index in range(1, max_row + 1):
            for col in range(1, max_col + 1):
                cell = sheet.cell(row=row_index, column=col)
                if needle in str(cell.value or ""):
                    apply_xlsx_cell_style(cell, rule)

    if formatting.get("auto_width") is not False:
        apply_xlsx_auto_width(sheet)


def apply_xlsx_cell_style(cell: Any, style: dict[str, Any]) -> None:
    if not isinstance(style, dict):
        return
    font_changed = False
    font = copy(cell.font)
    if "bold" in style:
        font.bold = bool(style.get("bold"))
        font_changed = True
    if "italic" in style:
        font.italic = bool(style.get("italic"))
        font_changed = True
    font_color = str(style.get("font_color") or "").strip()
    if font_color:
        font.color = font_color
        font_changed = True
    if font_changed:
        cell.font = font
    fill_color = str(style.get("fill_color") or style.get("highlight_color") or "").strip()
    if fill_color:
        from openpyxl.styles import PatternFill  # type: ignore

        cell.fill = PatternFill(fill_type="solid", fgColor=fill_color)


def apply_xlsx_auto_width(sheet: Any) -> None:
    from openpyxl.utils import get_column_letter  # type: ignore

    max_col = int(getattr(sheet, "max_column", 0) or 0)
    max_row = int(getattr(sheet, "max_row", 0) or 0)
    for col in range(1, max_col + 1):
        width = 10
        for row_index in range(1, max_row + 1):
            width = max(width, min(42, len(str(sheet.cell(row=row_index, column=col).value or "")) + 2))
        sheet.column_dimensions[get_column_letter(col)].width = width


def xlsx_row_matches(service: Any, *, sheet: Any, row_index: int, headers: list[str], rule: dict[str, Any]) -> bool:
    where = rule.get("where") if isinstance(rule.get("where"), dict) else {}
    match_text = str(rule.get("match_text") or where.get("match_text") or where.get("contains") or "").strip()
    if match_text:
        max_col = int(getattr(sheet, "max_column", 0) or 0)
        return any(match_text in str(sheet.cell(row=row_index, column=col).value or "") for col in range(1, max_col + 1))

    column_selector = where.get("column") or where.get("match_header") or rule.get("column") or rule.get("match_header")
    if column_selector is None:
        return False
    col = service._resolve_column_index({"column": column_selector, "match_header": column_selector}, headers)
    if not col:
        return False
    raw_value = sheet.cell(row=row_index, column=col).value
    if "eq" in where:
        return str(raw_value or "") == str(where.get("eq") or "")
    if "ne" in where:
        return str(raw_value or "") != str(where.get("ne") or "")
    number = service._coerce_float(raw_value)
    if number is None:
        return False
    for op, comparator in (
        ("lt", lambda a, b: a < b),
        ("lte", lambda a, b: a <= b),
        ("gt", lambda a, b: a > b),
        ("gte", lambda a, b: a >= b),
    ):
        if op in where:
            threshold = service._coerce_float(where.get(op))
            return threshold is not None and comparator(number, threshold)
    return False


def apply_docx_table_formatting(service: Any, table: Any, rows: list[list[str]], formatting: dict[str, Any]) -> None:
    if not formatting or not rows:
        return
    headers = rows[0]
    row_count = len(table.rows)
    col_count = max((len(row.cells) for row in table.rows), default=0)

    header_style = formatting.get("header") if isinstance(formatting.get("header"), dict) else {}
    if header_style and row_count:
        for cell in table.rows[0].cells:
            apply_docx_cell_style(service, cell, header_style)

    for rule in formatting.get("columns") or []:
        if not isinstance(rule, dict):
            continue
        col = service._resolve_column_index(rule, headers)
        if not col or col > col_count:
            continue
        for row in table.rows:
            apply_docx_cell_style(service, row.cells[col - 1], rule)

    for rule in formatting.get("rows") or []:
        if not isinstance(rule, dict):
            continue
        start, end = service._resolve_row_range(rule, max_row=row_count)
        if not start:
            continue
        for row_index in range(start, end + 1):
            for cell in table.rows[row_index - 1].cells:
                apply_docx_cell_style(service, cell, rule)

    for rule in formatting.get("cells") or []:
        if not isinstance(rule, dict):
            continue
        row_index = service._coerce_int(rule.get("row") or rule.get("row_index") or rule.get("index"))
        col = service._resolve_column_index(rule, headers)
        if not row_index or not col or row_index > row_count or col > col_count:
            continue
        apply_docx_cell_style(service, table.rows[row_index - 1].cells[col - 1], rule)

    for rule in formatting.get("highlights") or []:
        if not isinstance(rule, dict):
            continue
        needle = str(rule.get("text") or rule.get("contains") or "").strip()
        if not needle:
            continue
        for row in table.rows:
            for cell in row.cells:
                if needle in str(cell.text or ""):
                    apply_docx_cell_style(service, cell, rule)


def apply_docx_cell_style(service: Any, cell: Any, style: dict[str, Any]) -> None:
    if not isinstance(style, dict):
        return
    fill_color = str(style.get("fill_color") or "").strip()
    if fill_color:
        shade_docx_cell(cell, fill_color)
    for paragraph in cell.paragraphs:
        apply_docx_runs_style(service, paragraph, style)


def apply_docx_paragraph_rules(service: Any, paragraph: Any, formatting: dict[str, Any]) -> None:
    if not formatting:
        return
    text = str(getattr(paragraph, "text", "") or "")
    matched_styles: list[dict[str, Any]] = []
    for rule in formatting.get("highlights") or []:
        if not isinstance(rule, dict):
            continue
        needle = str(rule.get("text") or rule.get("contains") or "").strip()
        if needle and needle in text:
            matched_styles.append(rule)
    for rule in formatting.get("paragraphs") or []:
        if not isinstance(rule, dict):
            continue
        needle = str(rule.get("contains") or rule.get("text") or "").strip()
        index = service._coerce_int(rule.get("index") or rule.get("paragraph_index"))
        if needle and needle in text:
            matched_styles.append(rule)
        elif index is None and not needle:
            matched_styles.append(rule)
    for style in matched_styles:
        apply_docx_runs_style(service, paragraph, style)


def apply_docx_runs_style(service: Any, paragraph: Any, style: dict[str, Any]) -> None:
    if not isinstance(style, dict):
        return
    runs = list(getattr(paragraph, "runs", []) or [])
    if not runs and str(getattr(paragraph, "text", "") or ""):
        runs = [paragraph.add_run("")]
    for run in runs:
        apply_docx_run_style(service, run, style)


def apply_docx_run_style(service: Any, run: Any, style: dict[str, Any]) -> None:
    if "bold" in style:
        run.bold = bool(style.get("bold"))
    if "italic" in style:
        run.italic = bool(style.get("italic"))
    font_color = str(style.get("font_color") or "").strip()
    if font_color:
        from docx.shared import RGBColor  # type: ignore

        run.font.color.rgb = RGBColor.from_string(font_color)
    highlight = str(style.get("highlight_color") or style.get("fill_color") or "").strip()
    highlight_color = docx_highlight_color(highlight)
    if highlight_color is not None:
        run.font.highlight_color = highlight_color


def shade_docx_cell(cell: Any, fill_color: str) -> None:
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_color)
    tc_pr.append(shd)


def docx_highlight_color(color: str) -> Any | None:
    if not color:
        return None
    from docx.enum.text import WD_COLOR_INDEX  # type: ignore

    color_upper = color.upper()
    if color_upper in {"FFFF00", "FFF2CC", "FFD966"}:
        return WD_COLOR_INDEX.YELLOW
    if color_upper in {"FF0000", "F4CCCC", "EA9999"}:
        return WD_COLOR_INDEX.RED
    if color_upper in {"00FF00", "D9EAD3", "B6D7A8"}:
        return WD_COLOR_INDEX.BRIGHT_GREEN
    if color_upper in {"0000FF", "CFE2F3", "9FC5E8"}:
        return WD_COLOR_INDEX.BLUE
    if color_upper in {"D9D9D9", "CCCCCC", "999999"}:
        return WD_COLOR_INDEX.GRAY_25
    return WD_COLOR_INDEX.YELLOW


def write_pdf(service: Any, *, output_path: Path, title: str, content: str) -> None:
    if importlib.util.find_spec("reportlab") is None:
        raise RuntimeError("缺少 reportlab，无法生成 PDF。")
    from reportlab.lib.pagesizes import A4  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
    from reportlab.pdfbase import pdfmetrics  # type: ignore
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont  # type: ignore
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # type: ignore

    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        font_name = "STSong-Light"
    except Exception:
        font_name = "Helvetica"

    doc = SimpleDocTemplate(str(output_path), pagesize=A4)
    styles = getSampleStyleSheet()
    for style_name in ("Title", "Heading1", "Heading2", "BodyText"):
        styles[style_name].fontName = font_name
    story = []
    if title:
        story.append(Paragraph(service._escape_pdf_text(title), styles["Title"]))
        story.append(Spacer(1, 12))
    for block in service._split_markdown_blocks(content):
        if not block:
            continue
        if block.startswith("#"):
            text = block.lstrip("#").strip()
            story.append(Paragraph(service._escape_pdf_text(text), styles["Heading1"]))
        else:
            text = "<br/>".join(service._escape_pdf_text(line) for line in block.splitlines())
            story.append(Paragraph(text, styles["BodyText"]))
        story.append(Spacer(1, 8))
    doc.build(story)


def write_pcm16_wav(path: Path, *, samples: array, sample_rate: int, channels: int = 1) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with wave.open(str(path), "wb") as writer:
        writer.setnchannels(max(1, int(channels or 1)))
        writer.setsampwidth(2)
        writer.setframerate(max(1, int(sample_rate or 44100)))
        writer.writeframes(samples.tobytes())


def render_voice_dataset_readme(service: Any, manifest: dict[str, Any]) -> str:
    stats = manifest.get("stats") if isinstance(manifest.get("stats"), dict) else {}
    lines = [
        f"# {manifest.get('title') or 'Voice Dataset'}",
        "",
        f"- Profile: {manifest.get('profile') or 'gpt_sovits'}",
        f"- Slices: {stats.get('slice_count') or 0}",
        f"- Recommended: {stats.get('recommended_count') or 0}",
        f"- Flagged: {stats.get('flagged_count') or 0}",
        f"- Total duration: {service._format_duration_label(stats.get('total_duration_seconds')) or '0:00'}",
        "",
        "## Issue Slices",
    ]
    issue_slices = manifest.get("issue_slices") if isinstance(manifest.get("issue_slices"), dict) else {}
    if not issue_slices:
        lines.append("- None")
    for flag, items in issue_slices.items():
        if not isinstance(items, list) or not items:
            continue
        lines.append(f"- {flag}: " + ", ".join(str(item.get("filename") or "") for item in items[:30] if item.get("filename")))
    lines.extend(["", "Full metadata is stored in `manifest.json`."])
    return "\n".join(lines)


def render_transcript_output(
    service: Any,
    *,
    transcripts: list[dict[str, Any]],
    output_format: str,
    title: str,
    with_timestamps: bool,
) -> str:
    if output_format == "json":
        return json.dumps(
            {
                "title": title,
                "transcripts": transcripts,
            },
            ensure_ascii=False,
            indent=2,
        )
    if output_format == "srt":
        return render_srt_transcripts(service, transcripts)
    if output_format == "vtt":
        return render_vtt_transcripts(service, transcripts)
    if output_format == "txt":
        return render_plain_transcripts(service, transcripts, with_timestamps=with_timestamps)
    return render_markdown_transcripts(service, transcripts, title=title, with_timestamps=with_timestamps)


def render_markdown_transcripts(
    service: Any,
    transcripts: list[dict[str, Any]],
    *,
    title: str,
    with_timestamps: bool,
) -> str:
    lines = [f"# {title}", ""]
    for transcript in transcripts:
        source = transcript.get("source") if isinstance(transcript.get("source"), dict) else {}
        source_label = source.get("handle") or source.get("title") or f"source_{transcript.get('source_index')}"
        lines.extend([f"## {source_label}", ""])
        if with_timestamps:
            for segment in list(transcript.get("segments") or []):
                lines.append(
                    f"- [{service._format_timestamp_label(segment.get('start'))} - "
                    f"{service._format_timestamp_label(segment.get('end'))}] {segment.get('text') or ''}"
                )
        else:
            text = str(transcript.get("text") or "").strip()
            if text:
                lines.append(text)
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def render_plain_transcripts(service: Any, transcripts: list[dict[str, Any]], *, with_timestamps: bool) -> str:
    lines: list[str] = []
    for transcript in transcripts:
        source = transcript.get("source") if isinstance(transcript.get("source"), dict) else {}
        source_label = source.get("handle") or source.get("title") or f"source_{transcript.get('source_index')}"
        lines.append(f"【{source_label}】")
        if with_timestamps:
            for segment in list(transcript.get("segments") or []):
                lines.append(
                    f"[{service._format_timestamp_label(segment.get('start'))} - "
                    f"{service._format_timestamp_label(segment.get('end'))}] {segment.get('text') or ''}"
                )
        else:
            lines.append(str(transcript.get("text") or "").strip())
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def render_srt_transcripts(service: Any, transcripts: list[dict[str, Any]]) -> str:
    lines: list[str] = []
    cue_index = 1
    offset = 0.0
    for transcript in transcripts:
        for segment in list(transcript.get("segments") or []):
            start = offset + float(segment.get("start") or 0)
            end = offset + float(segment.get("end") or start)
            lines.extend(
                [
                    str(cue_index),
                    f"{service._format_srt_timestamp(start)} --> {service._format_srt_timestamp(end)}",
                    str(segment.get("text") or "").strip(),
                    "",
                ]
            )
            cue_index += 1
        offset += float(
            transcript.get("duration_seconds") or service._estimate_transcript_duration(list(transcript.get("segments") or [])) or 0
        )
    return "\n".join(lines).strip() + "\n"


def render_vtt_transcripts(service: Any, transcripts: list[dict[str, Any]]) -> str:
    lines = ["WEBVTT", ""]
    offset = 0.0
    for transcript in transcripts:
        source = transcript.get("source") if isinstance(transcript.get("source"), dict) else {}
        source_label = source.get("handle") or source.get("title") or f"source_{transcript.get('source_index')}"
        lines.extend([f"NOTE {source_label}", ""])
        for segment in list(transcript.get("segments") or []):
            start = offset + float(segment.get("start") or 0)
            end = offset + float(segment.get("end") or start)
            lines.extend(
                [
                    f"{service._format_vtt_timestamp(start)} --> {service._format_vtt_timestamp(end)}",
                    str(segment.get("text") or "").strip(),
                    "",
                ]
            )
        offset += float(
            transcript.get("duration_seconds") or service._estimate_transcript_duration(list(transcript.get("segments") or [])) or 0
        )
    return "\n".join(lines).strip() + "\n"
