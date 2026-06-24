from __future__ import annotations

import importlib.util
import json
import time
import zipfile
from pathlib import Path
from typing import Any


def send_generated_file(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str = "latest",
    targets: list[str] | tuple[str, ...] | None = None,
    timestamp: int | None = None,
) -> dict[str, Any]:
    resolved_targets = service._normalize_send_targets(target=target, targets=targets)
    generated_files: list[dict[str, Any]] = []
    unresolved: list[str] = []
    missing_on_disk: list[str] = []
    ambiguous_targets: list[str] = []

    for item in resolved_targets:
        generated = service._resolve_generated_file(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=item,
        )
        if generated is None:
            ambiguous_label = format_generated_file_ambiguity(
                service,
                profile_user_id=profile_user_id,
                session_id=session_id,
                target=item,
                statuses=["ready"],
            )
            if ambiguous_label:
                ambiguous_targets.append(ambiguous_label)
            else:
                unresolved.append(item)
            continue

        absolute_path = service.absolute_path(generated)
        if not absolute_path.exists() or not absolute_path.is_file():
            missing_on_disk.append(str(generated.get("generated_handle") or item).strip() or item)
            continue

        updated = service.store.update_generated_file(
            profile_user_id=profile_user_id,
            session_id=session_id,
            generated_id=str(generated.get("generated_id") or ""),
            delivery_status="pending",
            updated_at=timestamp,
        ) or generated
        updated["absolute_path"] = str(absolute_path)
        generated_files.append(updated)

    if not generated_files:
        return {
            "ok": False,
            "generated": None,
            "generated_files": [],
            "send_to_user": False,
            "unresolved": unresolved,
            "missing_on_disk": missing_on_disk,
            "ambiguous_targets": ambiguous_targets,
            "error": (
                "generated_file_ambiguous"
                if ambiguous_targets
                else "generated_file_not_found"
                if unresolved
                else "generated_file_missing_on_disk"
            ),
            "followup_context": service._build_send_followup_missing(
                requested_targets=resolved_targets,
                unresolved=unresolved,
                missing_on_disk=missing_on_disk,
                ambiguous_targets=ambiguous_targets,
            ),
        }

    return {
        "ok": True,
        "generated": generated_files[0],
        "generated_files": generated_files,
        "send_to_user": True,
        "unresolved": unresolved,
        "missing_on_disk": missing_on_disk,
        "ambiguous_targets": ambiguous_targets,
        "followup_context": service._build_send_followup_batch(
            generated_files=generated_files,
            unresolved=unresolved,
            missing_on_disk=missing_on_disk,
            ambiguous_targets=ambiguous_targets,
        ),
    }


def send_file(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str = "latest",
    targets: list[str] | tuple[str, ...] | None = None,
    timestamp: int | None = None,
) -> dict[str, Any]:
    resolved_targets = service._normalize_send_targets(target=target, targets=targets)
    files: list[dict[str, Any]] = []
    unresolved: list[str] = []
    missing_on_disk: list[str] = []
    ambiguous_targets: list[str] = []

    for item in resolved_targets:
        file_ref, error = service._resolve_sendable_file(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=item,
            timestamp=timestamp,
        )
        if file_ref is None:
            if error == "missing_on_disk":
                missing_on_disk.append(item)
            elif error == "ambiguous_generated_file":
                ambiguous_label = format_generated_file_ambiguity(
                    service,
                    profile_user_id=profile_user_id,
                    session_id=session_id,
                    target=item,
                    statuses=["ready"],
                )
                ambiguous_targets.append(ambiguous_label or item)
            elif error == "ambiguous_attachment_file":
                ambiguity = service.attachment_service.format_attachment_ambiguity(
                    profile_user_id=profile_user_id,
                    session_id=session_id,
                    target=item,
                    kind="any",
                    statuses=["ready"],
                )
                ambiguous_targets.append(ambiguity or item)
            else:
                unresolved.append(item)
            continue
        files.append(file_ref)

    if not files:
        return {
            "ok": False,
            "files": [],
            "unresolved": unresolved,
            "missing_on_disk": missing_on_disk,
            "ambiguous_targets": ambiguous_targets,
            "send_to_user": False,
            "error": (
                "file_target_ambiguous"
                if ambiguous_targets
                else "file_not_found"
                if unresolved
                else "file_missing_on_disk"
            ),
            "followup_context": service._build_send_file_followup_missing(
                requested_targets=resolved_targets,
                unresolved=unresolved,
                missing_on_disk=missing_on_disk,
                ambiguous_targets=ambiguous_targets,
            ),
        }

    return {
        "ok": True,
        "files": files,
        "unresolved": unresolved,
        "missing_on_disk": missing_on_disk,
        "ambiguous_targets": ambiguous_targets,
        "send_to_user": True,
        "followup_context": service._build_send_file_followup_batch(
            files=files,
            unresolved=unresolved,
            missing_on_disk=missing_on_disk,
            ambiguous_targets=ambiguous_targets,
        ),
    }


def inspect_generated_file(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str = "latest",
    section: str = "content",
    max_chars: int = 12000,
) -> dict[str, Any]:
    normalized_section = service._normalize_generated_inspection_section(section)
    char_budget = service._normalize_inspection_max_chars(max_chars)
    generated = service._resolve_generated_file(
        profile_user_id=profile_user_id,
        session_id=session_id,
        target=target,
    )
    if generated is None:
        return {
            "ok": False,
            "generated": None,
            "inspection": None,
            "error": "generated_file_not_found",
            "followup_context": (
                f"你刚刚想查看 {target or 'latest'}，但没有找到对应的生成文件。"
                "请自然向用户确认要查看哪一份，不要重复调用 inspect_generated_file。"
            ),
        }

    absolute_path = service.absolute_path(generated)
    generated["absolute_path"] = str(absolute_path)
    if normalized_section != "summary" and (not absolute_path.exists() or not absolute_path.is_file()):
        return {
            "ok": False,
            "generated": generated,
            "inspection": None,
            "error": "generated_file_missing_on_disk",
            "followup_context": (
                f"你刚刚想查看 {service._generated_display_name(generated)}，但本地文件本体已经找不到了。"
                "你仍然可以参考生成文件工作台里的摘要；请自然告诉用户文件记录还在，但文件本体缺失。"
            ),
        }

    inspection = service._inspect_generated_file_content(
        generated=generated,
        path=absolute_path,
        section=normalized_section,
        max_chars=char_budget,
    )
    return {
        "ok": True,
        "generated": generated,
        "inspection": inspection,
        "followup_context": service._build_generated_inspection_followup(
            generated=generated,
            inspection=inspection,
        ),
    }


def apply_style_to_existing_file(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str = "latest",
    instruction: str = "",
    output_title: str = "",
    formatting: dict[str, Any] | None = None,
    send_to_user: bool = True,
    target_type: str = "",
    timestamp: int | None = None,
) -> dict[str, Any]:
    effective_ts = int(timestamp or time.time())
    style_rules = service._normalize_formatting(formatting)
    if not style_rules:
        return {
            "ok": False,
            "generated": None,
            "send_to_user": False,
            "error": "missing_formatting",
            "followup_context": (
                "你刚刚想给已有文件套样式，但没有提供可执行的 formatting 规则。"
                "请先明确要标红/高亮/加粗的列、行、关键词或单元格，不要重复调用工具。"
            ),
        }

    source = service._resolve_existing_file_for_style(
        profile_user_id=profile_user_id,
        session_id=session_id,
        target=target,
        target_type=target_type,
    )
    if source is None:
        return {
            "ok": False,
            "generated": None,
            "send_to_user": False,
            "error": "source_file_not_found",
            "followup_context": (
                f"你刚刚想给 {target or 'latest'} 套样式，但没有找到对应的附件或生成文件。"
                "请自然向用户确认要处理哪份文件，不要重复调用 apply_style_to_existing_file。"
            ),
        }

    source_path = source.get("absolute_path")
    if not isinstance(source_path, Path) or not source_path.exists() or not source_path.is_file():
        return {
            "ok": False,
            "generated": None,
            "send_to_user": False,
            "error": "source_file_missing_on_disk",
            "followup_context": (
                f"你刚刚想给 {source.get('handle') or target} 套样式，"
                "但本地文件本体已经找不到了。请自然告诉用户文件记录还在，但文件本体缺失。"
            ),
        }

    source_format = service._normalize_output_format(source.get("output_format"))
    if source_format not in {"xlsx", "docx"}:
        return {
            "ok": False,
            "generated": None,
            "send_to_user": False,
            "error": f"unsupported_style_format:{source_format or 'unknown'}",
            "followup_context": (
                f"你刚刚想直接给 {source.get('handle') or target} 套样式，"
                f"但当前只支持对 docx/xlsx 做无内容重写的样式加工；这份文件格式是 {source_format or 'unknown'}。"
                "如果用户要改写内容或换格式，请改用 compose_file 或 revise_generated_file。"
            ),
        }

    title = service._normalize_title(output_title) or f"{source.get('title') or '生成文件'} 样式版"
    output_path = service._build_output_path(
        profile_user_id=profile_user_id,
        session_id=session_id,
        title=title,
        output_format=source_format,
        timestamp=effective_ts,
    )
    try:
        if source_format == "xlsx":
            service._style_existing_xlsx(
                source_path=source_path,
                output_path=output_path,
                formatting=style_rules,
            )
        elif source_format == "docx":
            service._style_existing_docx(
                source_path=source_path,
                output_path=output_path,
                formatting=style_rules,
            )
    except Exception as exc:
        return {
            "ok": False,
            "generated": None,
            "send_to_user": False,
            "error": str(exc),
            "followup_context": (
                f"你刚刚尝试给「{source.get('title') or target}」套样式但失败：{str(exc)[:180]}。"
                "请自然告诉用户失败原因；如果是缺少依赖，可以提醒先安装对应 Python 库。"
            ),
        }

    source_id = str(source.get("source_id") or "").strip()
    source_ids = [source_id] if source_id else []
    for extra_id in list(source.get("extra_source_ids") or [])[:8]:
        text = str(extra_id or "").strip()
        if text and text not in source_ids:
            source_ids.append(text)

    content_card = service._build_style_content_card(
        title=title,
        output_format=source_format,
        source=source,
        instruction=instruction,
        formatting=style_rules,
    )
    generated = service.store.add_generated_file(
        profile_user_id=profile_user_id,
        session_id=session_id,
        output_title=title,
        output_format=source_format,
        storage_relpath=service._storage_relpath(output_path),
        mime_type=service._mime_type_for_format(source_format),
        file_ext=source_format,
        file_size=output_path.stat().st_size,
        source_ids=source_ids,
        content_card=content_card,
        summary=str(content_card.get("summary") or "").strip(),
        created_by_tool="apply_style_to_existing_file",
        version_of_generated_id=str(source.get("source_id") or "").strip() if source.get("source_type") == "generated" else "",
        version_no=int(source.get("version_no") or 1) + 1 if source.get("source_type") == "generated" else 1,
        delivery_status="pending" if send_to_user else "not_requested",
        timestamp=effective_ts,
    )
    generated["absolute_path"] = str(service.absolute_path(generated))
    return {
        "ok": True,
        "generated": generated,
        "source": source,
        "send_to_user": bool(send_to_user),
        "followup_context": service._build_style_followup(
            source=source,
            generated=generated,
            send_to_user=send_to_user,
        ),
    }


def manage_generated_files(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    action: str,
    targets: list[str] | tuple[str, ...] | set[str] | str | None = None,
    reason: str = "",
    timestamp: int | None = None,
) -> dict[str, Any]:
    effective_ts = int(timestamp or time.time())
    normalized_action = service._normalize_generated_file_action(action)
    if not normalized_action:
        return {
            "ok": False,
            "managed": [],
            "unresolved": [],
            "action": "",
            "followup_context": (
                "你刚刚想管理生成文件，但 action 不明确。"
                "可用 action 是 archive、delete 或 purge；请自然向用户确认。"
            ),
        }

    normalized_targets = service._normalize_targets(targets)
    resolved, unresolved = service._resolve_generated_file_targets(
        profile_user_id=profile_user_id,
        session_id=session_id,
        targets=normalized_targets,
    )
    if not resolved:
        missing = f"未找到：{', '.join(unresolved[:5])}。" if unresolved else ""
        return {
            "ok": False,
            "managed": [],
            "unresolved": unresolved,
            "action": normalized_action,
            "followup_context": (
                "你刚刚想清理生成文件，但没有找到可管理的生成物。"
                f"{missing}请自然向用户确认要处理哪一份，不要重复调用 manage_generated_file。"
            ),
        }

    managed: list[dict[str, Any]] = []
    for item in resolved:
        item_id = str(item.get("generated_id") or "").strip()
        if not item_id:
            continue
        deleted_file = False
        delete_error = ""
        if normalized_action in {"delete", "purge"}:
            deleted_file, delete_error = service._delete_generated_file_on_disk(item)

        update_kwargs: dict[str, Any] = {
            "profile_user_id": profile_user_id,
            "session_id": session_id,
            "generated_id": item_id,
            "status": "removed",
            "delivery_status": normalized_action,
            "updated_at": effective_ts,
        }
        if normalized_action == "purge":
            update_kwargs.update(
                {
                    "summary": "",
                    "content_card": {
                        "summary": "生成文件内容已清理。",
                        "purged": True,
                        "purged_at": effective_ts,
                        "reason": str(reason or "").strip()[:200],
                    },
                    "storage_relpath": "",
                    "file_size": 0,
                }
            )
        updated = service.store.update_generated_file(**update_kwargs) or dict(item, status="removed")
        updated["file_deleted"] = deleted_file
        if delete_error:
            updated["delete_error"] = delete_error
        managed.append(updated)

    names = "、".join(service._generated_display_name(item) for item in managed[:3])
    if len(managed) > 3:
        names += f" 等 {len(managed)} 个"
    missing = f"未找到：{', '.join(unresolved[:5])}。" if unresolved else ""
    reason_text = f"原因：{str(reason).strip()}。" if str(reason or "").strip() else ""
    action_label = {
        "archive": "归档隐藏",
        "delete": "删除本地文件并归档",
        "purge": "彻底清理内容并归档",
    }.get(normalized_action, normalized_action)
    return {
        "ok": bool(managed),
        "managed": managed,
        "unresolved": unresolved,
        "action": normalized_action,
        "followup_context": (
            f"你刚刚已经对 {len(managed)} 个生成文件执行了「{action_label}」"
            f"（{names}）。{missing}{reason_text}"
            "这些生成物不会继续显示在生成文件工作台里；用户原始附件不会被删除。"
            "请基于这个既成事实自然回应，不要重复调用 manage_generated_file。"
        ),
    }


def resolve_generated_file(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str,
) -> dict[str, Any] | None:
    normalized = str(target or "").strip()
    if not normalized or normalized.lower() in {"latest", "current", "最近", "当前"}:
        items = service.store.list_generated_files(
            profile_user_id=profile_user_id,
            session_id=session_id,
            statuses=["ready"],
            limit=1,
        )
        return items[0] if items else None
    return service.store.find_generated_file(
        profile_user_id=profile_user_id,
        session_id=session_id,
        query=normalized,
        statuses=["ready"],
    )


def resolve_generated_file_targets(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    targets: list[str],
) -> tuple[list[dict[str, Any]], list[str]]:
    if not targets:
        targets = ["latest"]
    lowered = {target.lower() for target in targets}
    if lowered & {"all", "全部", "*"}:
        return (
            service.store.list_generated_files(
                profile_user_id=profile_user_id,
                session_id=session_id,
                statuses=["ready", "failed"],
                limit=200,
            ),
            [],
        )

    resolved: list[dict[str, Any]] = []
    unresolved: list[str] = []
    seen: set[str] = set()
    for target in targets:
        item = service._resolve_generated_file_any_status(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=target,
        )
        if item is None:
            unresolved.append(target)
            continue
        item_id = str(item.get("generated_id") or "").strip()
        if item_id and item_id not in seen:
            seen.add(item_id)
            resolved.append(item)
    return resolved, unresolved


def resolve_generated_file_any_status(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str,
) -> dict[str, Any] | None:
    normalized = str(target or "").strip()
    if not normalized or normalized.lower() in {"latest", "current", "最近", "当前"}:
        items = service.store.list_generated_files(
            profile_user_id=profile_user_id,
            session_id=session_id,
            statuses=["ready", "failed"],
            limit=1,
        )
        return items[0] if items else None
    return service.store.find_generated_file(
        profile_user_id=profile_user_id,
        session_id=session_id,
        query=normalized,
        statuses=["ready", "failed"],
    )


def find_generated_file_ambiguity_candidates(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str,
    statuses: list[str] | tuple[str, ...] | set[str] | None,
) -> list[dict[str, Any]]:
    normalized = str(target or "").strip()
    if not normalized:
        return []
    matches = service.store.find_generated_file_matches(
        profile_user_id=profile_user_id,
        session_id=session_id,
        query=normalized,
        statuses=statuses,
        limit=8,
    )
    if not matches:
        return []
    best_rank = int(matches[0].get("_match_rank") or 99)
    if best_rank <= 2:
        return []
    candidates = [
        {key: value for key, value in item.items() if key != "_match_rank"}
        for item in matches
        if int(item.get("_match_rank") or 99) == best_rank
    ]
    return candidates if len(candidates) > 1 else []


def format_generated_file_ambiguity(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str,
    statuses: list[str] | tuple[str, ...] | set[str] | None,
) -> str:
    candidates = find_generated_file_ambiguity_candidates(
        service,
        profile_user_id=profile_user_id,
        session_id=session_id,
        target=target,
        statuses=statuses,
    )
    if not candidates:
        return ""
    labels = [
        service._generated_display_name(item)
        for item in candidates[:4]
    ]
    return f"{target} -> {'、'.join(label for label in labels if label)}"


def normalize_targets(service: Any, value: list[str] | tuple[str, ...] | set[str] | str | None) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        raw_items = [item.strip() for item in value.replace("，", ",").replace("、", ",").split(",")]
    elif isinstance(value, (list, tuple, set)):
        raw_items = list(value)
    else:
        raw_items = [value]
    targets: list[str] = []
    for item in raw_items:
        text = str(item or "").strip()
        if text and text not in targets:
            targets.append(text[:120])
    return targets[:50]


def normalize_generated_file_action(service: Any, value: Any) -> str:
    action = str(value or "").strip().lower()
    aliases = {
        "hide": "archive",
        "archive": "archive",
        "remove": "archive",
        "clear": "archive",
        "收起": "archive",
        "归档": "archive",
        "隐藏": "archive",
        "delete": "delete",
        "unlink": "delete",
        "删除": "delete",
        "删掉": "delete",
        "purge": "purge",
        "destroy": "purge",
        "彻底删除": "purge",
        "彻底清理": "purge",
    }
    return aliases.get(action, "")


def delete_generated_file_on_disk(service: Any, item: dict[str, Any]) -> tuple[bool, str]:
    path = service.absolute_path(item)
    if not service.is_managed_storage_path(path):
        return False, "unsafe_path"
    if not path.exists():
        return False, "missing"
    if not path.is_file():
        return False, "not_file"
    try:
        path.unlink()
        return True, ""
    except Exception as exc:
        return False, str(exc)[:180]


def resolve_sendable_file(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str,
    timestamp: int | None,
) -> tuple[dict[str, Any] | None, str]:
    normalized = str(target or "").strip() or "latest"
    lowered = normalized.lower()
    if lowered in {"latest", "current", "最近", "当前"}:
        return service._resolve_latest_sendable_file(
            profile_user_id=profile_user_id,
            session_id=session_id,
            timestamp=timestamp,
        )

    prefer_generated = lowered.startswith(("gen_", "generated::"))
    prefer_attachment = lowered.startswith(("file_", "img_", "image_", "audio_", "video_", "attachment::"))
    if prefer_generated:
        return service._resolve_generated_sendable_file(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=normalized,
            timestamp=timestamp,
        )
    if prefer_attachment:
        return service._resolve_attachment_sendable_file(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=normalized,
        )

    generated, generated_error = service._resolve_generated_sendable_file(
        profile_user_id=profile_user_id,
        session_id=session_id,
        target=normalized,
        timestamp=timestamp,
    )
    if generated is not None or generated_error in {"missing_on_disk", "ambiguous_generated_file"}:
        return generated, generated_error
    return service._resolve_attachment_sendable_file(
        profile_user_id=profile_user_id,
        session_id=session_id,
        target=normalized,
    )


def resolve_latest_sendable_file(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    timestamp: int | None,
) -> tuple[dict[str, Any] | None, str]:
    generated_items = service.store.list_generated_files(
        profile_user_id=profile_user_id,
        session_id=session_id,
        statuses=["ready"],
        limit=1,
    )
    attachment_items = service.store.list_attachment_inbox_items(
        profile_user_id=profile_user_id,
        session_id=session_id,
        statuses=["ready"],
        limit=1,
    )
    generated = generated_items[0] if generated_items else None
    attachment = attachment_items[0] if attachment_items else None
    if generated is None and attachment is None:
        return None, "not_found"
    if generated is not None and attachment is not None:
        generated_ts = int(generated.get("updated_at") or generated.get("created_at") or 0)
        attachment_ts = int(attachment.get("updated_at") or attachment.get("created_at") or 0)
        if generated_ts >= attachment_ts:
            return service._generated_item_to_sendable_file(generated, timestamp=timestamp)
        return service._attachment_item_to_sendable_file(attachment)
    if generated is not None:
        return service._generated_item_to_sendable_file(generated, timestamp=timestamp)
    return service._attachment_item_to_sendable_file(attachment or {})


def resolve_generated_sendable_file(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str,
    timestamp: int | None,
) -> tuple[dict[str, Any] | None, str]:
    generated = service._resolve_generated_file(
        profile_user_id=profile_user_id,
        session_id=session_id,
        target=target,
    )
    if generated is None:
        if find_generated_file_ambiguity_candidates(
            service,
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=target,
            statuses=["ready"],
        ):
            return None, "ambiguous_generated_file"
        return None, "not_found"
    return service._generated_item_to_sendable_file(generated, timestamp=timestamp)


def resolve_attachment_sendable_file(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str,
) -> tuple[dict[str, Any] | None, str]:
    attachment = service.attachment_service.resolve_attachment(
        profile_user_id=profile_user_id,
        session_id=session_id,
        target=target,
        kind="any",
    )
    if attachment is None:
        ambiguity = service.attachment_service.format_attachment_ambiguity(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=target,
            kind="any",
            statuses=["ready"],
        )
        if ambiguity:
            return None, "ambiguous_attachment_file"
        return None, "not_found"
    return service._attachment_item_to_sendable_file(attachment)


def generated_item_to_sendable_file(
    service: Any,
    generated: dict[str, Any],
    *,
    timestamp: int | None,
) -> tuple[dict[str, Any] | None, str]:
    absolute_path = service.absolute_path(generated)
    if not absolute_path.exists() or not absolute_path.is_file():
        return None, "missing_on_disk"
    updated = service.store.update_generated_file(
        profile_user_id=str(generated.get("profile_user_id") or ""),
        session_id=str(generated.get("session_id") or ""),
        generated_id=str(generated.get("generated_id") or ""),
        delivery_status="pending",
        updated_at=timestamp,
    ) or generated
    updated["absolute_path"] = str(absolute_path)
    output_format = str(updated.get("output_format") or absolute_path.suffix.lstrip(".")).strip().lower()
    title = str(updated.get("output_title") or updated.get("generated_handle") or absolute_path.stem).strip()
    name = f"{title}.{output_format}" if output_format and not title.lower().endswith(f".{output_format}") else title
    return (
        {
            "source_type": "generated",
            "source_id": str(updated.get("generated_id") or "").strip(),
            "generated_id": str(updated.get("generated_id") or "").strip(),
            "handle": str(updated.get("generated_handle") or "").strip(),
            "title": title,
            "name": name or absolute_path.name,
            "absolute_path": str(absolute_path),
            "file_ext": output_format,
            "mime_type": str(updated.get("mime_type") or service._mime_type_for_format(output_format)),
            "file_size": int(updated.get("file_size") or absolute_path.stat().st_size),
            "generated_file": updated,
        },
        "",
    )


def attachment_item_to_sendable_file(service: Any, attachment: dict[str, Any]) -> tuple[dict[str, Any] | None, str]:
    if not attachment:
        return None, "not_found"
    absolute_path = service.attachment_service.resolve_storage_path(attachment)
    if absolute_path is None or not absolute_path.exists() or not absolute_path.is_file():
        return None, "missing_on_disk"
    handle = str(attachment.get("attachment_handle") or attachment.get("attachment_id") or "").strip()
    origin_name = str(attachment.get("origin_name") or "").strip()
    ext = str(attachment.get("file_ext") or absolute_path.suffix.lstrip(".")).strip().lstrip(".")
    title = origin_name or handle or absolute_path.stem
    if ext and title and not title.lower().endswith(f".{ext.lower()}"):
        name = f"{title}.{ext}"
    else:
        name = title or absolute_path.name
    return (
        {
            "source_type": "attachment",
            "source_id": str(attachment.get("attachment_id") or "").strip(),
            "attachment_id": str(attachment.get("attachment_id") or "").strip(),
            "handle": handle,
            "title": title,
            "name": name or absolute_path.name,
            "absolute_path": str(absolute_path),
            "file_ext": ext,
            "mime_type": str(attachment.get("mime_type") or ""),
            "file_size": int(attachment.get("file_size") or absolute_path.stat().st_size),
            "attachment": attachment,
        },
        "",
    )


def normalize_generated_inspection_section(service: Any, value: Any) -> str:
    text = str(value or "").strip()
    lowered = text.lower()
    if lowered.startswith(("file:", "member:")):
        prefix, _, rest = text.partition(":")
        return f"file:{rest.strip()[:240]}"
    aliases = {
        "": "content",
        "full": "content",
        "all": "content",
        "content": "content",
        "body": "content",
        "正文": "content",
        "全文": "content",
        "全部": "content",
        "head": "head",
        "start": "head",
        "开头": "head",
        "前面": "head",
        "tail": "tail",
        "end": "tail",
        "结尾": "tail",
        "末尾": "tail",
        "summary": "summary",
        "摘要": "summary",
        "概览": "summary",
        "manifest": "manifest",
        "manifest.json": "file:manifest.json",
        "readme": "file:README.md",
        "readme.md": "file:README.md",
        "file_list": "file_list",
        "list": "file_list",
        "members": "file_list",
        "文件列表": "file_list",
        "列表": "file_list",
    }
    return aliases.get(lowered, lowered if lowered in {"content", "head", "tail", "summary", "manifest", "file_list"} else "content")


def normalize_inspection_max_chars(service: Any, value: Any) -> int:
    try:
        parsed = int(float(str(value).strip()))
    except Exception:
        parsed = 12000
    return max(500, min(40000, parsed))


def inspect_generated_file_content(
    service: Any,
    *,
    generated: dict[str, Any],
    path: Path,
    section: str,
    max_chars: int,
) -> dict[str, Any]:
    output_format = str(generated.get("output_format") or path.suffix.lstrip(".")).strip().lower()
    if section == "summary":
        content = service._render_generated_summary_inspection(generated=generated, path=path)
        return {
            "section": "summary",
            "content": service._clip_inspection_text(content, max_chars=max_chars),
            "truncated": len(content) > max_chars,
            "source_kind": "metadata",
        }

    if output_format == "zip" or path.suffix.lower() == ".zip":
        return service._inspect_generated_zip(
            generated=generated,
            path=path,
            section=section,
            max_chars=max_chars,
        )

    text = service._read_generated_text_material(path=path, output_format=output_format, max_chars=max_chars * 3)
    if not text:
        content = service._render_generated_binary_inspection(generated=generated, path=path, output_format=output_format)
        return {
            "section": section,
            "content": service._clip_inspection_text(content, max_chars=max_chars),
            "truncated": False,
            "source_kind": "binary_or_unsupported",
        }

    sliced = service._slice_inspection_text(text, section=section, max_chars=max_chars)
    return {
        "section": section,
        "content": sliced["content"],
        "truncated": sliced["truncated"],
        "source_kind": output_format or "text",
    }


def read_generated_text_material(
    service: Any,
    *,
    path: Path,
    output_format: str,
    max_chars: int,
    text_inspect_formats: set[str],
) -> str:
    suffix = (output_format or path.suffix.lstrip(".")).lower()
    try:
        if suffix in text_inspect_formats:
            return service._read_plain_text_file(path, max_chars=max_chars)
        if suffix == "docx":
            return service._read_generated_docx(path=path, max_chars=max_chars)
        if suffix == "xlsx":
            return service._read_generated_xlsx(path=path, max_chars=max_chars)
        if suffix == "pdf":
            return service._read_generated_pdf(path=path, max_chars=max_chars)
    except Exception as exc:
        return f"读取生成文件失败：{str(exc)[:300]}"
    return ""


def read_plain_text_file(service: Any, path: Path, *, max_chars: int) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "cp936", "latin-1"):
        try:
            text = path.read_text(encoding=encoding)
            return text[: max_chars + 1]
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="replace")[: max_chars + 1]


def read_generated_docx(service: Any, *, path: Path, max_chars: int) -> str:
    if importlib.util.find_spec("docx") is None:
        return "缺少 python-docx，无法直接读取 Word 正文。"
    from docx import Document  # type: ignore

    document = Document(str(path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = str(paragraph.text or "").strip()
        if text:
            lines.append(text)
        if len("\n".join(lines)) >= max_chars:
            break
    for table in document.tables:
        if len("\n".join(lines)) >= max_chars:
            break
        for row in table.rows:
            cells = [str(cell.text or "").strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")
            if len("\n".join(lines)) >= max_chars:
                break
    return "\n".join(lines)[: max_chars + 1]


def read_generated_xlsx(service: Any, *, path: Path, max_chars: int) -> str:
    if importlib.util.find_spec("openpyxl") is None:
        return "缺少 openpyxl，无法直接读取 Excel 内容。"
    from openpyxl import load_workbook  # type: ignore

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"## Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if cell is None else str(cell) for cell in row]
            if any(cell.strip() for cell in cells):
                lines.append("| " + " | ".join(cells) + " |")
            if len("\n".join(lines)) >= max_chars:
                return "\n".join(lines)[: max_chars + 1]
    return "\n".join(lines)[: max_chars + 1]


def read_generated_pdf(service: Any, *, path: Path, max_chars: int) -> str:
    if importlib.util.find_spec("pypdf") is None:
        return "缺少 pypdf，无法直接读取 PDF 文本。"
    from pypdf import PdfReader  # type: ignore

    reader = PdfReader(str(path))
    lines: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = str(page.extract_text() or "").strip()
        if text:
            lines.extend([f"## Page {index}", text])
        if len("\n".join(lines)) >= max_chars:
            break
    return "\n".join(lines)[: max_chars + 1]


def inspect_generated_zip(
    service: Any,
    *,
    generated: dict[str, Any],
    path: Path,
    section: str,
    max_chars: int,
) -> dict[str, Any]:
    try:
        with zipfile.ZipFile(path, "r") as archive:
            if section == "file_list":
                content = service._render_zip_file_list(archive)
                sliced = service._slice_inspection_text(content, section="content", max_chars=max_chars)
                return {
                    "section": "file_list",
                    "content": sliced["content"],
                    "truncated": sliced["truncated"],
                    "source_kind": "zip",
                }
            if section == "manifest":
                section = "file:manifest.json"
            if section.startswith("file:"):
                member_name = service._resolve_zip_member_name(archive, section.partition(":")[2].strip())
                content = service._read_zip_member_for_inspection(archive, member_name=member_name, max_chars=max_chars)
                sliced = service._slice_inspection_text(content, section="content", max_chars=max_chars)
                return {
                    "section": f"file:{member_name}" if member_name else section,
                    "content": sliced["content"],
                    "truncated": sliced["truncated"],
                    "source_kind": "zip_member",
                }

            member = service._resolve_zip_member_name(archive, "manifest.json") or service._resolve_zip_member_name(archive, "README.md")
            if member:
                content = service._read_zip_member_for_inspection(archive, member_name=member, max_chars=max_chars)
            else:
                content = service._render_zip_file_list(archive)
            sliced = service._slice_inspection_text(content, section="content", max_chars=max_chars)
            return {
                "section": f"file:{member}" if member else "file_list",
                "content": sliced["content"],
                "truncated": sliced["truncated"],
                "source_kind": "zip",
            }
    except Exception as exc:
        content = (
            f"读取 zip 生成物 {service._generated_display_name(generated)} 失败：{str(exc)[:300]}。"
        )
        return {
            "section": section,
            "content": content,
            "truncated": False,
            "source_kind": "zip_error",
        }


def render_zip_file_list(service: Any, archive: zipfile.ZipFile) -> str:
    lines = ["zip 文件列表："]
    for info in archive.infolist()[:400]:
        if info.is_dir():
            continue
        lines.append(f"- {info.filename}（{service._format_file_size(int(info.file_size))}）")
    if len(archive.infolist()) > 400:
        lines.append(f"... 还有 {len(archive.infolist()) - 400} 个条目未显示。")
    return "\n".join(lines)


def resolve_zip_member_name(service: Any, archive: zipfile.ZipFile, target: str) -> str:
    normalized = str(target or "").replace("\\", "/").strip().lstrip("/")
    if not normalized or ".." in normalized.split("/"):
        return ""
    names = [info.filename for info in archive.infolist() if not info.is_dir()]
    for name in names:
        if name == normalized:
            return name
    lowered = normalized.lower()
    for name in names:
        if name.lower() == lowered:
            return name
    for name in names:
        if name.lower().endswith("/" + lowered) or lowered in name.lower():
            return name
    return ""


def read_zip_member_for_inspection(
    service: Any,
    archive: zipfile.ZipFile,
    *,
    member_name: str,
    max_chars: int,
    text_inspect_formats: set[str],
) -> str:
    if not member_name:
        return "没有在 zip 里找到对应文件。"
    info = archive.getinfo(member_name)
    suffix = Path(member_name).suffix.lower().lstrip(".")
    basename = Path(member_name).name.lower()
    if suffix not in text_inspect_formats and basename not in {"readme", "license"}:
        return (
            f"{member_name} 是 zip 内的二进制或不适合直接展开的文件，"
            f"大小 {service._format_file_size(int(info.file_size))}。"
        )
    data = archive.read(member_name)
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "cp936", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = data.decode(errors="replace")
    if suffix == "json":
        try:
            text = json.dumps(json.loads(text), ensure_ascii=False, indent=2)
        except Exception:
            pass
    return text[: max_chars + 1]


def slice_inspection_text(service: Any, text: str, *, section: str, max_chars: int) -> dict[str, Any]:
    raw = str(text or "")
    if section == "tail" and len(raw) > max_chars:
        return {
            "content": "（以下是文件末尾片段）\n" + raw[-max_chars:],
            "truncated": True,
        }
    if section == "head" and len(raw) > max_chars:
        return {
            "content": raw[:max_chars] + "\n\n（内容未显示完；如需继续查看，可指定 tail 或更大的 max_chars。）",
            "truncated": True,
        }
    return {
        "content": service._clip_inspection_text(raw, max_chars=max_chars),
        "truncated": len(raw) > max_chars,
    }


def clip_inspection_text(service: Any, text: str, *, max_chars: int) -> str:
    raw = str(text or "")
    if len(raw) <= max_chars:
        return raw
    return raw[:max_chars] + "\n\n（内容已按 max_chars 截断；如需更多，请指定更大的 max_chars 或查看 tail。）"


def resolve_existing_file_for_style(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str,
    target_type: str,
) -> dict[str, Any] | None:
    normalized = str(target or "").strip()
    normalized_type = str(target_type or "").strip().lower()
    lowered = normalized.lower()
    looks_like_attachment = lowered.startswith(("file_", "img_", "image_", "audio_"))
    looks_like_generated = lowered.startswith(("gen_", "generated::"))

    prefer_attachment = normalized_type in {"attachment", "inbox", "file"} or looks_like_attachment
    prefer_generated = normalized_type in {"generated", "gen"} or looks_like_generated

    if prefer_attachment:
        attachment = service._resolve_attachment_style_source(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=normalized,
        )
        if attachment is not None:
            return attachment
        if prefer_generated:
            return None

    if not prefer_attachment:
        generated = service._resolve_generated_style_source(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=normalized,
        )
        if generated is not None:
            return generated

    if not prefer_generated:
        attachment = service._resolve_attachment_style_source(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=normalized,
        )
        if attachment is not None:
            return attachment
    return None


def resolve_generated_style_source(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str,
) -> dict[str, Any] | None:
    generated = service._resolve_generated_file(
        profile_user_id=profile_user_id,
        session_id=session_id,
        target=target,
    )
    if generated is None:
        return None
    content_card = generated.get("content_card") if isinstance(generated.get("content_card"), dict) else {}
    media_info = content_card.get("media_info") if isinstance(content_card.get("media_info"), dict) else {}
    absolute_path = service.absolute_path(generated)
    return {
        "source_type": "generated",
        "source_id": str(generated.get("generated_id") or "").strip(),
        "handle": str(generated.get("generated_handle") or "").strip(),
        "title": str(generated.get("output_title") or "").strip() or "生成文件",
        "summary": str(generated.get("summary") or content_card.get("summary") or "").strip(),
        "preview": str(content_card.get("content_preview") or "").strip(),
        "output_format": str(generated.get("output_format") or absolute_path.suffix.lstrip(".")).strip().lower(),
        "absolute_path": absolute_path,
        "extra_source_ids": list(generated.get("source_ids") or []),
        "version_no": int(generated.get("version_no") or 1),
        "media_info": media_info,
    }


def resolve_attachment_style_source(
    service: Any,
    *,
    profile_user_id: str,
    session_id: str,
    target: str,
) -> dict[str, Any] | None:
    attachment = service.attachment_service.resolve_attachment(
        profile_user_id=profile_user_id,
        session_id=session_id,
        target=target,
        kind="any",
    )
    if attachment is None:
        return None
    source_path = None
    if hasattr(service.attachment_service, "resolve_storage_path"):
        try:
            source_path = service.attachment_service.resolve_storage_path(attachment)  # type: ignore[attr-defined]
        except Exception:
            source_path = None
    detail = attachment.get("detail") if isinstance(attachment.get("detail"), dict) else {}
    media_info = detail.get("media_info") if isinstance(detail.get("media_info"), dict) else {}
    title = (
        str(attachment.get("summary_title") or "").strip()
        or str(attachment.get("origin_name") or "").strip()
        or str(attachment.get("attachment_handle") or "").strip()
        or "工作台材料"
    )
    output_format = ""
    if isinstance(source_path, Path):
        output_format = source_path.suffix.lower().lstrip(".")
    if not output_format:
        output_format = str(attachment.get("file_ext") or detail.get("file_kind") or "").strip().lower().lstrip(".")
    return {
        "source_type": "attachment",
        "source_id": str(attachment.get("attachment_id") or "").strip(),
        "handle": str(attachment.get("attachment_handle") or "").strip(),
        "title": title,
        "summary": str(attachment.get("short_hint") or detail.get("summary") or "").strip(),
        "preview": str(detail.get("text_preview") or detail.get("content_preview") or "").strip(),
        "output_format": output_format,
        "absolute_path": source_path,
        "extra_source_ids": [],
        "version_no": 0,
        "media_info": media_info,
    }


def sendable_file_label(service: Any, file_ref: dict[str, Any]) -> str:
    source_type = str(file_ref.get("source_type") or "").strip()
    handle = str(file_ref.get("handle") or "").strip()
    title = str(file_ref.get("title") or file_ref.get("name") or "").strip()
    ext = str(file_ref.get("file_ext") or "").strip()
    type_label = "生成文件" if source_type == "generated" else "工作台材料" if source_type == "attachment" else "文件"
    label = f"{type_label} {handle}" if handle else type_label
    if title:
        label += f"《{title}》"
    if ext:
        label += f"（{ext}）"
    return label


def normalize_send_targets(
    service: Any,
    *,
    target: str = "latest",
    targets: list[str] | tuple[str, ...] | None = None,
) -> list[str]:
    raw_items: list[Any] = []
    if targets:
        raw_items.extend(list(targets))
    elif target:
        raw_items.append(target)
    else:
        raw_items.append("latest")

    normalized: list[str] = []
    for item in raw_items:
        if isinstance(item, str):
            chunks = item.replace("，", ",").replace("、", ",").split(",")
        else:
            chunks = [str(item or "")]
        for chunk in chunks:
            text = str(chunk or "").strip()
            if not text:
                continue
            if text not in normalized:
                normalized.append(text[:120])
    return normalized or ["latest"]
