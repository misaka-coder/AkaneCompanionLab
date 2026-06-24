from __future__ import annotations

import csv
import importlib.util
import time
import re
from pathlib import Path
from typing import Any, Callable

from .store import MemoryStore


ACTIVE_ATTACHMENT_STATUSES = ("ready", "pending_observation", "failed")
READY_ATTACHMENT_STATUSES = ("ready",)
AUTO_FOCUS_BATCH_SECONDS = 90
DEFAULT_WORKSPACE_CHAR_BUDGET = 24000
WORKSPACE_ITEM_CHAR_BUDGET = 12000
WORKSPACE_MAX_TARGETS = 30
AUTO_FOCUS_MAX_ITEMS = WORKSPACE_MAX_TARGETS


class AttachmentInboxService:
    """Session-scoped material workspace for user-provided files.

    Items here are short-lived source materials from clients such as QQ,
    desktop pet drag/drop, or URL imports. They are not gifts, character
    resources, generated outputs, or long-term memory.
    """

    def __init__(
        self,
        *,
        store: MemoryStore,
        base_dir: Path | None = None,
        legacy_base_dirs: list[Path] | tuple[Path, ...] | None = None,
        workspace_uri_resolver: Callable[[str], Path | None] | None = None,
    ) -> None:
        self.store = store
        self.base_dir = Path(base_dir) if base_dir is not None else None
        self.legacy_base_dirs = [
            Path(item)
            for item in list(legacy_base_dirs or [])
            if self.base_dir is None or Path(item) != self.base_dir
        ]
        self.workspace_uri_resolver = workspace_uri_resolver

    def create_pending(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        source: str,
        kind: str,
        origin_name: str = "",
        mime_type: str = "",
        file_ext: str = "",
        file_size: int = 0,
        storage_relpath: str = "",
        source_event_id: str = "",
        source_message_id: str = "",
        timestamp: int | None = None,
    ) -> dict[str, Any]:
        return self.store.add_attachment_inbox_item(
            profile_user_id=profile_user_id,
            session_id=session_id,
            source=source,
            kind=kind,
            status="pending_observation",
            origin_name=origin_name,
            mime_type=mime_type,
            file_ext=file_ext,
            file_size=file_size,
            storage_relpath=storage_relpath,
            source_event_id=source_event_id,
            source_message_id=source_message_id,
            timestamp=timestamp,
        )

    def mark_ready(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        attachment_id: str,
        summary_title: str = "",
        short_hint: str = "",
        detail: dict[str, Any] | None = None,
        timestamp: int | None = None,
    ) -> dict[str, Any] | None:
        return self.store.mark_attachment_inbox_item_ready(
            profile_user_id=profile_user_id,
            session_id=session_id,
            attachment_id=attachment_id,
            summary_title=summary_title,
            short_hint=short_hint,
            detail=detail if isinstance(detail, dict) else {},
            focus_batch_seconds=AUTO_FOCUS_BATCH_SECONDS,
            focus_max_items=AUTO_FOCUS_MAX_ITEMS,
            timestamp=timestamp,
        )

    def wait_for_attachments_settled(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        attachment_ids: list[str],
        timeout_seconds: float = 8.0,
        poll_interval_seconds: float = 0.2,
    ) -> dict[str, Any]:
        normalized_ids = [str(item or "").strip() for item in attachment_ids or [] if str(item or "").strip()]
        if not normalized_ids:
            return {"ok": True, "ready": [], "failed": [], "pending": [], "missing": []}

        deadline = time.time() + max(0.0, float(timeout_seconds or 0.0))
        poll_interval = min(1.0, max(0.05, float(poll_interval_seconds or 0.2)))
        latest: dict[str, dict[str, Any]] = {}
        while True:
            latest = self._load_attachment_items_by_id(
                profile_user_id=profile_user_id,
                session_id=session_id,
                attachment_ids=normalized_ids,
            )
            pending = [
                item_id
                for item_id in normalized_ids
                if str((latest.get(item_id) or {}).get("status") or "missing") == "pending_observation"
            ]
            if not pending or time.time() >= deadline:
                break
            time.sleep(min(poll_interval, max(0.0, deadline - time.time())))

        ready: list[str] = []
        failed: list[str] = []
        pending_final: list[str] = []
        missing: list[str] = []
        for item_id in normalized_ids:
            status = str((latest.get(item_id) or {}).get("status") or "missing")
            if status == "ready":
                ready.append(item_id)
            elif status == "failed":
                failed.append(item_id)
            elif status == "pending_observation":
                pending_final.append(item_id)
            else:
                missing.append(item_id)
        return {
            "ok": not pending_final,
            "ready": ready,
            "failed": failed,
            "pending": pending_final,
            "missing": missing,
        }

    def build_prompt_context(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        detail_limit: int = 2,
        index_limit: int = 12,
        pending_limit: int = 3,
    ) -> str:
        items = self.store.list_attachment_inbox_items(
            profile_user_id=profile_user_id,
            session_id=session_id,
            statuses=list(ACTIVE_ATTACHMENT_STATUSES),
            limit=max(60, detail_limit + index_limit + pending_limit + 12),
        )
        focused = sorted(
            [item for item in items if int(item.get("focus_rank") or 0) > 0],
            key=lambda item: int(item.get("focus_rank") or 0),
        )
        ready = [item for item in items if item.get("status") == "ready"]
        pending = [item for item in items if item.get("status") == "pending_observation"]
        ready_remote_sources = {
            str(item.get("source_event_id") or "").strip()
            for item in ready
            if str(item.get("source") or "").strip() == "remote_url"
            and str(item.get("source_event_id") or "").strip()
        }
        failed = [
            item
            for item in items
            if item.get("status") == "failed"
            and not (
                str(item.get("source") or "").strip() == "remote_url"
                and str(item.get("source_event_id") or "").strip() in ready_remote_sources
            )
        ]

        if not ready and not pending and not failed:
            return ""

        lines = [
            "【当前材料工作台】",
            "这些是当前会话临时放进工作台的原始材料，可能来自 QQ、桌宠本地拖拽、链接下载或其他客户端；只用于近几轮处理/讨论。",
            "不要把这些材料当成礼物、角色资源、生成成果或长期记忆。原始材料和工具生成成果是两类对象：要修改、转写、转换或整理材料时，先使用对应工具生成新成果；用户要拿到某个已有文件时，再按当前客户端可用的交付方式发送或打开它。",
        ]

        detailed_items = focused
        detailed_ids = {str(item.get("attachment_id") or "") for item in detailed_items}
        indexed_ready = [
            item
            for item in ready
            if str(item.get("attachment_id") or "") not in detailed_ids
        ][: max(0, int(index_limit or 6))]
        if detailed_items:
            lines.append("当前重点材料 Focus（这些材料会持续放在你眼前，直到重新同步或清理）：")
            rendered, budget_overflow = self._render_focus_items_with_budget(
                detailed_items,
                total_budget=DEFAULT_WORKSPACE_CHAR_BUDGET,
            )
            lines.extend(rendered)
            if budget_overflow:
                lines.append(
                    "工作台空间暂时不够完整展开："
                    + "、".join(self._compact_item_label(item) for item in budget_overflow[:8])
                    + "。这些仍在 Manifest，可用 read_attachment_section 指定页/行/sheet 展开。"
                )

        if indexed_ready:
            lines.append("旁边的材料清单 Manifest（只用于识别和选择，不代表已经阅读全文/全片）：")
            for index, item in enumerate(indexed_ready, start=1):
                lines.extend(self._render_manifest_item(index, item))

        if pending:
            lines.append("正在处理、还没有完整描述的材料：")
            for item in pending[: max(1, int(pending_limit or 3))]:
                lines.append(f"- {self._compact_item_label(item)}")

        if failed:
            lines.append("处理失败的材料：")
            for item in failed[:2]:
                label = self._compact_item_label(item)
                error = str(item.get("error_message") or "").strip()
                lines.append(f"- {label}" + (f"：{error}" if error else ""))

        overflow = max(0, len(ready) - len([item for item in detailed_items if item.get("status") == "ready"]) - len(indexed_ready))
        if overflow:
            lines.append(f"此外还有 {overflow} 个较早材料未展开。")
        lines.append(
            "需要收起暂时不分析的材料、重新指定重点材料，或切换对比对象时，使用 sync_attachment_workspace 整理当前材料工作台；"
            "聊完或用户说不用了，可用 clear_attachment_focus 清理。"
        )
        return "\n".join(lines)

    def sync_workspace(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        focus_targets: list[Any] | tuple[Any, ...] | set[Any] | str | None,
        kind: str = "any",
        reason: str = "",
        timestamp: int | None = None,
        max_focus: int = WORKSPACE_MAX_TARGETS,
    ) -> dict[str, Any]:
        effective_ts = int(timestamp or time.time())
        targets = self._normalize_targets(focus_targets)
        resolved: list[dict[str, Any]] = []
        unresolved: list[str] = []
        ambiguous_targets: list[str] = []
        for target in targets:
            resolution = self._resolve_target_result(
                profile_user_id=profile_user_id,
                session_id=session_id,
                target=target,
                kind=kind,
            )
            item = resolution.get("item") if isinstance(resolution, dict) else None
            if item is None:
                ambiguity = self._format_ambiguity_target(
                    target=target,
                    items=list(resolution.get("ambiguous_matches") or []) if isinstance(resolution, dict) else [],
                )
                if ambiguity:
                    ambiguous_targets.append(ambiguity)
                else:
                    unresolved.append(target)
                continue
            item_id = str(item.get("attachment_id") or "").strip()
            if item_id and item_id not in {str(existing.get("attachment_id") or "") for existing in resolved}:
                resolved.append(item)

        target_limit = max(1, min(WORKSPACE_MAX_TARGETS, int(max_focus or WORKSPACE_MAX_TARGETS)))
        overflow = resolved[target_limit:]
        selected = resolved[:target_limit]
        selected_ids = [str(item.get("attachment_id") or "") for item in selected if str(item.get("attachment_id") or "")]
        synced = self.store.sync_attachment_workspace_focus(
            profile_user_id=profile_user_id,
            session_id=session_id,
            attachment_ids=selected_ids,
            timestamp=effective_ts,
        )
        return {
            "ok": bool(synced or not targets),
            "focused": synced,
            "unresolved": unresolved,
            "ambiguous_targets": ambiguous_targets,
            "overflow": overflow,
            "followup_context": self._build_workspace_followup(
                focused=synced,
                unresolved=unresolved,
                ambiguous_targets=ambiguous_targets,
                overflow=overflow,
                reason=reason,
            ),
        }

    def inspect_attachment(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        target: str = "",
        kind: str = "any",
        timestamp: int | None = None,
    ) -> dict[str, Any]:
        resolution = self._resolve_target_result(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=target,
            kind=kind,
        )
        item = resolution.get("item") if isinstance(resolution, dict) else None
        if item is None:
            ambiguity = self._format_ambiguity_target(
                target=str(target or "").strip(),
                items=list(resolution.get("ambiguous_matches") or []) if isinstance(resolution, dict) else [],
            )
            return {
                "ok": False,
                "item": None,
                "ambiguous_targets": [ambiguity] if ambiguity else [],
                "followup_context": (
                    "你刚刚想查看某个附件，但当前没有找到明确匹配的图片或文件。请自然向用户确认是哪一个。"
                    if not ambiguity
                    else (
                        "你刚刚想查看某个附件，但当前有多个候选，请让用户确认："
                        f"{ambiguity}。不要自己替用户决定。"
                    )
                ),
            }

        touched = self.store.update_attachment_inbox_item(
            profile_user_id=profile_user_id,
            session_id=session_id,
            attachment_id=str(item.get("attachment_id") or ""),
            last_used_at=int(timestamp or time.time()),
            updated_at=int(timestamp or time.time()),
        ) or item
        return {
            "ok": True,
            "item": touched,
            "followup_context": self._build_inspect_followup(touched),
        }

    def read_section(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        target: str = "",
        section: str = "",
        kind: str = "any",
        timestamp: int | None = None,
    ) -> dict[str, Any]:
        resolution = self._resolve_target_result(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=target,
            kind=kind,
        )
        item = resolution.get("item") if isinstance(resolution, dict) else None
        if item is None:
            ambiguity = self._format_ambiguity_target(
                target=str(target or "").strip(),
                items=list(resolution.get("ambiguous_matches") or []) if isinstance(resolution, dict) else [],
            )
            return {
                "ok": False,
                "item": None,
                "content": "",
                "ambiguous_targets": [ambiguity] if ambiguity else [],
                "followup_context": (
                    "你刚刚想展开读取附件的一部分，但当前没有找到明确匹配的图片或文件。请自然向用户确认是哪一个。"
                    if not ambiguity
                    else (
                        "你刚刚想展开读取附件的一部分，但当前有多个候选，请让用户确认："
                        f"{ambiguity}。不要自己替用户决定。"
                    )
                ),
            }

        touched = self.store.update_attachment_inbox_item(
            profile_user_id=profile_user_id,
            session_id=session_id,
            attachment_id=str(item.get("attachment_id") or ""),
            last_used_at=int(timestamp or time.time()),
            updated_at=int(timestamp or time.time()),
        ) or item
        content = self._extract_section_content(touched, section=section)
        if not content:
            return {
                "ok": False,
                "item": touched,
                "content": "",
                "followup_context": (
                    "你刚刚想展开读取附件内容，但这个附件目前没有可展开的文本预览。"
                    "如果它是扫描 PDF、图片型文件或未支持格式，请自然告诉用户暂时只能看到摘要。"
                ),
            }
        label = self._compact_item_label(touched)
        section_text = str(section or "当前可用片段").strip()
        return {
            "ok": True,
            "item": touched,
            "content": content,
            "followup_context": (
                f"你刚刚展开读取了工作台材料 {label} 的「{section_text}」。\n"
                f"展开内容如下：\n{content}\n"
                "请基于这段展开内容自然回应；不要把材料全文默认写入长期记忆。"
            ),
        }

    def resolve_attachment(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        target: str = "",
        kind: str = "any",
    ) -> dict[str, Any] | None:
        return self._resolve_target(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=target,
            kind=kind,
        )

    def find_attachment_matches(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        target: str,
        kind: str = "any",
        statuses: list[str] | tuple[str, ...] | set[str] | None = None,
        limit: int = 8,
    ) -> list[dict[str, Any]]:
        items = self.store.find_attachment_inbox_item_matches(
            profile_user_id=profile_user_id,
            session_id=session_id,
            query=str(target or "").strip(),
            kind=kind,
            statuses=list(statuses) if statuses is not None else list(ACTIVE_ATTACHMENT_STATUSES),
            limit=limit,
        )
        return [{key: value for key, value in item.items() if key != "_match_rank"} for item in items]

    def format_attachment_ambiguity(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        target: str,
        kind: str = "any",
        statuses: list[str] | tuple[str, ...] | set[str] | None = None,
    ) -> str:
        resolution = self._resolve_target_result(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=target,
            kind=kind,
            statuses=list(statuses) if statuses is not None else list(ACTIVE_ATTACHMENT_STATUSES),
        )
        return self._format_ambiguity_target(
            target=str(target or "").strip(),
            items=list(resolution.get("ambiguous_matches") or []) if isinstance(resolution, dict) else [],
        )

    def resolve_storage_path(self, item: dict[str, Any]) -> Path | None:
        """Return the original stored file path for an inbox item, if safe.

        Generated-file tooling needs a public way to work from the original
        attachment bytes without reaching into the private resolver.
        """

        return self._resolve_storage_path(item)

    def clear_focus(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        target: str = "current",
        targets: list[Any] | tuple[Any, ...] | set[Any] | str | None = None,
        kind: str = "any",
        reason: str = "",
        timestamp: int | None = None,
    ) -> dict[str, Any]:
        effective_ts = int(timestamp or time.time())
        normalized_targets = self._normalize_targets(targets)
        unresolved: list[str] = []
        ambiguous_targets: list[str] = []
        if normalized_targets:
            lowered_targets = {item.lower() for item in normalized_targets}
            if lowered_targets & {"all", "全部", "*"}:
                cleared = self.store.clear_attachment_inbox_items(
                    profile_user_id=profile_user_id,
                    session_id=session_id,
                    target="all",
                    kind=kind,
                    timestamp=effective_ts,
                )
            else:
                resolved: list[dict[str, Any]] = []
                seen_ids: set[str] = set()
                for item_target in normalized_targets:
                    resolution = self._resolve_target_result(
                        profile_user_id=profile_user_id,
                        session_id=session_id,
                        target=item_target,
                        kind=kind,
                    )
                    item = resolution.get("item") if isinstance(resolution, dict) else None
                    if item is None:
                        ambiguity = self._format_ambiguity_target(
                            target=item_target,
                            items=list(resolution.get("ambiguous_matches") or []) if isinstance(resolution, dict) else [],
                        )
                        if ambiguity:
                            ambiguous_targets.append(ambiguity)
                        else:
                            unresolved.append(item_target)
                        continue
                    item_id = str(item.get("attachment_id") or "").strip()
                    if item_id and item_id not in seen_ids:
                        seen_ids.add(item_id)
                        resolved.append(item)

                cleared = []
                for item in resolved:
                    updated = self.store.update_attachment_inbox_item(
                        profile_user_id=profile_user_id,
                        session_id=session_id,
                        attachment_id=str(item.get("attachment_id") or ""),
                        status="cleared",
                        updated_at=effective_ts,
                    )
                    cleared.append(updated or dict(item, status="cleared", updated_at=effective_ts))
        else:
            cleared = self.store.clear_attachment_inbox_items(
                profile_user_id=profile_user_id,
                session_id=session_id,
                target=target,
                kind=kind,
                timestamp=effective_ts,
            )
        if not cleared:
            missing = f"没有找到这些目标：{', '.join(unresolved[:5])}。" if unresolved else ""
            return {
                "ok": False,
                "cleared": [],
                "unresolved": unresolved,
                "ambiguous_targets": ambiguous_targets,
                "followup_context": (
                    "你刚刚想移除工作台材料焦点，但没有找到可移除的材料。"
                    + (missing if missing else "")
                    + (
                        "这些目标不够明确，存在多个候选，请让用户确认："
                        + "；".join(ambiguous_targets[:5])
                        + "。"
                        if ambiguous_targets
                        else ""
                    )
                    + "请自然继续对话，不要重复调用工具。"
                ),
            }

        names = "、".join(self._display_name(item) for item in cleared[:3])
        if len(cleared) > 3:
            names += f" 等 {len(cleared)} 个"
        missing = f"未找到：{', '.join(unresolved[:5])}。" if unresolved else ""
        suffix = f"原因：{reason}。" if str(reason or "").strip() else ""
        return {
            "ok": True,
            "cleared": cleared,
            "unresolved": unresolved,
            "ambiguous_targets": ambiguous_targets,
            "followup_context": (
                f"你刚刚已经从工作台材料焦点中移除了 {len(cleared)} 个材料"
                f"（{names}）。{missing}"
                + (
                    "这些目标不够明确，存在多个候选，请让用户确认："
                    + "；".join(ambiguous_targets[:5])
                    + "。"
                    if ambiguous_targets
                    else ""
                )
                + f"{suffix}这些材料不会继续注入上下文，也不会作为礼物、角色资源或长期记忆保存。"
                "请自然继续回应，不要重复调用工具。"
            ),
        }

    def _resolve_target(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        target: str,
        kind: str,
    ) -> dict[str, Any] | None:
        result = self._resolve_target_result(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=target,
            kind=kind,
        )
        item = result.get("item") if isinstance(result, dict) else None
        return item if isinstance(item, dict) else None

    def _resolve_target_result(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        target: str,
        kind: str,
        statuses: list[str] | tuple[str, ...] | set[str] | None = None,
    ) -> dict[str, Any]:
        normalized_target = str(target or "").strip()
        lowered_target = normalized_target.lower()
        normalized_statuses = list(statuses) if statuses is not None else list(ACTIVE_ATTACHMENT_STATUSES)
        items = self.store.list_attachment_inbox_items(
            profile_user_id=profile_user_id,
            session_id=session_id,
            statuses=normalized_statuses,
            limit=200,
        )
        if not normalized_target or lowered_target in {"current", "latest", "最近", "当前"}:
            kind_filtered = self._filter_items_by_target_kind(items, normalized_target, kind)
            item = self._latest_item(kind_filtered)
            return {"item": item, "ambiguous_matches": []}

        matches = self._find_target_matches_from_items(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=normalized_target,
            items=items,
            kind=kind,
            statuses=normalized_statuses,
        )
        if not matches:
            return {"item": None, "ambiguous_matches": []}
        if len(matches) > 1:
            return {"item": None, "ambiguous_matches": matches}
        return {"item": matches[0], "ambiguous_matches": []}

    def _load_attachment_items_by_id(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        attachment_ids: list[str],
    ) -> dict[str, dict[str, Any]]:
        wanted = {str(item_id or "").strip() for item_id in attachment_ids or [] if str(item_id or "").strip()}
        if not wanted:
            return {}
        items = self.store.list_attachment_inbox_items(
            profile_user_id=profile_user_id,
            session_id=session_id,
            statuses=list(ACTIVE_ATTACHMENT_STATUSES),
            limit=max(80, len(wanted) + 20),
        )
        return {
            str(item.get("attachment_id") or "").strip(): item
            for item in items
            if str(item.get("attachment_id") or "").strip() in wanted
        }

    def _resolve_target_from_items(
        self,
        *,
        target: str,
        items: list[dict[str, Any]],
        kind: str,
    ) -> dict[str, Any] | None:
        matches = self._find_target_matches_from_items_local(
            target=target,
            items=items,
            kind=kind,
        )
        return matches[0] if len(matches) == 1 else None

    def _find_target_matches_from_items(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        target: str,
        items: list[dict[str, Any]],
        kind: str,
        statuses: list[str] | tuple[str, ...] | set[str] | None,
    ) -> list[dict[str, Any]]:
        local_matches = self._find_target_matches_from_items_local(
            target=target,
            items=items,
            kind=kind,
        )
        if local_matches:
            return local_matches

        raw_matches = self.store.find_attachment_inbox_item_matches(
            profile_user_id=profile_user_id,
            session_id=session_id,
            query=str(target or "").strip(),
            kind=kind,
            statuses=statuses,
            limit=8,
        )
        if not raw_matches:
            return []
        item_by_id = {
            str(item.get("attachment_id") or "").strip(): item
            for item in items
            if str(item.get("attachment_id") or "").strip()
        }
        matches: list[dict[str, Any]] = []
        for match in raw_matches:
            item_id = str(match.get("attachment_id") or "").strip()
            resolved = item_by_id.get(item_id)
            if resolved is not None:
                matches.append(dict(resolved, _match_rank=match.get("_match_rank")))
        if not matches:
            return []
        best_rank = int(matches[0].get("_match_rank") or 99)
        top_matches = [
            {key: value for key, value in item.items() if key != "_match_rank"}
            for item in matches
            if int(item.get("_match_rank") or 99) == best_rank
        ]
        return self._dedupe_attachment_items(top_matches)

    def _find_target_matches_from_items_local(
        self,
        *,
        target: str,
        items: list[dict[str, Any]],
        kind: str,
    ) -> list[dict[str, Any]]:
        normalized = str(target or "").strip()
        if not normalized:
            return []
        lowered = normalized.lower()
        kind_filtered = self._filter_items_by_target_kind(items, normalized, kind)

        exact_matches = [
            item
            for item in items
            if normalized == str(item.get("attachment_id") or "").strip()
            or lowered == str(item.get("attachment_handle") or "").strip().lower()
        ]
        if exact_matches:
            return self._dedupe_attachment_items(exact_matches[:1])

        if lowered in {"current", "latest", "最近", "当前", "最后一张", "最后一个"}:
            latest = self._latest_item(kind_filtered)
            return [latest] if latest is not None else []
        if lowered in {"oldest", "first", "最早", "第一张", "第一个"}:
            oldest = self._oldest_item(kind_filtered)
            return [oldest] if oldest is not None else []

        reverse_index = self._parse_reverse_ordinal(normalized)
        if reverse_index is not None:
            ordered = self._sort_by_sequence(kind_filtered)
            if 1 <= reverse_index <= len(ordered):
                return [ordered[-reverse_index]]
            return []

        sequence_no = self._parse_sequence_no(normalized) if self._looks_like_sequence_reference(normalized) else None
        if sequence_no is not None:
            sequence_matches = [
                item
                for item in kind_filtered
                if int(item.get("sequence_no") or 0) == sequence_no
            ]
            return self._dedupe_attachment_items(sequence_matches)
        return []

    def _auto_focus_recent_batch(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        anchor: dict[str, Any],
        timestamp: int | None = None,
    ) -> None:
        effective_ts = int(timestamp or anchor.get("updated_at") or time.time())
        threshold = effective_ts - AUTO_FOCUS_BATCH_SECONDS
        candidates = self.store.list_attachment_inbox_items(
            profile_user_id=profile_user_id,
            session_id=session_id,
            statuses=list(READY_ATTACHMENT_STATUSES),
            limit=80,
        )
        recent = [
            item
            for item in candidates
            if int(item.get("updated_at") or item.get("created_at") or 0) >= threshold
        ]
        anchor_id = str(anchor.get("attachment_id") or "").strip()
        if anchor_id and anchor_id not in {str(item.get("attachment_id") or "") for item in recent}:
            recent.append(anchor)
        recent = sorted(
            recent,
            key=lambda item: (
                int(item.get("created_at") or 0),
                int(item.get("sequence_no") or 0),
                int(item.get("updated_at") or 0),
            ),
        )[-AUTO_FOCUS_MAX_ITEMS:]
        attachment_ids = [
            str(item.get("attachment_id") or "").strip()
            for item in recent
            if str(item.get("attachment_id") or "").strip()
        ]
        if attachment_ids:
            self.store.sync_attachment_workspace_focus(
                profile_user_id=profile_user_id,
                session_id=session_id,
                attachment_ids=attachment_ids,
                timestamp=effective_ts,
            )

    def _render_focus_items_with_budget(
        self,
        items: list[dict[str, Any]],
        *,
        total_budget: int,
    ) -> tuple[list[str], list[dict[str, Any]]]:
        remaining = max(1200, int(total_budget or DEFAULT_WORKSPACE_CHAR_BUDGET))
        lines: list[str] = []
        overflow: list[dict[str, Any]] = []
        for index, item in enumerate(items, start=1):
            if remaining < 800 and lines:
                overflow.append(item)
                continue
            item_budget = min(WORKSPACE_ITEM_CHAR_BUDGET, max(800, remaining))
            item_lines = self._render_detail_item(index, item, char_budget=item_budget)
            rendered = "\n".join(item_lines)
            cost = len(rendered)
            if cost > remaining and lines:
                overflow.append(item)
                continue
            lines.extend(item_lines)
            remaining -= min(cost, remaining)
        return lines, overflow

    def _render_detail_item(self, index: int, item: dict[str, Any], *, char_budget: int = WORKSPACE_ITEM_CHAR_BUDGET) -> list[str]:
        detail = item.get("detail") if isinstance(item.get("detail"), dict) else {}
        lines = [f"{index}. {self._compact_item_label(item)}"]
        if item.get("status") == "pending_observation":
            lines.append("   状态：系统正在生成描述，还没有完整内容。")
            return lines
        if item.get("status") == "failed":
            error = str(item.get("error_message") or "").strip()
            readable_error = self._readable_failure_message(item) or error
            lines.append("   状态：处理失败。" + (f"原因：{readable_error}" if readable_error else ""))
            return lines
        lines.append("   状态：已放在当前工作台，可直接基于下面内容讨论。")
        source_label = self._source_label(item)
        if source_label:
            lines.append(f"   来源：{source_label}")
        kind = str(item.get("kind") or "").strip().lower()
        if kind == "image":
            lines.extend(self._render_image_focus_detail(detail, fallback_hint=str(item.get("short_hint") or "")))
            return self._clip_rendered_lines(lines, char_budget)

        media_info = detail.get("media_info") if isinstance(detail.get("media_info"), dict) else {}
        if media_info:
            lines.extend(self._render_media_info_lines(media_info))
            remote_source = detail.get("remote_source") if isinstance(detail.get("remote_source"), dict) else {}
            if remote_source:
                lines.extend(self._render_remote_source_lines(remote_source))
            lines.append("   说明：这是轻量媒体规格卡；如需转码、截取、调音量、淡入淡出、调速，可用 convert_media_file；如需降噪、去混响或净化人声，可用 clean_voice_track；如需转写文字稿/字幕，可用 transcribe_media；如需整理训练素材，可用 prepare_voice_dataset。")
            return self._clip_rendered_lines(lines, char_budget)

        tags = self._normalize_text_list(
            detail.get("mood_tags")
            or detail.get("tags")
            or detail.get("keywords")
            or []
        )
        if tags:
            lines.append(f"   标签：{', '.join(tags[:8])}")

        entities = self._normalize_text_list(detail.get("entities") or detail.get("objects") or [])
        if entities:
            lines.append(f"   要素：{', '.join(entities[:8])}")

        content_budget = max(600, char_budget - len("\n".join(lines)) - 160)
        content = self._extract_original_file_section(item, section="全文", max_chars=content_budget)
        if not content:
            content = str(detail.get("text_preview") or detail.get("content_preview") or "").strip()[:content_budget]
        if content:
            lines.append("   内容：")
            lines.append(self._indent_block(content, prefix="   "))
            if len(content) >= content_budget - 20:
                lines.append("   注意：当前工作台内容已接近预算边界；如需更具体位置，用 read_attachment_section 指定页、行或 sheet 展开。")
        else:
            hint = str(item.get("short_hint") or detail.get("summary") or detail.get("description") or "").strip()
            if hint:
                lines.append(f"   说明：{hint[:360]}")
        return self._clip_rendered_lines(lines, char_budget)

    def _render_image_focus_detail(self, detail: dict[str, Any], *, fallback_hint: str = "") -> list[str]:
        lines: list[str] = []
        title = str(detail.get("summary_title") or "").strip()
        if title:
            lines.append(f"   临时标题：{title}")
        description_parts = []
        for key in ("summary", "description", "scene_description", "appearance", "composition"):
            text = str(detail.get(key) or "").strip()
            if text and text not in description_parts:
                description_parts.append(text)
        fallback = str(fallback_hint or "").strip()
        if fallback and fallback not in description_parts:
            description_parts.append(fallback)
        if description_parts:
            lines.append("   视觉描述：" + "\n   ".join(description_parts))
        entities = self._normalize_text_list(detail.get("entities") or detail.get("objects") or [])
        if entities:
            lines.append(f"   要素：{', '.join(entities[:24])}")
        tags = self._normalize_text_list(detail.get("mood_tags") or detail.get("tags") or detail.get("keywords") or [])
        if tags:
            lines.append(f"   标签：{', '.join(tags[:24])}")
        uncertainty = str(detail.get("uncertainty") or "").strip()
        if uncertainty:
            lines.append(f"   不确定处：{uncertainty[:500]}")
        extra_keys = [
            key
            for key in sorted(detail.keys())
            if key
            not in {
                "summary_title",
                "summary",
                "description",
                "scene_description",
                "appearance",
                "composition",
                "entities",
                "objects",
                "mood_tags",
                "tags",
                "keywords",
                "uncertainty",
            }
        ]
        for key in extra_keys[:8]:
            value = detail.get(key)
            if isinstance(value, (str, int, float, bool)) and str(value).strip():
                lines.append(f"   {key}: {str(value).strip()[:500]}")
            elif isinstance(value, list) and value:
                values = self._normalize_text_list(value)
                if values:
                    lines.append(f"   {key}: {', '.join(values[:16])}")
        return lines or ["   视觉描述：当前图片已在工作台，但观察卡没有提供更多细节。"]

    def _render_media_info_lines(self, media_info: dict[str, Any], *, prefix: str = "   ") -> list[str]:
        lines: list[str] = []
        format_name = str(media_info.get("format_name") or "").strip()
        duration = self._format_duration_label(media_info.get("duration_seconds"))
        file_size = self._format_size(int(media_info.get("file_size") or 0)) if media_info.get("file_size") else ""
        basics = []
        if format_name:
            basics.append(f"容器/格式：{format_name}")
        if duration:
            basics.append(f"时长：{duration}")
        if file_size:
            basics.append(f"大小：{file_size}")
        if basics:
            lines.append(prefix + "媒体信息：" + "；".join(basics))
        audio = media_info.get("audio") if isinstance(media_info.get("audio"), dict) else {}
        if audio:
            audio_bits = [
                f"编码 {audio.get('codec') or '未知'}",
                f"{audio.get('sample_rate') or '未知'}Hz",
                f"{audio.get('channels') or '未知'}声道",
            ]
            if audio.get("bit_rate"):
                audio_bits.append(self._format_bitrate(int(audio.get("bit_rate") or 0)))
            lines.append(prefix + "音频：" + "，".join(bit for bit in audio_bits if bit) + "。")
        else:
            lines.append(prefix + "音频：未检测到音轨。")
        video = media_info.get("video") if isinstance(media_info.get("video"), dict) else {}
        if video:
            video_bits = [
                f"编码 {video.get('codec') or '未知'}",
                f"{video.get('width') or '?'}x{video.get('height') or '?'}",
            ]
            if isinstance(video.get("fps"), (int, float)) and video.get("fps") > 0:
                video_bits.append(f"{video.get('fps'):g}fps")
            lines.append(prefix + "视频：" + "，".join(video_bits) + "。")
        return lines

    def _render_remote_source_lines(self, remote_source: dict[str, Any], *, prefix: str = "   ") -> list[str]:
        platform = str(remote_source.get("platform") or remote_source.get("extractor") or "").strip()
        uploader = str(remote_source.get("uploader") or "").strip()
        webpage_url = str(remote_source.get("webpage_url") or remote_source.get("source_url") or "").strip()
        pieces = []
        if platform:
            pieces.append(f"平台 {platform}")
        if uploader:
            pieces.append(f"发布者 {uploader}")
        if webpage_url:
            pieces.append(f"链接 {webpage_url[:160]}")
        if not pieces:
            return []
        return [prefix + "来源：" + "；".join(pieces)]

    def _render_manifest_item(self, index: int, item: dict[str, Any]) -> list[str]:
        detail = item.get("detail") if isinstance(item.get("detail"), dict) else {}
        parts = [
            f"{index}. {self._compact_item_label(item)}",
        ]
        meta = []
        file_kind = str(detail.get("file_kind") or item.get("file_ext") or "").strip()
        if file_kind:
            meta.append(f"格式={file_kind}")
        file_size = int(item.get("file_size") or detail.get("file_size") or 0)
        if file_size:
            meta.append(f"大小={self._format_size(file_size)}")
        source_label = self._source_label(item)
        if source_label:
            meta.append(f"来源={source_label}")
        if detail.get("line_count") is not None:
            meta.append(f"行数={detail.get('line_count')}")
        if detail.get("page_count") is not None:
            meta.append(f"页数={detail.get('page_count')}")
        if detail.get("paragraph_count") is not None:
            meta.append(f"段落={detail.get('paragraph_count')}")
        media_info = detail.get("media_info") if isinstance(detail.get("media_info"), dict) else {}
        if media_info:
            duration = self._format_duration_label(media_info.get("duration_seconds"))
            if duration:
                meta.append(f"时长={duration}")
            audio = media_info.get("audio") if isinstance(media_info.get("audio"), dict) else {}
            if audio:
                audio_bits = [
                    str(audio.get("codec") or "").strip(),
                    f"{audio.get('sample_rate')}Hz" if audio.get("sample_rate") else "",
                    f"{audio.get('channels')}声道" if audio.get("channels") else "",
                ]
                meta.append("音频=" + ",".join(bit for bit in audio_bits if bit))
            video = media_info.get("video") if isinstance(media_info.get("video"), dict) else {}
            if video:
                video_bits = [
                    str(video.get("codec") or "").strip(),
                    f"{video.get('width')}x{video.get('height')}" if video.get("width") and video.get("height") else "",
                    f"{video.get('fps'):g}fps" if isinstance(video.get("fps"), (int, float)) else "",
                ]
                meta.append("视频=" + ",".join(bit for bit in video_bits if bit))
        remote_source = detail.get("remote_source") if isinstance(detail.get("remote_source"), dict) else {}
        if remote_source:
            platform = str(remote_source.get("platform") or remote_source.get("extractor") or "").strip()
            uploader = str(remote_source.get("uploader") or "").strip()
            if platform:
                meta.append(f"来源={platform}")
            if uploader:
                meta.append(f"发布者={uploader}")
        sheet_names = self._normalize_text_list(detail.get("sheet_names") or [])
        if sheet_names:
            meta.append("工作表=" + ",".join(sheet_names[:6]))
        entities = self._normalize_text_list(detail.get("entities") or detail.get("objects") or [])
        if entities:
            meta.append("要素=" + ",".join(entities[:8]))
        tags = self._normalize_text_list(detail.get("mood_tags") or detail.get("tags") or detail.get("keywords") or [])
        if tags:
            meta.append("标签=" + ",".join(tags[:8]))
        if item.get("status") == "failed":
            failure = self._readable_failure_message(item)
            if failure:
                meta.append(f"失败={failure}")
        if meta:
            parts.append("   " + "；".join(meta))
        return parts

    def _build_inspect_followup(self, item: dict[str, Any]) -> str:
        lines = [
            "你刚刚查看了一份工作台材料：",
            self._compact_item_label(item),
        ]
        source_label = self._source_label(item)
        if source_label:
            lines.append(f"来源：{source_label}")
        detail = item.get("detail") if isinstance(item.get("detail"), dict) else {}
        hint = str(item.get("short_hint") or detail.get("summary") or detail.get("description") or "").strip()
        if hint:
            lines.append(f"印象：{hint}")
        if item.get("status") == "failed":
            failure_message = self._readable_failure_message(item)
            if failure_message:
                lines.append(f"失败原因：{failure_message}")
        tags = self._normalize_text_list(detail.get("mood_tags") or detail.get("tags") or [])
        if tags:
            lines.append(f"标签：{', '.join(tags[:10])}")
        media_info = detail.get("media_info") if isinstance(detail.get("media_info"), dict) else {}
        if media_info:
            lines.append("媒体规格：")
            lines.extend(self._render_media_info_lines(media_info, prefix="  "))
        remote_source = detail.get("remote_source") if isinstance(detail.get("remote_source"), dict) else {}
        if remote_source:
            lines.extend(self._render_remote_source_lines(remote_source, prefix="  "))
        preview = str(detail.get("text_preview") or detail.get("content_preview") or "").strip()
        if preview:
            excerpt_limit = 600
            lines.append(f"文本摘录：{preview[:excerpt_limit]}")
            if detail.get("preview_is_truncated"):
                lines.append("注意：这只是预览，不是全文；如果要处理完整文件，应继续用 read_attachment_section 展开指定范围。")
            elif len(preview) > excerpt_limit:
                lines.append("注意：这只是本次提示词节选，不是全文；忠实转换/导出原附件时不要复制这段节选，应让文件工具读取原始附件。")
        lines.append("请基于这份材料信息自然回应；如果用户聊完了，可以稍后用 clear_attachment_focus 移除它。")
        return "\n".join(lines)

    def _extract_section_content(self, item: dict[str, Any], *, section: str) -> str:
        original = self._extract_original_file_section(item, section=section)
        if original:
            return original

        detail = item.get("detail") if isinstance(item.get("detail"), dict) else {}
        normalized_section = str(section or "").strip()
        sheets = detail.get("sheets")
        if isinstance(sheets, list) and sheets:
            sheet_content = self._extract_sheet_section(sheets, normalized_section)
            if sheet_content:
                return sheet_content[:2200]

        tables = detail.get("tables")
        if isinstance(tables, list) and tables and ("表" in normalized_section or "table" in normalized_section.lower()):
            table_content = self._extract_table_section(tables, normalized_section)
            if table_content:
                return table_content[:2200]

        preview = str(detail.get("text_preview") or detail.get("content_preview") or "").strip()
        if not preview:
            return ""
        lines = preview.splitlines()
        line_range = self._parse_line_range(normalized_section)
        if line_range is not None:
            start, end = line_range
            selected = lines[max(0, start - 1) : max(start, end)]
            return "\n".join(selected).strip()[:2200]

        paragraph_index = self._parse_section_index(normalized_section)
        paragraphs = [block.strip() for block in re.split(r"\n\s*\n", preview) if block.strip()]
        if paragraph_index is not None and paragraphs:
            if 1 <= paragraph_index <= len(paragraphs):
                return paragraphs[paragraph_index - 1][:2200]
            return ""
        if normalized_section and normalized_section not in {"全文", "全部", "预览", "当前可用片段"}:
            lowered = normalized_section.lower()
            for index, line in enumerate(lines):
                if lowered in line.lower():
                    start = max(0, index - 2)
                    end = min(len(lines), index + 8)
                    return "\n".join(lines[start:end]).strip()[:2200]
        return preview[:2200]

    def read_material_for_generation(self, item: dict[str, Any], *, max_chars: int = 30000) -> str:
        """Return a larger, bounded source excerpt for backend rendering.

        This is still not long-term memory: it is only used when Akane asks the
        file generator to create an output from an attachment.
        """

        content = self._extract_original_file_section(item, section="全文", max_chars=max_chars)
        if content:
            return content[:max_chars]
        detail = item.get("detail") if isinstance(item.get("detail"), dict) else {}
        preview = str(detail.get("text_preview") or detail.get("content_preview") or "").strip()
        return preview[:max_chars]

    def _extract_original_file_section(self, item: dict[str, Any], *, section: str, max_chars: int = 12000) -> str:
        source_path = self._resolve_storage_path(item)
        if source_path is None:
            return ""
        suffix = source_path.suffix.lower()
        normalized_section = str(section or "").strip()
        try:
            if suffix in {".txt", ".md", ".markdown", ".log", ".json", ".toml", ".yaml", ".yml", ".csv", ".ini", ".cfg", ".conf", ".py", ".js", ".ts", ".tsx", ".jsx", ".html", ".css", ".xml", ".sql", ".java", ".c", ".cpp", ".h", ".hpp", ".cs", ".go", ".rs"}:
                return self._read_text_file_section(source_path, section=normalized_section, max_chars=max_chars)
            if suffix == ".docx":
                return self._read_docx_section(source_path, section=normalized_section, max_chars=max_chars)
            if suffix == ".xlsx":
                return self._read_xlsx_section(source_path, section=normalized_section, max_chars=max_chars)
            if suffix == ".pdf":
                return self._read_pdf_section(source_path, section=normalized_section, max_chars=max_chars)
        except Exception:
            return ""
        return ""

    def _resolve_storage_path(self, item: dict[str, Any]) -> Path | None:
        relpath = str(item.get("storage_relpath") or "").strip()
        if not relpath:
            return None
        if relpath.lower().startswith("workspace:"):
            if self.workspace_uri_resolver is None:
                return None
            try:
                candidate = self.workspace_uri_resolver(relpath)
            except Exception:
                return None
            if isinstance(candidate, Path) and candidate.exists() and candidate.is_file():
                return candidate
            return None
        storage_roots = [
            root
            for root in [self.base_dir, *self.legacy_base_dirs]
            if isinstance(root, Path)
        ]
        for storage_root in storage_roots:
            candidate = (storage_root / Path(relpath)).resolve()
            try:
                base = storage_root.resolve()
                candidate.relative_to(base)
            except Exception:
                continue
            if candidate.exists() and candidate.is_file():
                return candidate
        return None

    def _read_text_file_section(self, path: Path, *, section: str, max_chars: int) -> str:
        payload = path.read_bytes()[: max(256 * 1024, max_chars * 4)]
        text = ""
        for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
            try:
                text = payload.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if not text:
            text = payload.decode("utf-8", errors="replace")
        if path.suffix.lower() == ".csv":
            text = self._render_csv_text(text)
        return self._slice_plain_text_section(text, section=section, max_chars=max_chars)

    def _render_csv_text(self, text: str) -> str:
        rows: list[str] = []
        try:
            reader = csv.reader(text.splitlines())
            for index, row in enumerate(reader, start=1):
                if index > 500:
                    rows.append("...（后续 CSV 行已省略，可按行号继续展开）")
                    break
                rows.append(" | ".join(str(cell).strip() for cell in row))
        except Exception:
            return text
        return "\n".join(rows)

    def _read_docx_section(self, path: Path, *, section: str, max_chars: int) -> str:
        if importlib.util.find_spec("docx") is None:
            return ""
        from docx import Document  # type: ignore

        document = Document(str(path))
        if "表" in section or "table" in section.lower():
            index = self._parse_section_index(section) or 1
            if not (1 <= index <= len(document.tables)):
                return ""
            table = document.tables[index - 1]
            lines = [f"[表格 {index}] 行数约 {len(table.rows)}，列数约 {len(table.columns)}"]
            for row in table.rows[:120]:
                lines.append(" | ".join(str(cell.text or "").strip() for cell in row.cells if str(cell.text or "").strip()))
            return "\n".join(lines).strip()[:max_chars]
        paragraphs = [str(paragraph.text or "").strip() for paragraph in document.paragraphs]
        paragraphs = [text for text in paragraphs if text]
        return self._slice_plain_text_section("\n".join(paragraphs), section=section, max_chars=max_chars)

    def _read_xlsx_section(self, path: Path, *, section: str, max_chars: int) -> str:
        if importlib.util.find_spec("openpyxl") is None:
            return ""
        from openpyxl import load_workbook  # type: ignore

        workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
        try:
            sheet_name = self._select_sheet_name(workbook.sheetnames, section)
            sheet = workbook[sheet_name]
            line_range = self._parse_line_range(section)
            if line_range is None:
                start_row, end_row = 1, min(int(sheet.max_row or 1), 80)
            else:
                start_row, end_row = line_range
                end_row = min(end_row, start_row + 120)
            lines = [f"[{sheet_name}] 行 {start_row}-{end_row} / 共 {sheet.max_row or '?'} 行"]
            for row in sheet.iter_rows(min_row=start_row, max_row=end_row, values_only=True):
                values = [self._cell_to_text(value) for value in list(row)[:40]]
                if any(values):
                    lines.append(" | ".join(values))
                if sum(len(line) for line in lines) > max_chars:
                    break
            return "\n".join(lines).strip()[:max_chars]
        finally:
            workbook.close()

    def _read_pdf_section(self, path: Path, *, section: str, max_chars: int) -> str:
        if importlib.util.find_spec("pypdf") is None:
            return ""
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        normalized_section = str(section or "").strip().lower()
        if normalized_section in {"全文", "全部", "完整", "all", "full"}:
            start_page, end_page = 1, page_count
        else:
            page_range = self._parse_page_range(section)
            if page_range is None:
                start_page, end_page = 1, min(page_count, 5)
            else:
                start_page, end_page = page_range
                start_page = max(1, min(page_count, start_page))
                end_page = max(start_page, min(page_count, end_page))
        if page_count <= 0:
            start_page, end_page = 1, min(page_count, 5)
        chunks = [f"[PDF 页 {start_page}-{end_page} / 共 {page_count} 页]"]
        for page_no in range(start_page, end_page + 1):
            text = str(reader.pages[page_no - 1].extract_text() or "").strip()
            if text:
                chunks.append(f"--- 第 {page_no} 页 ---\n{text}")
            if sum(len(chunk) for chunk in chunks) > max_chars:
                break
        return "\n\n".join(chunks).strip()[:max_chars]

    def _slice_plain_text_section(self, text: str, *, section: str, max_chars: int) -> str:
        normalized = str(section or "").strip()
        if normalized and normalized not in {"全文", "全部", "预览", "当前可用片段"}:
            line_range = self._parse_line_range(normalized)
            if line_range is not None:
                start, end = line_range
                lines = text.splitlines()
                return "\n".join(lines[max(0, start - 1) : max(start, end)]).strip()[:max_chars]

            paragraph_index = self._parse_section_index(normalized)
            paragraphs = [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]
            if paragraph_index is not None and paragraphs:
                if 1 <= paragraph_index <= len(paragraphs):
                    return paragraphs[paragraph_index - 1][:max_chars]
                return ""

            lowered = normalized.lower()
            lines = text.splitlines()
            for index, line in enumerate(lines):
                if lowered in line.lower():
                    start = max(0, index - 6)
                    end = min(len(lines), index + 20)
                    return "\n".join(lines[start:end]).strip()[:max_chars]
        return text.strip()[:max_chars]

    def _select_sheet_name(self, sheet_names: list[str], section: str) -> str:
        lowered = str(section or "").strip().lower()
        index = self._parse_section_index(section)
        if index is not None and 1 <= index <= len(sheet_names):
            return sheet_names[index - 1]
        for name in sheet_names:
            if name.lower() in lowered:
                return name
        return sheet_names[0]

    def _parse_page_range(self, section: str) -> tuple[int, int] | None:
        text = str(section or "").strip()
        match = re.search(r"(?:第\s*)?(\d+)\s*(?:页|page)?\s*(?:-|到|至|~)\s*(?:第\s*)?(\d+)\s*(?:页|page)?", text, re.IGNORECASE)
        if match:
            start = int(match.group(1))
            end = int(match.group(2))
            if start > end:
                start, end = end, start
            return start, min(end, start + 10)
        match = re.search(r"(?:第\s*)?(\d+)\s*(?:页|page)", text, re.IGNORECASE)
        if match:
            page = int(match.group(1))
            return page, page
        return None

    def _cell_to_text(self, value: Any) -> str:
        if value is None:
            return ""
        return str(value).strip()

    def _extract_sheet_section(self, sheets: list[Any], section: str) -> str:
        selected = None
        index = self._parse_section_index(section)
        lowered = str(section or "").strip().lower()
        if index is not None and 1 <= index <= len(sheets):
            selected = sheets[index - 1]
        if selected is None and lowered:
            for sheet in sheets:
                if not isinstance(sheet, dict):
                    continue
                name = str(sheet.get("name") or "").strip()
                if name and name.lower() in lowered:
                    selected = sheet
                    break
        if selected is None:
            selected = sheets[0]
        if not isinstance(selected, dict):
            return ""
        rows = selected.get("preview_rows") if isinstance(selected.get("preview_rows"), list) else []
        output = [f"[{selected.get('name') or 'Sheet'}]"]
        for row in rows[:12]:
            if isinstance(row, list):
                output.append(" | ".join(str(cell) for cell in row if str(cell).strip()))
        return "\n".join(output).strip()

    def _extract_table_section(self, tables: list[Any], section: str) -> str:
        index = self._parse_section_index(section) or 1
        if not (1 <= index <= len(tables)):
            return ""
        table = tables[index - 1]
        if not isinstance(table, dict):
            return ""
        rows = table.get("preview_rows") if isinstance(table.get("preview_rows"), list) else []
        output = [f"[表格 {index}] 行数约 {table.get('row_count') or '?'}，列数约 {table.get('column_count') or '?'}"]
        for row in rows[:12]:
            if isinstance(row, list):
                output.append(" | ".join(str(cell) for cell in row if str(cell).strip()))
        return "\n".join(output).strip()

    def _parse_line_range(self, section: str) -> tuple[int, int] | None:
        text = str(section or "").strip()
        match = re.search(r"(?:第\s*)?(\d+)\s*(?:行|line)?\s*(?:-|到|至|~)\s*(?:第\s*)?(\d+)\s*(?:行|line)?", text, re.IGNORECASE)
        if match:
            start = int(match.group(1))
            end = int(match.group(2))
            if start > end:
                start, end = end, start
            return start, min(end, start + 80)
        match = re.search(r"(?:第\s*)?(\d+)\s*(?:行|line)", text, re.IGNORECASE)
        if match:
            start = int(match.group(1))
            return start, min(start + 20, start + 80)
        return None

    def _parse_section_index(self, section: str) -> int | None:
        text = str(section or "").strip()
        match = re.search(r"(?:第\s*)?(\d+)\s*(?:页|段|节|章|个|张|表|sheet|page|paragraph|section)?", text, re.IGNORECASE)
        if match:
            return int(match.group(1))
        return self._parse_chinese_number(text)

    def _build_workspace_followup(
        self,
        *,
        focused: list[dict[str, Any]],
        unresolved: list[str],
        ambiguous_targets: list[str],
        overflow: list[dict[str, Any]],
        reason: str,
    ) -> str:
        if not focused:
            lines = ["你刚刚同步了材料工作台，但没有成功放入任何材料。"]
        else:
            lines = [
                f"你刚刚把 {len(focused)} 个材料放到了当前工作台，接下来请只细看这些材料："
            ]
            rendered, budget_overflow = self._render_focus_items_with_budget(
                focused,
                total_budget=max(12000, DEFAULT_WORKSPACE_CHAR_BUDGET // 2),
            )
            lines.extend(rendered)
            if budget_overflow:
                lines.append(
                    "以下材料已经进入工作台，但本次工具回执空间不够完整展开："
                    + "、".join(self._compact_item_label(item) for item in budget_overflow[:8])
                    + "。下一轮系统提示词会继续按工作台预算挂载；必要时可用 read_attachment_section 精确展开。"
                )
        if unresolved:
            lines.append(f"没有找到这些目标：{', '.join(unresolved[:5])}。")
        if ambiguous_targets:
            lines.append("这些目标不够明确，存在多个候选，请让用户确认：" + "；".join(ambiguous_targets[:5]) + "。")
        if overflow:
            lines.append(
                f"本次目标超过工作台软上限 {WORKSPACE_MAX_TARGETS} 个，以下目标没有展开："
                + "、".join(self._compact_item_label(item) for item in overflow[:5])
            )
        if str(reason or "").strip():
            lines.append(f"本次同步原因：{str(reason).strip()[:160]}")
        lines.append("请基于当前材料工作台自然回应，不要再次重复同步同一批材料。")
        return "\n".join(lines)

    def _clip_rendered_lines(self, lines: list[str], char_budget: int) -> list[str]:
        budget = max(300, int(char_budget or WORKSPACE_ITEM_CHAR_BUDGET))
        rendered: list[str] = []
        used = 0
        for line in lines:
            text = str(line or "")
            remaining = budget - used
            if remaining <= 0:
                break
            if len(text) + 1 <= remaining:
                rendered.append(text)
                used += len(text) + 1
                continue
            if remaining > 80:
                rendered.append(text[: remaining - 32] + "\n   ...（此附件内容因工作台预算暂时截断）")
            break
        return rendered

    def _indent_block(self, text: str, *, prefix: str) -> str:
        return "\n".join(prefix + line if line else prefix.rstrip() for line in str(text or "").splitlines())

    def _format_size(self, size: int) -> str:
        value = max(0, int(size or 0))
        if value < 1024:
            return f"{value}B"
        if value < 1024 * 1024:
            return f"{value / 1024:.1f}KB"
        return f"{value / (1024 * 1024):.1f}MB"

    def _format_duration_label(self, value: Any) -> str:
        if not isinstance(value, (int, float)):
            return ""
        seconds = max(0, int(round(float(value))))
        hours, remainder = divmod(seconds, 3600)
        minutes, secs = divmod(remainder, 60)
        if hours:
            return f"{hours}:{minutes:02d}:{secs:02d}"
        return f"{minutes}:{secs:02d}"

    def _format_bitrate(self, value: int) -> str:
        if value <= 0:
            return ""
        if value >= 1_000_000:
            return f"{value / 1_000_000:.2f}Mbps"
        return f"{value / 1000:.0f}kbps"

    def _source_label(self, item: dict[str, Any]) -> str:
        source = str(item.get("source") or "").strip().lower()
        if not source:
            return ""
        labels = {
            "qq": "QQ",
            "desktop_pet": "桌宠本地",
            "remote_url": "链接下载",
            "web": "Web",
        }
        return labels.get(source, str(item.get("source") or "").strip()[:40])

    def _compact_item_label(self, item: dict[str, Any]) -> str:
        attachment_id = str(item.get("attachment_handle") or item.get("attachment_id") or "").strip()
        kind = self._kind_label(item.get("kind"))
        name = self._display_name(item)
        origin = str(item.get("origin_name") or "").strip()
        seq_label = self._sequence_label(item)
        seq_part = f"（{seq_label}）" if seq_label else ""
        if origin and origin != name:
            return f"[{attachment_id}]{seq_part} {kind}《{name}》({origin})"
        return f"[{attachment_id}]{seq_part} {kind}《{name}》"

    def _display_name(self, item: dict[str, Any]) -> str:
        return (
            str(item.get("summary_title") or "").strip()
            or str(item.get("origin_name") or "").strip()
            or str(item.get("attachment_id") or "").strip()
            or "未命名附件"
        )

    def _item_summary(self, item: dict[str, Any]) -> str:
        detail = item.get("detail") if isinstance(item.get("detail"), dict) else {}
        if item.get("status") == "failed":
            return self._readable_failure_message(item) or str(item.get("error_message") or "").strip()[:160]
        return str(item.get("short_hint") or detail.get("summary") or detail.get("description") or "").strip()[:160]

    def _readable_failure_message(self, item: dict[str, Any]) -> str:
        detail = item.get("detail") if isinstance(item.get("detail"), dict) else {}
        failure = detail.get("failure") if isinstance(detail.get("failure"), dict) else {}
        reason = str(failure.get("reason") or failure.get("message") or "").strip()
        if reason:
            return reason[:220]
        error = str(item.get("error_message") or "").strip()
        if not error:
            return ""
        lowered = error.lower()
        if "bad request" in lowered or "400 client error" in lowered:
            return "QQ 临时链接返回 400，请稍后重试，或让系统尝试从 NapCat 本地缓存重新读取。"
        if "视觉模型" in error or "vision" in lowered:
            return error[:220]
        if "timeout" in lowered or "timed out" in lowered:
            return "读取附件超时，可能是网络或 NapCat 缓存暂时不可用。"
        if "too large" in lowered or "附件过大" in error:
            return error[:220]
        return error[:220]

    def _kind_label(self, value: Any) -> str:
        kind = str(value or "").strip().lower()
        if kind == "image":
            return "图片"
        if kind == "audio":
            return "音频"
        if kind == "document":
            return "文档"
        return "文件"

    def _sequence_label(self, item: dict[str, Any]) -> str:
        sequence_no = int(item.get("sequence_no") or 0)
        if sequence_no <= 0:
            return ""
        kind = str(item.get("kind") or "").strip().lower()
        if kind == "image":
            return f"第{sequence_no}张图"
        if kind == "audio":
            return f"第{sequence_no}段音频"
        return f"第{sequence_no}个文件"

    def _normalize_targets(self, value: Any) -> list[str]:
        if value is None:
            return []
        if isinstance(value, str):
            raw_items = re.split(r"[,，、;；\n]+", value)
        elif isinstance(value, (list, tuple, set)):
            raw_items = list(value)
        else:
            raw_items = [value]
        targets: list[str] = []
        for item in raw_items:
            text = str(item or "").strip()
            if text and text not in targets:
                targets.append(text[:120])
        return targets

    def _dedupe_attachment_items(self, items: list[dict[str, Any]]) -> list[dict[str, Any]]:
        output: list[dict[str, Any]] = []
        seen_ids: set[str] = set()
        for item in items:
            item_id = str(item.get("attachment_id") or "").strip()
            if item_id and item_id in seen_ids:
                continue
            if item_id:
                seen_ids.add(item_id)
            output.append(item)
        return output

    def _format_ambiguity_target(self, *, target: str, items: list[dict[str, Any]]) -> str:
        if not items:
            return ""
        labels = [self._compact_item_label(item) for item in items[:4]]
        labels = [label for label in labels if label]
        if not labels:
            return ""
        return f"{target} -> {'、'.join(labels)}"

    def _filter_items_by_target_kind(
        self,
        items: list[dict[str, Any]],
        target: str,
        kind: str,
    ) -> list[dict[str, Any]]:
        explicit_kind = self._infer_kind_from_target(target)
        normalized_kind = explicit_kind or str(kind or "any").strip().lower()
        if normalized_kind in {"photo", "picture", "pic", "img"}:
            normalized_kind = "image"
        if normalized_kind in {"doc", "txt", "pdf", "text"}:
            normalized_kind = "document"
        if normalized_kind in {"music", "song", "voice"}:
            normalized_kind = "audio"
        if normalized_kind not in {"image", "document", "audio", "file"}:
            return list(items)
        if normalized_kind == "document":
            return [item for item in items if item.get("kind") in {"document", "file"}]
        return [item for item in items if item.get("kind") == normalized_kind]

    def _infer_kind_from_target(self, target: str) -> str:
        text = str(target or "").strip().lower()
        if any(token in text for token in ("图", "张", "image", "img", "pic", "photo")):
            return "image"
        if any(token in text for token in ("文件", "文档", "file", "doc", "pdf", "txt")):
            return "document"
        if any(token in text for token in ("音频", "歌", "audio", "song", "music")):
            return "audio"
        return ""

    def _latest_item(self, items: list[dict[str, Any]]) -> dict[str, Any] | None:
        ordered = sorted(
            items,
            key=lambda item: (
                int(item.get("created_at") or 0),
                int(item.get("updated_at") or 0),
                int(item.get("sequence_no") or 0),
            ),
        )
        return ordered[-1] if ordered else None

    def _oldest_item(self, items: list[dict[str, Any]]) -> dict[str, Any] | None:
        ordered = sorted(
            items,
            key=lambda item: (
                int(item.get("created_at") or 0),
                int(item.get("updated_at") or 0),
                int(item.get("sequence_no") or 0),
            ),
        )
        return ordered[0] if ordered else None

    def _sort_by_sequence(self, items: list[dict[str, Any]]) -> list[dict[str, Any]]:
        return sorted(
            items,
            key=lambda item: (
                str(item.get("kind") or ""),
                int(item.get("sequence_no") or 0),
                int(item.get("created_at") or 0),
            ),
        )

    def _parse_sequence_no(self, target: str) -> int | None:
        text = str(target or "").strip()
        match = re.search(r"(?:第\s*)?(\d+)\s*(?:张|个|份|段|图|文件|文档|音频)?", text)
        if match:
            return int(match.group(1))
        chinese = self._parse_chinese_number(text)
        return chinese

    def _looks_like_sequence_reference(self, target: str) -> bool:
        text = str(target or "").strip().lower()
        if not text:
            return False
        text = re.sub(r"\s+", "", text)
        if re.fullmatch(r"(?:第)?\d+(?:张图|个文件|份文档|段音频|张|个|份|段|图|文件|文档|音频)?", text):
            return True
        if re.fullmatch(r"(?:第)?[一二两三四五六七八九十]+(?:张图|个文件|份文档|段音频|张|个|份|段|图|文件|文档|音频)?", text):
            return True
        return False

    def _parse_reverse_ordinal(self, target: str) -> int | None:
        text = str(target or "").strip()
        if "倒数" not in text:
            return None
        match = re.search(r"倒数(?:第)?\s*(\d+)", text)
        if match:
            return int(match.group(1))
        return self._parse_chinese_number(text.replace("倒数", ""))

    def _parse_chinese_number(self, text: str) -> int | None:
        compact = re.sub(r"[第张个份段页节章表图文件文档音频\s]", "", str(text or ""))
        compact = compact.strip()
        if not compact:
            return None
        table = {
            "一": 1,
            "二": 2,
            "两": 2,
            "三": 3,
            "四": 4,
            "五": 5,
            "六": 6,
            "七": 7,
            "八": 8,
            "九": 9,
            "十": 10,
        }
        if compact in table:
            return table[compact]
        if compact.startswith("十") and len(compact) == 2 and compact[1] in table:
            return 10 + int(table[compact[1]])
        if compact.endswith("十") and len(compact) == 2 and compact[0] in table:
            return int(table[compact[0]]) * 10
        if "十" in compact:
            left, right = compact.split("十", 1)
            if left in table and right in table:
                return int(table[left]) * 10 + int(table[right])
        return None

    def _normalize_text_list(self, value: Any) -> list[str]:
        if isinstance(value, str):
            raw_items = [item.strip() for item in value.replace("，", ",").split(",")]
        elif isinstance(value, (list, tuple, set)):
            raw_items = [str(item or "").strip() for item in value]
        else:
            return []
        result: list[str] = []
        for item in raw_items:
            if item and item not in result:
                result.append(item[:40])
        return result
