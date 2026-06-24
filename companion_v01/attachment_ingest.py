from __future__ import annotations

import csv
from dataclasses import dataclass
import importlib.util
import json
import mimetypes
import shutil
import subprocess
import time
import re
from pathlib import Path
from typing import Any, Callable
from urllib.parse import urlparse, unquote

import requests

import config

from .attachment_inbox import AttachmentInboxService
from .background_tasks import BackgroundTaskRunner
from .store import MemoryStore
from .vision_service import VisionObservationService


TEXT_SUFFIXES = {
    ".txt",
    ".md",
    ".markdown",
    ".log",
    ".lrc",
    ".srt",
    ".vtt",
    ".json",
    ".toml",
    ".yaml",
    ".yml",
    ".csv",
    ".ini",
    ".cfg",
    ".conf",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".html",
    ".css",
    ".xml",
    ".sql",
    ".java",
    ".c",
    ".cpp",
    ".h",
    ".hpp",
    ".cs",
    ".go",
    ".rs",
}

DOCUMENT_SUFFIXES = {
    ".pdf",
    ".docx",
    ".xlsx",
}

MEDIA_SUFFIXES = {
    ".mp3",
    ".wav",
    ".flac",
    ".m4a",
    ".aac",
    ".ogg",
    ".opus",
    ".mp4",
    ".mov",
    ".mkv",
    ".webm",
    ".avi",
}

AUDIO_MEDIA_SUFFIXES = {
    ".mp3",
    ".wav",
    ".flac",
    ".m4a",
    ".aac",
    ".ogg",
    ".opus",
}

IMAGE_SUFFIXES = {
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
}

REMOTE_MEDIA_DEFAULT_TIMEOUT = 180.0
REMOTE_MEDIA_DEFAULT_MAX_BYTES = 1024 * 1024 * 1024
REMOTE_MEDIA_DEFAULT_MAX_URLS = 8
REMOTE_MEDIA_DEFAULT_USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/123.0.0.0 Safari/537.36"
)
REMOTE_MEDIA_DEFAULT_REFERER = "https://www.bilibili.com/"


@dataclass(frozen=True)
class RemoteMediaDescriptor:
    source_url: str
    title: str
    ext: str
    mime_type: str
    kind: str
    download_mode: str
    webpage_url: str = ""
    extractor: str = ""
    extractor_key: str = ""
    uploader: str = ""
    channel: str = ""
    duration_seconds: float | None = None
    thumbnail_url: str = ""
    description: str = ""
    file_size_hint: int = 0

    @property
    def origin_name(self) -> str:
        suffix = f".{self.ext.lstrip('.')}" if str(self.ext).strip() else ""
        title = str(self.title or "remote_media").strip() or "remote_media"
        return f"{title}{suffix}"


class AttachmentIngestService:
    def __init__(
        self,
        *,
        base_dir: Path,
        store: MemoryStore,
        attachment_service: AttachmentInboxService,
        vision_service: VisionObservationService | None,
        background_tasks: BackgroundTaskRunner | None = None,
        legacy_base_dirs: list[Path] | tuple[Path, ...] | None = None,
        ensure_storage_ready: Callable[[], Any] | None = None,
        workspace_uri_resolver: Callable[[str], Path | None] | None = None,
    ) -> None:
        self.base_dir = Path(base_dir)
        self.base_dir.mkdir(parents=True, exist_ok=True)
        self.legacy_base_dirs = [
            Path(item)
            for item in list(legacy_base_dirs or [])
            if Path(item) != self.base_dir
        ]
        self.ensure_storage_ready = ensure_storage_ready
        self.workspace_uri_resolver = workspace_uri_resolver
        self.store = store
        self.attachment_service = attachment_service
        self.vision_service = vision_service
        self.background_tasks = background_tasks or BackgroundTaskRunner(
            {"attachment": int(getattr(config, "BACKGROUND_ATTACHMENT_WORKERS", 3) or 3)},
            default_workers=1,
        )

    def ingest_qq_attachments(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        attachments: list[dict[str, Any]],
        timestamp: int | None = None,
    ) -> list[dict[str, Any]]:
        effective_ts = int(timestamp or time.time())
        created_items: list[dict[str, Any]] = []
        for payload in attachments or []:
            if not isinstance(payload, dict):
                continue
            kind = self._normalize_kind(payload.get("kind"))
            if kind not in {"image", "document", "audio", "file"}:
                continue
            item = self.attachment_service.create_pending(
                profile_user_id=profile_user_id,
                session_id=session_id,
                source="qq",
                kind=kind,
                origin_name=self._clean_filename(payload.get("origin_name") or payload.get("name") or ""),
                mime_type=str(payload.get("mime_type") or "").strip(),
                file_ext=str(payload.get("file_ext") or "").strip(),
                file_size=self._safe_int(payload.get("file_size")),
                source_event_id=str(payload.get("source_event_id") or "").strip(),
                source_message_id=str(payload.get("source_message_id") or "").strip(),
                timestamp=effective_ts,
            )
            created_items.append(item)
            self.background_tasks.submit(
                lane="attachment",
                name=f"qq_attachment:{item.get('attachment_handle') or item.get('attachment_id')}",
                fn=self._process_qq_attachment,
                args=(item, dict(payload), effective_ts),
            )
        return created_items

    def ingest_local_file(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        source_path: Path | str,
        origin_name: str = "",
        mime_type: str = "",
        kind: str = "",
        source: str = "local",
        timestamp: int | None = None,
    ) -> dict[str, Any]:
        """Synchronously copy a trusted local file into the attachment inbox."""

        effective_ts = int(timestamp or time.time())
        local_path = Path(source_path)
        if not local_path.exists() or not local_path.is_file():
            raise FileNotFoundError(str(local_path))

        clean_name = self._clean_filename(origin_name or local_path.name)
        guessed_mime = str(mime_type or mimetypes.guess_type(clean_name or str(local_path))[0] or "").strip()
        suffix = local_path.suffix.lower() or Path(clean_name).suffix.lower()
        normalized_kind = self._normalize_kind(kind)
        if not kind:
            normalized_kind = "audio" if suffix in AUDIO_MEDIA_SUFFIXES or guessed_mime.startswith("audio/") else "file"

        item = self.attachment_service.create_pending(
            profile_user_id=profile_user_id,
            session_id=session_id,
            source=str(source or "local").strip() or "local",
            kind=normalized_kind,
            origin_name=clean_name,
            mime_type=guessed_mime,
            file_ext=suffix,
            file_size=local_path.stat().st_size,
            source_event_id="",
            source_message_id="",
            timestamp=effective_ts,
        )
        self._process_qq_attachment(
            item,
            {
                "path": str(local_path),
                "origin_name": clean_name,
                "mime_type": guessed_mime,
                "file_ext": suffix,
                "kind": normalized_kind,
            },
            effective_ts,
        )
        return (
            self.store.get_attachment_inbox_item(
                profile_user_id=profile_user_id,
                session_id=session_id,
                attachment_id=str(item.get("attachment_id") or ""),
            )
            or item
        )

    def register_workspace_file(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        workspace_uri: str,
        timestamp: int | None = None,
    ) -> dict[str, Any]:
        normalized_uri = str(workspace_uri or "").strip()
        if not normalized_uri.lower().startswith("workspace:"):
            raise ValueError("workspace registration requires a workspace:/ URI")
        if self.workspace_uri_resolver is None:
            raise RuntimeError("workspace file resolver is unavailable")
        source_path = self.workspace_uri_resolver(normalized_uri)
        if source_path is None or not source_path.exists() or not source_path.is_file():
            raise FileNotFoundError(normalized_uri)

        effective_ts = int(timestamp or time.time())
        mime_type = str(mimetypes.guess_type(source_path.name)[0] or "").strip()
        file_ext = source_path.suffix.lower()
        kind = self._infer_local_kind(source_path=source_path, mime_type=mime_type)
        existing = self.store.get_attachment_inbox_item_by_storage_relpath(
            profile_user_id=profile_user_id,
            session_id=session_id,
            storage_relpath=normalized_uri,
            statuses=["ready", "pending_observation", "failed"],
        )
        if existing is not None and str(existing.get("status") or "") == "pending_observation":
            return {
                "status": "already_registered",
                "workspace_uri": normalized_uri,
                "item": existing,
            }

        if existing is None:
            item = self.attachment_service.create_pending(
                profile_user_id=profile_user_id,
                session_id=session_id,
                source="workspace",
                kind=kind,
                origin_name=self._clean_filename(source_path.name),
                mime_type=mime_type,
                file_ext=file_ext,
                file_size=source_path.stat().st_size,
                storage_relpath=normalized_uri,
                timestamp=effective_ts,
            )
            registration_status = "registered"
        else:
            item = self.store.update_attachment_inbox_item(
                profile_user_id=profile_user_id,
                session_id=session_id,
                attachment_id=str(existing.get("attachment_id") or ""),
                status="pending_observation",
                short_hint="正在重新读取这个工作区文件。",
                detail={},
                error_message="",
                mime_type=mime_type,
                file_ext=file_ext,
                file_size=source_path.stat().st_size,
                storage_relpath=normalized_uri,
                updated_at=effective_ts,
            ) or existing
            registration_status = "refreshed"

        self._process_qq_attachment(
            item,
            {
                "workspace_uri": normalized_uri,
                "origin_name": source_path.name,
                "mime_type": mime_type,
                "file_ext": file_ext,
                "kind": kind,
            },
            effective_ts,
        )
        latest = (
            self.store.get_attachment_inbox_item(
                profile_user_id=profile_user_id,
                session_id=session_id,
                attachment_id=str(item.get("attachment_id") or ""),
            )
            or item
        )
        return {
            "status": registration_status,
            "workspace_uri": normalized_uri,
            "item": latest,
        }

    def _process_qq_attachment(
        self,
        item: dict[str, Any],
        payload: dict[str, Any],
        timestamp: int,
    ) -> None:
        try:
            source_path = self._materialize_attachment_file(item=item, payload=payload)
            if source_path is None:
                self._mark_failed(item, "没有可下载或可读取的附件地址。", timestamp=timestamp)
                return

            workspace_uri = str(payload.get("workspace_uri") or "").strip()
            relpath = workspace_uri or self._storage_relpath(source_path)
            mime_type = str(payload.get("mime_type") or mimetypes.guess_type(str(source_path))[0] or "").strip()
            file_ext = source_path.suffix.lower()
            self.store.update_attachment_inbox_item(
                profile_user_id=str(item.get("profile_user_id") or ""),
                session_id=str(item.get("session_id") or ""),
                attachment_id=str(item.get("attachment_id") or ""),
                storage_relpath=relpath,
                mime_type=mime_type,
                file_ext=file_ext,
                file_size=source_path.stat().st_size if source_path.exists() else 0,
                updated_at=timestamp,
            )

            kind = str(item.get("kind") or "").strip().lower()
            if kind == "image" and self.vision_service is not None:
                scheduled = self.vision_service.schedule_attachment_image_observation(
                    attachment={
                        **item,
                        "storage_relpath": relpath,
                        "mime_type": mime_type,
                        "file_ext": file_ext,
                    },
                    source_path=source_path,
                )
                if scheduled is None:
                    self._mark_failed(item, "图片已接收，但视觉模型尚未配置或无法启动。", timestamp=timestamp)
                elif str(scheduled.get("status") or "").strip().lower() == "ready":
                    self._mark_image_ready_from_observation(
                        item=item,
                        observation=scheduled,
                        timestamp=timestamp,
                    )
                return

            file_card = self._build_file_card(source_path=source_path, item=item, mime_type=mime_type)
            self.attachment_service.mark_ready(
                profile_user_id=str(item.get("profile_user_id") or ""),
                session_id=str(item.get("session_id") or ""),
                attachment_id=str(item.get("attachment_id") or ""),
                summary_title=file_card["summary_title"],
                short_hint=file_card["short_hint"],
                detail=file_card["detail"],
                timestamp=timestamp,
            )
        except Exception as exc:
            self._mark_failed(item, str(exc), timestamp=timestamp)

    def retry_attachment(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        target: str = "latest",
        kind: str = "any",
        timestamp: int | None = None,
    ) -> dict[str, Any]:
        effective_ts = int(timestamp or time.time())
        item = self.attachment_service.resolve_attachment(
            profile_user_id=profile_user_id,
            session_id=session_id,
            target=target,
            kind=kind,
        )
        if item is None:
            return {
                "ok": False,
                "status": "not_found",
                "item": None,
                "followup_context": "你刚刚想重试某个附件，但当前没有找到明确匹配的临时图片或文件。请自然向用户确认是哪一个。",
            }

        current_status = str(item.get("status") or "").strip()
        if current_status == "ready":
            return {
                "ok": True,
                "status": "already_ready",
                "item": item,
                "followup_context": (
                    "这个附件其实已经处理成功了，可以直接查看。"
                    "请自然告诉用户不用再重试，并基于已有附件内容继续回应。"
                ),
            }
        if current_status == "pending_observation":
            return {
                "ok": True,
                "status": "already_pending",
                "item": item,
                "followup_context": (
                    "这个附件仍在处理中，还不需要重复重试。"
                    "请自然告诉用户稍等一下，不要再次调用 retry_attachment。"
                ),
            }

        previous_error = str(item.get("error_message") or "").strip()
        retry_item = self.store.update_attachment_inbox_item(
            profile_user_id=profile_user_id,
            session_id=session_id,
            attachment_id=str(item.get("attachment_id") or ""),
            status="pending_observation",
            short_hint="正在重新读取这个附件。",
            error_message="",
            detail={
                "retry": {
                    "previous_status": current_status,
                    "previous_error": previous_error[:500],
                    "retried_at": effective_ts,
                }
            },
            updated_at=effective_ts,
        ) or dict(item, status="pending_observation", error_message="", updated_at=effective_ts)

        payload = self._build_retry_payload(retry_item, previous_error=previous_error)
        self.background_tasks.submit(
            lane="attachment",
            name=f"retry_attachment:{retry_item.get('attachment_handle') or retry_item.get('attachment_id')}",
            fn=self._process_qq_attachment,
            args=(retry_item, payload, effective_ts),
        )
        return {
            "ok": True,
            "status": "retry_started",
            "item": retry_item,
            "followup_context": (
                f"你刚刚已经开始重新处理工作台材料 {retry_item.get('attachment_handle') or ''}。"
                "请自然告诉用户你在重新试一次；结果会在处理完成后进入当前材料工作台。"
                "不要立刻断言已经成功，也不要重复调用 retry_attachment。"
            ),
        }

    def fetch_media_from_urls(
        self,
        *,
        profile_user_id: str,
        session_id: str,
        urls: list[str] | tuple[str, ...] | set[str] | str,
        preferred_title: str = "",
        timestamp: int | None = None,
    ) -> dict[str, Any]:
        effective_ts = int(timestamp or time.time())
        normalized_urls = self._normalize_remote_urls(urls)
        if not normalized_urls:
            return {
                "ok": False,
                "items": [],
                "failed": [],
                "followup_context": (
                    "你刚刚想通过链接获取视频或音频，但这次没有拿到明确可用的 http/https 链接。"
                    "请自然向用户确认具体链接，不要重复调用 fetch_media_from_url。"
                ),
            }

        items: list[dict[str, Any]] = []
        failures: list[dict[str, str]] = []
        single_title = str(preferred_title or "").strip()
        for index, url in enumerate(normalized_urls, start=1):
            title_hint = single_title if len(normalized_urls) == 1 else ""
            item: dict[str, Any] | None = None
            try:
                descriptor = self._fetch_remote_media_descriptor(
                    url=url,
                    preferred_title=title_hint,
                )
                item = self.attachment_service.create_pending(
                    profile_user_id=profile_user_id,
                    session_id=session_id,
                    source="remote_url",
                    kind=descriptor.kind,
                    origin_name=self._clean_filename(descriptor.origin_name),
                    mime_type=str(descriptor.mime_type or "").strip(),
                    file_ext=f".{descriptor.ext.lstrip('.')}" if descriptor.ext else "",
                    file_size=max(0, int(descriptor.file_size_hint or 0)),
                    source_event_id=url,
                    source_message_id="",
                    timestamp=effective_ts,
                )
                ready_item = self._process_remote_media_download(
                    item=item,
                    descriptor=descriptor,
                    timestamp=effective_ts,
                )
                if ready_item is not None:
                    self._clear_stale_failed_remote_items(
                        ready_item=ready_item,
                        source_url=url,
                        timestamp=effective_ts,
                    )
                    items.append(ready_item)
                    continue
                self._mark_failed(item, "下载完成后没有生成可用的媒体文件。", timestamp=effective_ts)
                failures.append(
                    {
                        "url": url,
                        "error": "下载完成后没有生成可用的媒体文件。",
                        "index_label": f"第{index}个链接",
                    }
                )
            except Exception as exc:
                error_message = self._humanize_remote_fetch_error(str(exc))
                if item is not None:
                    self._mark_failed(item, error_message, timestamp=effective_ts)
                failures.append(
                    {
                        "url": url,
                        "error": error_message,
                        "index_label": f"第{index}个链接",
                    }
                )

        return {
            "ok": bool(items),
            "items": items,
            "failed": failures,
            "followup_context": self._build_remote_media_followup(
                items=items,
                failures=failures,
            ),
        }

    def _clear_stale_failed_remote_items(
        self,
        *,
        ready_item: dict[str, Any],
        source_url: str,
        timestamp: int,
    ) -> None:
        normalized_url = str(source_url or "").strip()
        if not normalized_url:
            return
        profile_user_id = str(ready_item.get("profile_user_id") or "").strip()
        session_id = str(ready_item.get("session_id") or "").strip()
        ready_id = str(ready_item.get("attachment_id") or "").strip()
        if not profile_user_id or not session_id:
            return

        failed_items = self.store.list_attachment_inbox_items(
            profile_user_id=profile_user_id,
            session_id=session_id,
            statuses=["failed"],
            limit=200,
        )
        for failed in failed_items:
            failed_id = str(failed.get("attachment_id") or "").strip()
            if not failed_id or failed_id == ready_id:
                continue
            if str(failed.get("source") or "").strip() != "remote_url":
                continue
            if str(failed.get("source_event_id") or "").strip() != normalized_url:
                continue
            self.store.update_attachment_inbox_item(
                profile_user_id=profile_user_id,
                session_id=session_id,
                attachment_id=failed_id,
                status="cleared",
                updated_at=timestamp,
            )

    def _process_remote_media_download(
        self,
        *,
        item: dict[str, Any],
        descriptor: RemoteMediaDescriptor,
        timestamp: int,
    ) -> dict[str, Any] | None:
        source_path = self._download_remote_media(
            item=item,
            descriptor=descriptor,
        )
        relpath = self._storage_relpath(source_path)
        mime_type = str(descriptor.mime_type or mimetypes.guess_type(str(source_path))[0] or "").strip()
        file_ext = source_path.suffix.lower()
        self.store.update_attachment_inbox_item(
            profile_user_id=str(item.get("profile_user_id") or ""),
            session_id=str(item.get("session_id") or ""),
            attachment_id=str(item.get("attachment_id") or ""),
            storage_relpath=relpath,
            mime_type=mime_type,
            file_ext=file_ext,
            file_size=source_path.stat().st_size if source_path.exists() else 0,
            updated_at=timestamp,
        )

        file_card = self._build_file_card(source_path=source_path, item=item, mime_type=mime_type)
        remote_card = self._enrich_remote_media_card(
            file_card=file_card,
            descriptor=descriptor,
        )
        return self.attachment_service.mark_ready(
            profile_user_id=str(item.get("profile_user_id") or ""),
            session_id=str(item.get("session_id") or ""),
            attachment_id=str(item.get("attachment_id") or ""),
            summary_title=str(remote_card.get("summary_title") or file_card.get("summary_title") or ""),
            short_hint=str(remote_card.get("short_hint") or file_card.get("short_hint") or ""),
            detail=remote_card.get("detail") if isinstance(remote_card.get("detail"), dict) else {},
            timestamp=timestamp,
        )

    def _normalize_remote_urls(self, value: list[str] | tuple[str, ...] | set[str] | str) -> list[str]:
        if isinstance(value, str):
            raw_items = re.split(r"[\s,，;；]+", value)
        elif isinstance(value, (list, tuple, set)):
            raw_items = list(value)
        else:
            raw_items = [value]
        seen: set[str] = set()
        normalized: list[str] = []
        limit = max(1, int(getattr(config, "REMOTE_MEDIA_MAX_URLS_PER_CALL", REMOTE_MEDIA_DEFAULT_MAX_URLS) or REMOTE_MEDIA_DEFAULT_MAX_URLS))
        for item in raw_items:
            text = str(item or "").strip()
            if not text:
                continue
            parsed = urlparse(text)
            if parsed.scheme not in {"http", "https"}:
                continue
            if text in seen:
                continue
            seen.add(text)
            normalized.append(text[:1000])
            if len(normalized) >= limit:
                break
        return normalized

    def _fetch_remote_media_descriptor(
        self,
        *,
        url: str,
        preferred_title: str = "",
    ) -> RemoteMediaDescriptor:
        direct_descriptor = self._build_direct_media_descriptor(
            url=url,
            preferred_title=preferred_title,
        )
        if direct_descriptor is not None:
            return direct_descriptor

        descriptor = self._extract_remote_media_with_yt_dlp(
            url=url,
            preferred_title=preferred_title,
        )
        if descriptor is not None:
            return descriptor

        raise RuntimeError(
            "当前环境没有可用的公开视频下载器；请先安装 yt-dlp，或者提供一个可直接下载的音频/视频文件链接。"
        )

    def _build_direct_media_descriptor(
        self,
        *,
        url: str,
        preferred_title: str = "",
    ) -> RemoteMediaDescriptor | None:
        parsed = urlparse(url)
        suffix = Path(unquote(parsed.path or "")).suffix.lower()
        if suffix not in MEDIA_SUFFIXES:
            return None
        title = str(preferred_title or Path(parsed.path).stem or parsed.netloc or "remote_media").strip() or "remote_media"
        mime_type = mimetypes.guess_type(parsed.path or "")[0] or ""
        if not mime_type:
            mime_type = "audio/mpeg" if suffix in AUDIO_MEDIA_SUFFIXES else "video/mp4"
        kind = "audio" if suffix in AUDIO_MEDIA_SUFFIXES or mime_type.startswith("audio/") else "file"
        return RemoteMediaDescriptor(
            source_url=url,
            webpage_url=url,
            title=title,
            ext=suffix.lstrip("."),
            mime_type=mime_type,
            kind=kind,
            download_mode="direct",
            extractor="direct",
            extractor_key="direct",
        )

    def _extract_remote_media_with_yt_dlp(
        self,
        *,
        url: str,
        preferred_title: str = "",
    ) -> RemoteMediaDescriptor | None:
        if importlib.util.find_spec("yt_dlp") is None:
            return None
        try:
            from yt_dlp import YoutubeDL  # type: ignore
        except Exception as exc:
            raise RuntimeError(f"yt-dlp 无法导入：{exc}") from exc

        options = {
            "quiet": True,
            "no_warnings": True,
            "noplaylist": True,
            "skip_download": True,
        }
        options.update(
            self._yt_dlp_common_options(
                timeout=float(getattr(config, "REMOTE_MEDIA_DOWNLOAD_TIMEOUT", REMOTE_MEDIA_DEFAULT_TIMEOUT) or REMOTE_MEDIA_DEFAULT_TIMEOUT)
            )
        )
        try:
            with YoutubeDL(options) as ydl:
                info = ydl.extract_info(url, download=False)
        except Exception as exc:
            raise RuntimeError(f"链接解析失败：{exc}") from exc

        if not isinstance(info, dict):
            raise RuntimeError("链接解析失败：没有拿到可用的媒体信息。")
        if str(info.get("_type") or "").strip().lower() == "playlist" or isinstance(info.get("entries"), list):
            raise RuntimeError("暂时只支持单个音频/视频链接，不支持整条播放列表或合集。")

        ext = str(info.get("ext") or "").strip().lower()
        if not ext:
            webpage_url = str(info.get("webpage_url") or url).strip()
            ext = Path(unquote(urlparse(webpage_url).path or "")).suffix.lower().lstrip(".")
        mime_type = mimetypes.guess_type(f"sample.{ext or 'bin'}")[0] or ""
        if not mime_type:
            mime_type = "audio/mpeg" if str(info.get("vcodec") or "").strip().lower() == "none" else "video/mp4"
        title = str(preferred_title or info.get("title") or "").strip()
        if not title:
            title = Path(unquote(urlparse(str(info.get("webpage_url") or url).strip()).path or "")).stem or "remote_media"
        is_audio_only = str(info.get("vcodec") or "").strip().lower() == "none" and str(info.get("acodec") or "").strip().lower() not in {"", "none"}
        return RemoteMediaDescriptor(
            source_url=url,
            webpage_url=str(info.get("webpage_url") or url).strip(),
            title=title,
            ext=ext or ("mp3" if is_audio_only else "mp4"),
            mime_type=mime_type,
            kind="audio" if is_audio_only else "file",
            download_mode="yt_dlp",
            extractor=str(info.get("extractor") or urlparse(url).netloc).strip(),
            extractor_key=str(info.get("extractor_key") or "").strip(),
            uploader=str(info.get("uploader") or "").strip(),
            channel=str(info.get("channel") or info.get("uploader_id") or "").strip(),
            duration_seconds=self._safe_float(info.get("duration")),
            thumbnail_url=str(info.get("thumbnail") or "").strip(),
            description=str(info.get("description") or "").strip()[:1000],
            file_size_hint=self._safe_int(info.get("filesize") or info.get("filesize_approx")),
        )

    def _download_remote_media(
        self,
        *,
        item: dict[str, Any],
        descriptor: RemoteMediaDescriptor,
    ) -> Path:
        target_dir = self._workspace_date_dir(item)
        target_dir.mkdir(parents=True, exist_ok=True)
        handle = self._safe_path_part(item.get("attachment_handle") or item.get("attachment_id"))
        timeout = float(getattr(config, "REMOTE_MEDIA_DOWNLOAD_TIMEOUT", REMOTE_MEDIA_DEFAULT_TIMEOUT) or REMOTE_MEDIA_DEFAULT_TIMEOUT)
        max_bytes = int(getattr(config, "REMOTE_MEDIA_MAX_BYTES", REMOTE_MEDIA_DEFAULT_MAX_BYTES) or REMOTE_MEDIA_DEFAULT_MAX_BYTES)

        if descriptor.download_mode == "direct":
            suffix = f".{descriptor.ext.lstrip('.')}" if descriptor.ext else ".bin"
            target_path = self._available_attachment_path(
                target_dir=target_dir,
                handle=handle,
                origin_name=descriptor.origin_name,
                suffix=suffix,
            )
            self._download_to_path(
                url=descriptor.source_url,
                target_path=target_path,
                timeout=timeout,
                max_bytes=max_bytes,
                headers={
                    "User-Agent": "Mozilla/5.0 AkaneCompanionLab/1.0",
                    "Accept": "*/*",
                },
            )
            return target_path

        download_stem = self._available_attachment_stem(
            target_dir=target_dir,
            handle=handle,
            origin_name=descriptor.origin_name,
        )
        return self._download_remote_media_with_yt_dlp(
            descriptor=descriptor,
            target_dir=target_dir,
            handle=download_stem,
            timeout=timeout,
            max_bytes=max_bytes,
        )

    def _download_remote_media_with_yt_dlp(
        self,
        *,
        descriptor: RemoteMediaDescriptor,
        target_dir: Path,
        handle: str,
        timeout: float,
        max_bytes: int,
    ) -> Path:
        try:
            from yt_dlp import YoutubeDL  # type: ignore
        except Exception as exc:
            raise RuntimeError(f"yt-dlp 无法导入：{exc}") from exc

        ffmpeg_path = shutil.which("ffmpeg")
        options: dict[str, Any] = {
            "quiet": True,
            "no_warnings": True,
            "noplaylist": True,
            "overwrites": True,
            "outtmpl": {"default": str(target_dir / f"{handle}.%(ext)s")},
            "cachedir": False,
            "writesubtitles": False,
            "writeautomaticsub": False,
            "writethumbnail": False,
            "writeinfojson": False,
            "restrictfilenames": False,
            "windowsfilenames": False,
        }
        options.update(self._yt_dlp_common_options(timeout=timeout))
        if ffmpeg_path:
            options["ffmpeg_location"] = str(Path(ffmpeg_path).parent)
        if max_bytes > 0:
            options["max_filesize"] = max_bytes
        try:
            with YoutubeDL(options) as ydl:
                ydl.download([descriptor.source_url])
        except Exception as exc:
            raise RuntimeError(f"下载失败：{exc}") from exc

        downloaded = self._locate_downloaded_remote_media_file(
            target_dir=target_dir,
            handle=handle,
        )
        if downloaded is None:
            raise RuntimeError("下载完成后没有找到可用的媒体文件。")
        if max_bytes > 0 and downloaded.stat().st_size > max_bytes:
            raise RuntimeError(f"下载后的媒体文件过大，当前限制为 {max_bytes} bytes。")
        return downloaded

    def _yt_dlp_common_options(self, *, timeout: float | None = None) -> dict[str, Any]:
        user_agent = (
            str(getattr(config, "REMOTE_MEDIA_YTDLP_USER_AGENT", "") or "").strip()
            or REMOTE_MEDIA_DEFAULT_USER_AGENT
        )
        referer = (
            str(getattr(config, "REMOTE_MEDIA_YTDLP_REFERER", "") or "").strip()
            or REMOTE_MEDIA_DEFAULT_REFERER
        )
        headers = {
            "User-Agent": user_agent,
            "Accept": "*/*",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
        }
        if referer:
            headers["Referer"] = referer
        options: dict[str, Any] = {
            "http_headers": headers,
        }
        if timeout is not None:
            options["socket_timeout"] = timeout
        cookiefile = str(getattr(config, "REMOTE_MEDIA_YTDLP_COOKIEFILE", "") or "").strip()
        if cookiefile:
            options["cookiefile"] = str(Path(cookiefile).expanduser())
        else:
            browser_spec = self._parse_ytdlp_browser_cookie_spec(
                str(getattr(config, "REMOTE_MEDIA_YTDLP_COOKIES_FROM_BROWSER", "") or "").strip()
            )
            if browser_spec is not None:
                options["cookiesfrombrowser"] = browser_spec
        return options

    def _parse_ytdlp_browser_cookie_spec(self, value: str) -> tuple[str, str | None, str | None, str | None] | None:
        raw = str(value or "").strip()
        if not raw:
            return None
        container: str | None = None
        if "::" in raw:
            raw, container = raw.split("::", 1)
            container = container.strip() or None
        profile: str | None = None
        if ":" in raw:
            raw, profile = raw.split(":", 1)
            profile = profile.strip() or None
        keyring: str | None = None
        if "+" in raw:
            raw, keyring = raw.split("+", 1)
            keyring = keyring.strip().upper() or None
        browser = raw.strip().lower()
        if not browser:
            return None
        return (browser, profile, keyring, container)

    def _locate_downloaded_remote_media_file(self, *, target_dir: Path, handle: str) -> Path | None:
        ignored_suffixes = {
            ".part",
            ".ytdl",
            ".json",
            ".description",
            ".vtt",
            ".srt",
            ".ass",
            ".lrc",
            ".jpg",
            ".jpeg",
            ".png",
            ".webp",
        }
        candidates: list[Path] = []
        for candidate in target_dir.glob(f"{handle}.*"):
            if not candidate.is_file():
                continue
            if candidate.suffix.lower() in ignored_suffixes:
                continue
            candidates.append(candidate)
        if not candidates:
            return None
        return max(candidates, key=lambda path: path.stat().st_size)

    def _enrich_remote_media_card(
        self,
        *,
        file_card: dict[str, Any],
        descriptor: RemoteMediaDescriptor,
    ) -> dict[str, Any]:
        summary_title = str(file_card.get("summary_title") or descriptor.title or descriptor.origin_name).strip()
        detail = dict(file_card.get("detail") or {}) if isinstance(file_card.get("detail"), dict) else {}
        detail["remote_source"] = {
            "platform": str(descriptor.extractor or descriptor.extractor_key or urlparse(descriptor.webpage_url or descriptor.source_url).netloc).strip(),
            "extractor": str(descriptor.extractor or "").strip(),
            "extractor_key": str(descriptor.extractor_key or "").strip(),
            "uploader": str(descriptor.uploader or descriptor.channel or "").strip(),
            "webpage_url": str(descriptor.webpage_url or descriptor.source_url).strip(),
            "source_url": str(descriptor.source_url or "").strip(),
            "thumbnail_url": str(descriptor.thumbnail_url or "").strip(),
        }
        if descriptor.description and not detail.get("description"):
            detail["description"] = descriptor.description
        if descriptor.duration_seconds is not None:
            media_info = detail.get("media_info") if isinstance(detail.get("media_info"), dict) else {}
            if media_info and not media_info.get("duration_seconds"):
                media_info["duration_seconds"] = round(float(descriptor.duration_seconds), 3)
                detail["media_info"] = media_info

        source_bits = []
        platform = str(detail["remote_source"].get("platform") or "").strip()
        uploader = str(detail["remote_source"].get("uploader") or "").strip()
        if platform:
            source_bits.append(f"来源 {platform}")
        if uploader:
            source_bits.append(f"发布者 {uploader}")
        short_hint = str(file_card.get("short_hint") or detail.get("summary") or "").strip()
        if source_bits:
            suffix = "；".join(source_bits)
            short_hint = f"{short_hint}；{suffix}" if short_hint else suffix
        return {
            "summary_title": summary_title,
            "short_hint": short_hint[:320],
            "detail": detail,
        }

    def _build_remote_media_followup(
        self,
        *,
        items: list[dict[str, Any]],
        failures: list[dict[str, str]],
    ) -> str:
        if items and not failures:
            labels = "、".join(
                str(item.get("attachment_handle") or item.get("summary_title") or item.get("origin_name") or "").strip()
                for item in items[:6]
                if str(item.get("attachment_handle") or item.get("summary_title") or item.get("origin_name") or "").strip()
            )
            return (
                f"你刚刚已经把链接里的媒体素材放进当前材料工作台了：{labels or '已下载媒体'}。"
                "请基于这个既成事实自然回应，不要重复调用 fetch_media_from_url。"
                "如果用户想继续处理内容，可以直接使用 inspect_attachment、inspect_media_info、transcribe_media、convert_media_file 或 send_file。"
            )
        if failures and not items:
            reasons = "；".join(
                f"{entry.get('index_label') or '某个链接'}：{str(entry.get('error') or '').strip()[:120]}"
                for entry in failures[:4]
            )
            install_hint = self._remote_fetch_failure_needs_ytdlp_hint(failures)
            retry_hint = (
                "请自然告诉用户失败原因，并提醒需要安装或修复 yt-dlp 环境。"
                if install_hint
                else "请自然告诉用户失败原因；如果像是网络、权限或平台临时问题，可以建议换公开链接或稍后重试。"
            )
            return (
                "你刚刚尝试通过链接获取媒体，但这次没有成功。"
                + (f"失败情况：{reasons}。" if reasons else "")
                + retry_hint
            )
        success_labels = "、".join(
            str(item.get("attachment_handle") or item.get("summary_title") or item.get("origin_name") or "").strip()
            for item in items[:6]
            if str(item.get("attachment_handle") or item.get("summary_title") or item.get("origin_name") or "").strip()
        )
        reasons = "；".join(
            f"{entry.get('index_label') or '某个链接'}：{str(entry.get('error') or '').strip()[:120]}"
            for entry in failures[:4]
        )
        return (
            f"你刚刚通过链接成功拿到这些媒体：{success_labels or '部分媒体'}。"
            + (f"同时还有部分失败：{reasons}。" if reasons else "")
            + "请自然向用户说明哪些已经到手、哪些还没成功；不要重复调用 fetch_media_from_url。"
        )

    def _remote_fetch_failure_needs_ytdlp_hint(self, failures: list[dict[str, str]]) -> bool:
        for entry in failures:
            error = str(entry.get("error") or "").strip().lower()
            if "yt-dlp 无法导入" in error or "没有可用的公开视频下载器" in error:
                return True
            if "no module named" in error and "yt" in error and "dlp" in error:
                return True
        return False

    def _humanize_remote_fetch_error(self, error: str) -> str:
        text = str(error or "").strip()
        lowered = text.lower()
        if "playlist" in lowered or "合集" in text:
            return "这个链接更像播放列表/合集，当前只支持单个视频或音频链接。"
        if "could not copy" in lowered and "cookie database" in lowered:
            return (
                "yt-dlp 没能复制浏览器 Cookie 数据库，通常是浏览器仍在运行并锁住了 Cookie 文件；"
                "请完全关闭对应浏览器后台进程后重试，或改用 REMOTE_MEDIA_YTDLP_COOKIEFILE 指向导出的 cookies.txt。"
            )
        if "yt-dlp" in lowered:
            return text[:240]
        if "attachment handle allocation failed" in lowered or (
            "unique constraint failed" in lowered and "attachment_inbox_items" in lowered
        ):
            return "这次把链接素材放进工作台时出了点临时问题，请稍后再试一次。"
        if "unsupported url" in lowered or "unsupported" in lowered:
            return "这个链接当前下载器还不认识，可能不是公开可抓取的媒体页面。"
        if "private" in lowered or "login" in lowered or "sign in" in lowered:
            return "这个链接可能需要登录、会员或额外权限，当前不能直接抓取。"
        if "412" in lowered or "precondition failed" in lowered:
            return (
                "远端拒绝了这次媒体信息请求，像是平台风控或前置校验失败；"
                "可以稍后重试、换原始公开链接，或在 .env 配置 REMOTE_MEDIA_YTDLP_COOKIEFILE 后再试。"
            )
        if "403" in lowered or "forbidden" in lowered:
            return "远端拒绝了这次下载请求，可能有权限或地区限制。"
        if "cookie" in lowered and ("not found" in lowered or "no such file" in lowered or "cannot" in lowered):
            return "配置的 yt-dlp Cookie 文件不可用，请检查 REMOTE_MEDIA_YTDLP_COOKIEFILE 路径是否正确。"
        if "404" in lowered or "not found" in lowered:
            return "这个链接对应的页面或媒体文件似乎不存在了。"
        if "timeout" in lowered or "timed out" in lowered:
            return "下载超时了，可能是网络不稳定，或者远端响应太慢。"
        if "过大" in text or "max_filesize" in lowered:
            return text[:240]
        return text[:240] or "链接媒体获取失败，原因未知。"

    def _build_retry_payload(self, item: dict[str, Any], *, previous_error: str = "") -> dict[str, Any]:
        origin_name = self._clean_filename(item.get("origin_name") or "")
        payload = {
            "kind": str(item.get("kind") or "file").strip(),
            "origin_name": origin_name,
            "name": origin_name,
            "file": origin_name,
            "mime_type": str(item.get("mime_type") or "").strip(),
            "file_ext": str(item.get("file_ext") or "").strip(),
            "file_size": self._safe_int(item.get("file_size")),
        }
        storage_relpath = str(item.get("storage_relpath") or "").strip()
        if storage_relpath.lower().startswith("workspace:") and self.workspace_uri_resolver is not None:
            candidate = self.workspace_uri_resolver(storage_relpath)
            if candidate is not None and candidate.exists() and candidate.is_file():
                payload["workspace_uri"] = storage_relpath
                return payload
        if storage_relpath:
            for storage_root in [self.base_dir, *self.legacy_base_dirs]:
                candidate = (storage_root / Path(storage_relpath)).resolve()
                try:
                    candidate.relative_to(storage_root.resolve())
                except Exception:
                    continue
                if candidate.exists() and candidate.is_file():
                    payload["path"] = str(candidate)
                    return payload
        recovered_url = self._extract_url(previous_error)
        if recovered_url:
            payload["url"] = recovered_url
        return payload

    def _extract_url(self, text: str) -> str:
        match = re.search(r"https?://[^\s)]+", str(text or ""))
        return match.group(0).strip().rstrip("。.,，") if match else ""

    def _materialize_attachment_file(self, *, item: dict[str, Any], payload: dict[str, Any]) -> Path | None:
        workspace_uri = str(payload.get("workspace_uri") or "").strip()
        if workspace_uri:
            if self.workspace_uri_resolver is None:
                raise RuntimeError("workspace file resolver is unavailable")
            source_path = self.workspace_uri_resolver(workspace_uri)
            if source_path is None or not source_path.exists() or not source_path.is_file():
                raise FileNotFoundError(workspace_uri)
            return source_path
        if self.ensure_storage_ready is not None:
            self.ensure_storage_ready()
        origin_name = self._clean_filename(
            payload.get("origin_name")
            or payload.get("name")
            or payload.get("file")
            or payload.get("filename")
            or item.get("origin_name")
            or ""
        )
        suffix = self._guess_suffix(origin_name=origin_name, payload=payload, kind=str(item.get("kind") or "file"))
        target_dir = self._workspace_date_dir(item)
        target_dir.mkdir(parents=True, exist_ok=True)
        try:
            target_dir.resolve().relative_to(self.base_dir.resolve())
        except Exception:
            raise RuntimeError("attachment destination escaped the managed workspace") from None
        handle = self._safe_path_part(item.get("attachment_handle") or item.get("attachment_id"))
        target_path = self._available_attachment_path(
            target_dir=target_dir,
            handle=handle,
            origin_name=origin_name,
            suffix=suffix,
        )
        try:
            target_path.resolve(strict=False).relative_to(self.base_dir.resolve())
        except Exception:
            raise RuntimeError("attachment destination escaped the managed workspace") from None

        local_path = str(payload.get("path") or payload.get("local_path") or "").strip()
        if local_path:
            source = Path(local_path)
            if source.exists() and source.is_file():
                shutil.copyfile(source, target_path)
                return target_path

        onebot_cached = self._copy_from_onebot_cache(
            item=item,
            payload=payload,
            target_path=target_path,
            origin_name=origin_name,
        )
        if onebot_cached is not None:
            return onebot_cached

        url = str(payload.get("url") or "").strip()
        if url:
            self._download_to_path(url=url, target_path=target_path)
            return target_path

        raw_file = str(payload.get("file") or "").strip()
        if raw_file:
            possible = Path(raw_file)
            if possible.exists() and possible.is_file():
                shutil.copyfile(possible, target_path)
                return target_path
        return None

    def _copy_from_onebot_cache(
        self,
        *,
        item: dict[str, Any],
        payload: dict[str, Any],
        target_path: Path,
        origin_name: str,
    ) -> Path | None:
        file_token = (
            str(payload.get("file") or "").strip()
            or str(payload.get("file_id") or "").strip()
            or str(payload.get("origin_name") or "").strip()
            or str(origin_name or "").strip()
            or str(item.get("origin_name") or "").strip()
        )
        if not file_token:
            return None

        kind = self._normalize_kind(item.get("kind") or payload.get("kind"))
        endpoints = ["/get_image"] if kind == "image" else []
        endpoints.append("/get_file")
        base_url = str(getattr(config, "QQ_ONEBOT_HTTP_URL", "http://127.0.0.1:3001") or "").strip().rstrip("/")
        if not base_url:
            return None

        timeout = float(getattr(config, "QQ_ATTACHMENT_DOWNLOAD_TIMEOUT", 20.0) or 20.0)
        for endpoint in endpoints:
            try:
                response = requests.post(f"{base_url}{endpoint}", json={"file": file_token}, timeout=timeout)
                response.raise_for_status()
                payload_data = response.json()
            except Exception:
                continue

            if not isinstance(payload_data, dict):
                continue
            if str(payload_data.get("status") or "").lower() not in {"ok", "async"} and int(payload_data.get("retcode") or 0) != 0:
                continue
            data = payload_data.get("data") if isinstance(payload_data.get("data"), dict) else {}
            for key in ("path", "local_path", "file"):
                cached_path = str(data.get(key) or "").strip()
                if not cached_path:
                    continue
                source = Path(cached_path)
                if source.exists() and source.is_file():
                    shutil.copyfile(source, target_path)
                    return target_path
        return None

    def _download_to_path(
        self,
        *,
        url: str,
        target_path: Path,
        timeout: float | None = None,
        max_bytes: int | None = None,
        headers: dict[str, str] | None = None,
    ) -> None:
        timeout_value = float(timeout if timeout is not None else (getattr(config, "QQ_ATTACHMENT_DOWNLOAD_TIMEOUT", 20.0) or 20.0))
        max_bytes_value = int(max_bytes if max_bytes is not None else (getattr(config, "QQ_ATTACHMENT_MAX_BYTES", 20 * 1024 * 1024) or 20 * 1024 * 1024))
        request_headers = headers or {
            "User-Agent": "Mozilla/5.0 AkaneCompanionLab/1.0",
            "Accept": "image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.8",
        }
        with requests.get(url, stream=True, timeout=timeout_value, headers=request_headers) as response:
            response.raise_for_status()
            total = 0
            with target_path.open("wb") as handle:
                for chunk in response.iter_content(chunk_size=64 * 1024):
                    if not chunk:
                        continue
                    total += len(chunk)
                    if total > max_bytes_value:
                        raise RuntimeError(f"附件过大，当前限制为 {max_bytes_value} bytes。")
                    handle.write(chunk)

    def _build_file_card(self, *, source_path: Path, item: dict[str, Any], mime_type: str) -> dict[str, Any]:
        origin_name = self._clean_filename(item.get("origin_name") or source_path.name)
        title = origin_name or source_path.name
        suffix = source_path.suffix.lower()
        file_size = source_path.stat().st_size if source_path.exists() else 0
        if suffix in DOCUMENT_SUFFIXES:
            return self._build_document_card(
                source_path=source_path,
                title=title,
                suffix=suffix,
                mime_type=mime_type,
                file_size=file_size,
            )
        if suffix in MEDIA_SUFFIXES or mime_type.startswith(("audio/", "video/")):
            return self._build_media_card(
                source_path=source_path,
                title=title,
                suffix=suffix,
                mime_type=mime_type,
                file_size=file_size,
            )
        if suffix not in TEXT_SUFFIXES:
            return {
                "summary_title": title,
                "short_hint": f"一个 {suffix or '未知类型'} 文件，约 {file_size} bytes；当前只保留文件名和基本信息，暂未解析内容。",
                "detail": {
                    "summary": "暂不支持直接解析该文件内容。",
                    "file_kind": suffix.lstrip(".") or "binary",
                    "mime_type": mime_type,
                    "file_size": file_size,
                    "text_preview": "",
                },
            }

        text, encoding = self._read_text_preview(source_path)
        lines = text.splitlines()
        preview = text[:4000]
        detail: dict[str, Any] = {
            "summary": f"文本文件 {title}，约 {len(text)} 字符，{len(lines)} 行。",
            "file_kind": self._file_kind_from_suffix(suffix),
            "mime_type": mime_type,
            "encoding": encoding,
            "file_size": file_size,
            "line_count": len(lines),
            "text_preview": preview,
            "preview_is_truncated": len(text) > len(preview),
        }
        self._enrich_structured_text_detail(detail=detail, text=text, suffix=suffix)
        hint_parts = [detail["summary"]]
        if detail.get("headings"):
            hint_parts.append("标题：" + "、".join(list(detail["headings"])[:4]))
        if detail.get("columns"):
            hint_parts.append("列：" + "、".join(list(detail["columns"])[:6]))
        if detail.get("top_level_keys"):
            hint_parts.append("字段：" + "、".join(list(detail["top_level_keys"])[:6]))
        return {
            "summary_title": title,
            "short_hint": " ".join(hint_parts)[:240],
            "detail": detail,
        }

    def _build_media_card(
        self,
        *,
        source_path: Path,
        title: str,
        suffix: str,
        mime_type: str,
        file_size: int,
    ) -> dict[str, Any]:
        media_info = self._probe_media_info(source_path)
        file_kind = suffix.lstrip(".") or ("audio" if mime_type.startswith("audio/") else "video" if mime_type.startswith("video/") else "media")
        detail: dict[str, Any] = {
            "summary": f"媒体文件 {title}，格式 {file_kind}，约 {file_size} bytes。",
            "file_kind": file_kind,
            "mime_type": mime_type,
            "file_size": file_size,
            "text_preview": "",
        }
        hint_parts = [detail["summary"]]
        if media_info:
            detail["media_info"] = media_info
            duration = self._format_duration_label(media_info.get("duration_seconds"))
            if duration:
                hint_parts.append(f"时长 {duration}")
            audio = media_info.get("audio") if isinstance(media_info.get("audio"), dict) else {}
            if audio:
                audio_bits = [
                    str(audio.get("codec") or "未知音频编码"),
                    f"{audio.get('sample_rate')}Hz" if audio.get("sample_rate") else "",
                    f"{audio.get('channels')}声道" if audio.get("channels") else "",
                ]
                hint_parts.append("音频 " + "，".join(bit for bit in audio_bits if bit))
            video = media_info.get("video") if isinstance(media_info.get("video"), dict) else {}
            if video:
                size = ""
                if video.get("width") and video.get("height"):
                    size = f"{video.get('width')}x{video.get('height')}"
                fps = f"{video.get('fps'):g}fps" if isinstance(video.get("fps"), (int, float)) else ""
                hint_parts.append("视频 " + "，".join(bit for bit in (video.get("codec"), size, fps) if bit))
        else:
            detail["media_probe_status"] = "unavailable"
            hint_parts.append("暂未读取到媒体规格")
        return {
            "summary_title": title,
            "short_hint": "；".join(hint_parts)[:240],
            "detail": detail,
        }

    def _build_document_card(
        self,
        *,
        source_path: Path,
        title: str,
        suffix: str,
        mime_type: str,
        file_size: int,
    ) -> dict[str, Any]:
        if suffix == ".docx":
            return self._build_docx_card(
                source_path=source_path,
                title=title,
                mime_type=mime_type,
                file_size=file_size,
            )
        if suffix == ".xlsx":
            return self._build_xlsx_card(
                source_path=source_path,
                title=title,
                mime_type=mime_type,
                file_size=file_size,
            )
        if suffix == ".pdf":
            return self._build_pdf_card(
                source_path=source_path,
                title=title,
                mime_type=mime_type,
                file_size=file_size,
            )
        return self._build_unsupported_document_card(
            title=title,
            suffix=suffix,
            mime_type=mime_type,
            file_size=file_size,
            reason="暂不支持解析该文档格式。",
        )

    def _build_docx_card(self, *, source_path: Path, title: str, mime_type: str, file_size: int) -> dict[str, Any]:
        if importlib.util.find_spec("docx") is None:
            return self._build_unsupported_document_card(
                title=title,
                suffix=".docx",
                mime_type=mime_type,
                file_size=file_size,
                reason="缺少 python-docx，暂时只能看到文件名，不能读取 Word 内容。",
            )
        try:
            from docx import Document  # type: ignore

            document = Document(str(source_path))
            paragraphs = [str(paragraph.text or "").strip() for paragraph in document.paragraphs]
            paragraphs = [text for text in paragraphs if text]
            headings: list[str] = []
            for paragraph in document.paragraphs:
                text = str(paragraph.text or "").strip()
                style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "")
                if text and style_name.lower().startswith("heading"):
                    headings.append(text[:80])
                if len(headings) >= 8:
                    break

            table_summaries: list[dict[str, Any]] = []
            for index, table in enumerate(document.tables[:5], start=1):
                rows = []
                for row in table.rows[:3]:
                    rows.append([str(cell.text or "").strip()[:80] for cell in row.cells[:8]])
                table_summaries.append(
                    {
                        "index": index,
                        "row_count": len(table.rows),
                        "column_count": len(table.columns),
                        "preview_rows": rows,
                    }
                )

            preview = "\n".join(paragraphs[:20])[:4000]
            detail = {
                "summary": f"Word 文档 {title}，约 {len(paragraphs)} 段，{len(document.tables)} 个表格。",
                "file_kind": "docx",
                "mime_type": mime_type,
                "file_size": file_size,
                "paragraph_count": len(paragraphs),
                "table_count": len(document.tables),
                "headings": headings,
                "tables": table_summaries,
                "text_preview": preview,
                "preview_is_truncated": len(paragraphs) > 20 or len("\n".join(paragraphs[:20])) > 4000 or len(document.tables) > 5,
            }
            hint_parts = [detail["summary"]]
            if headings:
                hint_parts.append("标题：" + "、".join(headings[:4]))
            elif paragraphs:
                hint_parts.append("开头：" + " ".join(paragraphs[:2])[:100])
            return {
                "summary_title": title,
                "short_hint": " ".join(hint_parts)[:240],
                "detail": detail,
            }
        except Exception as exc:
            return self._build_unsupported_document_card(
                title=title,
                suffix=".docx",
                mime_type=mime_type,
                file_size=file_size,
                reason=f"Word 文档解析失败：{str(exc)[:160]}",
            )

    def _build_xlsx_card(self, *, source_path: Path, title: str, mime_type: str, file_size: int) -> dict[str, Any]:
        if importlib.util.find_spec("openpyxl") is None:
            return self._build_unsupported_document_card(
                title=title,
                suffix=".xlsx",
                mime_type=mime_type,
                file_size=file_size,
                reason="缺少 openpyxl，暂时只能看到文件名，不能读取 Excel 内容。",
            )
        try:
            from openpyxl import load_workbook  # type: ignore

            workbook = load_workbook(filename=str(source_path), read_only=True, data_only=True)
            sheet_cards: list[dict[str, Any]] = []
            for sheet_name in workbook.sheetnames[:5]:
                sheet = workbook[sheet_name]
                preview_rows: list[list[str]] = []
                for row in sheet.iter_rows(min_row=1, max_row=5, values_only=True):
                    values = [self._cell_to_text(value)[:80] for value in list(row)[:10]]
                    if any(values):
                        preview_rows.append(values)
                columns = [cell for cell in (preview_rows[0] if preview_rows else []) if cell]
                sheet_cards.append(
                    {
                        "name": str(sheet_name),
                        "max_row": int(sheet.max_row or 0),
                        "max_column": int(sheet.max_column or 0),
                        "columns": columns[:12],
                        "preview_rows": preview_rows,
                    }
                )
            workbook.close()

            sheet_names = [item["name"] for item in sheet_cards]
            first_columns = []
            if sheet_cards:
                first_columns = list(sheet_cards[0].get("columns") or [])
            detail = {
                "summary": f"Excel 工作簿 {title}，包含 {len(sheet_names)} 个工作表。",
                "file_kind": "xlsx",
                "mime_type": mime_type,
                "file_size": file_size,
                "sheet_count": len(sheet_names),
                "sheet_names": sheet_names,
                "sheets": sheet_cards,
                "columns": first_columns,
                "text_preview": self._render_xlsx_preview(sheet_cards),
                "preview_is_truncated": any(
                    int(sheet.get("max_row") or 0) > 5 or int(sheet.get("max_column") or 0) > 10
                    for sheet in sheet_cards
                ) or len(sheet_names) > 5,
            }
            hint_parts = [detail["summary"]]
            if sheet_names:
                hint_parts.append("工作表：" + "、".join(sheet_names[:5]))
            if first_columns:
                hint_parts.append("首表列：" + "、".join(first_columns[:6]))
            return {
                "summary_title": title,
                "short_hint": " ".join(hint_parts)[:240],
                "detail": detail,
            }
        except Exception as exc:
            return self._build_unsupported_document_card(
                title=title,
                suffix=".xlsx",
                mime_type=mime_type,
                file_size=file_size,
                reason=f"Excel 文件解析失败：{str(exc)[:160]}",
            )

    def _build_pdf_card(self, *, source_path: Path, title: str, mime_type: str, file_size: int) -> dict[str, Any]:
        if importlib.util.find_spec("pypdf") is None:
            return self._build_unsupported_document_card(
                title=title,
                suffix=".pdf",
                mime_type=mime_type,
                file_size=file_size,
                reason="缺少 pypdf，暂时只能看到 PDF 文件名；安装依赖后可读取页数和文本层预览。",
            )
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(str(source_path))
            page_count = len(reader.pages)
            page_texts: list[str] = []
            for page in reader.pages[:3]:
                page_texts.append(str(page.extract_text() or "").strip())
            text_preview = "\n\n".join(text for text in page_texts if text)[:4000]
            has_text = bool(text_preview.strip())
            summary = f"PDF 文档 {title}，共 {page_count} 页。"
            if not has_text:
                summary += " 目前没有读到明显文本层，可能是扫描件或图片型 PDF。"
            detail = {
                "summary": summary,
                "file_kind": "pdf",
                "mime_type": mime_type,
                "file_size": file_size,
                "page_count": page_count,
                "has_text_layer": has_text,
                "text_preview": text_preview,
                "preview_is_truncated": page_count > 3 or len("\n\n".join(text for text in page_texts if text)) > 4000,
            }
            return {
                "summary_title": title,
                "short_hint": summary[:240],
                "detail": detail,
            }
        except Exception as exc:
            return self._build_unsupported_document_card(
                title=title,
                suffix=".pdf",
                mime_type=mime_type,
                file_size=file_size,
                reason=f"PDF 解析失败：{str(exc)[:160]}",
            )

    def _build_unsupported_document_card(
        self,
        *,
        title: str,
        suffix: str,
        mime_type: str,
        file_size: int,
        reason: str,
    ) -> dict[str, Any]:
        file_kind = suffix.lstrip(".") or "document"
        return {
            "summary_title": title,
            "short_hint": f"{file_kind.upper()} 文件，约 {file_size} bytes；{reason}",
            "detail": {
                "summary": reason,
                "file_kind": file_kind,
                "mime_type": mime_type,
                "file_size": file_size,
                "parser_status": "unsupported_or_unavailable",
                "text_preview": "",
            },
        }

    def _cell_to_text(self, value: Any) -> str:
        if value is None:
            return ""
        return str(value).strip()

    def _render_xlsx_preview(self, sheets: list[dict[str, Any]]) -> str:
        lines: list[str] = []
        for sheet in sheets[:3]:
            name = str(sheet.get("name") or "Sheet")
            lines.append(f"[{name}]")
            for row in list(sheet.get("preview_rows") or [])[:5]:
                if isinstance(row, list):
                    lines.append(" | ".join(str(cell) for cell in row if str(cell).strip()))
        return "\n".join(lines)[:4000]

    def _read_text_preview(self, path: Path) -> tuple[str, str]:
        max_bytes = int(getattr(config, "QQ_TEXT_ATTACHMENT_MAX_READ_BYTES", 256 * 1024) or 256 * 1024)
        payload = path.read_bytes()[:max_bytes]
        for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
            try:
                return payload.decode(encoding), encoding
            except UnicodeDecodeError:
                continue
        return payload.decode("utf-8", errors="replace"), "utf-8-replace"

    def _enrich_structured_text_detail(self, *, detail: dict[str, Any], text: str, suffix: str) -> None:
        if suffix in {".md", ".markdown"}:
            headings = []
            for line in text.splitlines():
                stripped = line.strip()
                if stripped.startswith("#"):
                    heading = stripped.lstrip("#").strip()
                    if heading:
                        headings.append(heading[:80])
                if len(headings) >= 8:
                    break
            if headings:
                detail["headings"] = headings
            return

        if suffix == ".json":
            try:
                payload = json.loads(text)
                if isinstance(payload, dict):
                    detail["top_level_keys"] = [str(key)[:60] for key in list(payload.keys())[:12]]
                elif isinstance(payload, list):
                    detail["summary"] += f" 顶层是数组，约 {len(payload)} 项。"
            except Exception:
                return
            return

        if suffix == ".csv":
            try:
                rows = list(csv.reader(text.splitlines()[:3]))
                if rows:
                    detail["columns"] = [str(item).strip()[:60] for item in rows[0] if str(item).strip()][:20]
            except Exception:
                return

    def _mark_failed(self, item: dict[str, Any], error: str, *, timestamp: int) -> None:
        readable_error = self._humanize_failure(error, kind=str(item.get("kind") or ""))
        self.store.update_attachment_inbox_item(
            profile_user_id=str(item.get("profile_user_id") or ""),
            session_id=str(item.get("session_id") or ""),
            attachment_id=str(item.get("attachment_id") or ""),
            status="failed",
            error_message=str(error or "")[:500],
            short_hint=readable_error,
            detail={
                "failure": {
                    "reason": readable_error,
                    "raw_error": str(error or "")[:500],
                    "failed_at": timestamp,
                }
            },
            updated_at=timestamp,
        )

    def _mark_image_ready_from_observation(
        self,
        *,
        item: dict[str, Any],
        observation: dict[str, Any],
        timestamp: int,
    ) -> None:
        card = observation.get("observation") if isinstance(observation.get("observation"), dict) else {}
        summary = str(card.get("summary") or observation.get("summary") or "").strip()
        summary_title = (
            str(card.get("summary_title") or "").strip()
            or str(item.get("origin_name") or "").strip()
            or str(item.get("attachment_handle") or "").strip()
            or "图片附件"
        )
        self.attachment_service.mark_ready(
            profile_user_id=str(item.get("profile_user_id") or ""),
            session_id=str(item.get("session_id") or ""),
            attachment_id=str(item.get("attachment_id") or ""),
            summary_title=summary_title,
            short_hint=summary,
            detail=card,
            timestamp=timestamp,
        )

    def _humanize_failure(self, error: str, *, kind: str = "") -> str:
        text = str(error or "").strip()
        lowered = text.lower()
        normalized_kind = self._normalize_kind(kind)
        if "bad request" in lowered or "400 client error" in lowered:
            if normalized_kind == "image":
                return "QQ 临时图片链接返回 400；可以稍后重试，系统会优先尝试从 NapCat 本地缓存读取。"
            return "QQ 临时文件链接返回 400；可以稍后重试，系统会优先尝试从 NapCat 本地缓存读取。"
        if "视觉模型" in text or "vision" in lowered:
            return text[:240]
        if "timeout" in lowered or "timed out" in lowered:
            return "读取附件超时，可能是网络、NapCat 缓存或上游接口暂时不可用。"
        if "附件过大" in text or "too large" in lowered:
            return text[:240]
        if "没有可下载" in text:
            return "没有拿到可下载地址或本地缓存路径，可能需要重新发送附件。"
        return text[:240] or "附件处理失败，原因未知。"

    def _probe_media_info(self, source_path: Path) -> dict[str, Any] | None:
        ffprobe_path = shutil.which("ffprobe")
        if not ffprobe_path:
            return None
        command = [
            ffprobe_path,
            "-v",
            "error",
            "-print_format",
            "json",
            "-show_format",
            "-show_streams",
            str(source_path),
        ]
        try:
            completed = subprocess.run(command, capture_output=True, text=True, timeout=15, check=False)
        except Exception:
            return None
        if completed.returncode != 0:
            return None
        try:
            parsed = json.loads(str(completed.stdout or "{}"))
        except Exception:
            return None
        if not isinstance(parsed, dict):
            return None
        return self._normalize_media_info(parsed, source_path=source_path)

    def _normalize_media_info(self, data: dict[str, Any], *, source_path: Path) -> dict[str, Any]:
        format_info = data.get("format") if isinstance(data.get("format"), dict) else {}
        streams = data.get("streams") if isinstance(data.get("streams"), list) else []
        audio_streams = []
        video_streams = []
        for stream in streams:
            if not isinstance(stream, dict):
                continue
            kind = str(stream.get("codec_type") or "").strip().lower()
            if kind == "audio":
                audio_streams.append(
                    {
                        "codec": str(stream.get("codec_name") or "").strip(),
                        "sample_rate": self._safe_int(stream.get("sample_rate")),
                        "channels": self._safe_int(stream.get("channels")),
                        "bit_rate": self._safe_int(stream.get("bit_rate")),
                        "duration_seconds": self._safe_float(stream.get("duration")),
                    }
                )
            elif kind == "video":
                video_streams.append(
                    {
                        "codec": str(stream.get("codec_name") or "").strip(),
                        "width": self._safe_int(stream.get("width")),
                        "height": self._safe_int(stream.get("height")),
                        "fps": self._parse_frame_rate(stream.get("avg_frame_rate") or stream.get("r_frame_rate")),
                        "bit_rate": self._safe_int(stream.get("bit_rate")),
                        "duration_seconds": self._safe_float(stream.get("duration")),
                    }
                )
        duration = self._safe_float(format_info.get("duration"))
        if duration is None:
            duration_values = [
                item.get("duration_seconds")
                for item in [*audio_streams, *video_streams]
                if isinstance(item.get("duration_seconds"), (int, float))
            ]
            duration = max(duration_values) if duration_values else None
        return {
            "format_name": str(format_info.get("format_name") or source_path.suffix.lower().lstrip(".")).strip(),
            "duration_seconds": round(float(duration), 3) if duration is not None else None,
            "file_size": self._safe_int(format_info.get("size")),
            "bit_rate": self._safe_int(format_info.get("bit_rate")),
            "audio": audio_streams[0] if audio_streams else None,
            "video": video_streams[0] if video_streams else None,
            "audio_streams": audio_streams[:4],
            "video_streams": video_streams[:4],
        }

    def _parse_frame_rate(self, value: Any) -> float | None:
        text = str(value or "").strip()
        if not text or text == "0/0":
            return None
        if "/" in text:
            numerator, denominator = text.split("/", 1)
            try:
                den = float(denominator)
                if den == 0:
                    return None
                return round(float(numerator) / den, 3)
            except Exception:
                return None
        try:
            return round(float(text), 3)
        except Exception:
            return None

    def _safe_float(self, value: Any) -> float | None:
        try:
            number = float(str(value).strip())
        except Exception:
            return None
        return number if number >= 0 else None

    def _format_duration_label(self, value: Any) -> str:
        if not isinstance(value, (int, float)):
            return ""
        seconds = max(0, int(round(float(value))))
        hours, remainder = divmod(seconds, 3600)
        minutes, secs = divmod(remainder, 60)
        if hours:
            return f"{hours}:{minutes:02d}:{secs:02d}"
        return f"{minutes}:{secs:02d}"

    def _storage_relpath(self, path: Path) -> str:
        try:
            return str(path.relative_to(self.base_dir)).replace("\\", "/")
        except ValueError:
            return str(path)

    def _workspace_date_dir(self, item: dict[str, Any]) -> Path:
        timestamp = self._safe_int(item.get("created_at")) or int(time.time())
        date_slug = time.strftime("%Y-%m-%d", time.localtime(timestamp))
        target_dir = self.base_dir / date_slug
        try:
            target_dir.resolve(strict=False).relative_to(self.base_dir.resolve())
        except Exception:
            raise RuntimeError("attachment destination escaped the managed workspace") from None
        return target_dir

    def _available_attachment_path(
        self,
        *,
        target_dir: Path,
        handle: str,
        origin_name: str,
        suffix: str,
    ) -> Path:
        stem = self._available_attachment_stem(
            target_dir=target_dir,
            handle=handle,
            origin_name=origin_name,
            suffix=suffix,
        )
        return target_dir / f"{stem}{suffix}"

    def _available_attachment_stem(
        self,
        *,
        target_dir: Path,
        handle: str,
        origin_name: str,
        suffix: str = "",
    ) -> str:
        clean_suffix = str(suffix or "").lower()
        safe_origin_name = self._safe_path_part(origin_name) if origin_name else ""
        origin_stem = Path(safe_origin_name).stem if safe_origin_name else ""
        base_stem = f"{handle}__{origin_stem}" if origin_stem else handle
        candidate_stem = base_stem
        sequence = 2
        while any(target_dir.glob(f"{candidate_stem}.*")) or (
            clean_suffix and (target_dir / f"{candidate_stem}{clean_suffix}").exists()
        ):
            candidate_stem = f"{base_stem}_{sequence}"
            sequence += 1
        return candidate_stem

    def _guess_suffix(self, *, origin_name: str, payload: dict[str, Any], kind: str) -> str:
        suffix = Path(origin_name).suffix.lower()
        if suffix:
            return suffix
        url = str(payload.get("url") or "").strip()
        if url:
            path = unquote(urlparse(url).path)
            suffix = Path(path).suffix.lower()
            if suffix:
                return suffix
        mime_type = str(payload.get("mime_type") or "").strip()
        guessed = mimetypes.guess_extension(mime_type) if mime_type else ""
        if guessed:
            return guessed
        return ".png" if kind == "image" else ".bin"

    def _normalize_kind(self, value: Any) -> str:
        kind = str(value or "").strip().lower()
        if kind in {"photo", "picture", "pic", "img"}:
            return "image"
        if kind in {"doc", "text", "txt", "pdf"}:
            return "document"
        if kind in {"music", "song", "voice", "record"}:
            return "audio"
        if kind in {"image", "document", "audio", "file"}:
            return kind
        return "file"

    def _infer_local_kind(self, *, source_path: Path, mime_type: str) -> str:
        suffix = source_path.suffix.lower()
        normalized_mime = str(mime_type or "").strip().lower()
        if suffix in IMAGE_SUFFIXES or normalized_mime.startswith("image/"):
            return "image"
        if suffix in AUDIO_MEDIA_SUFFIXES or normalized_mime.startswith("audio/"):
            return "audio"
        if suffix in TEXT_SUFFIXES or suffix in DOCUMENT_SUFFIXES or normalized_mime.startswith("text/"):
            return "document"
        return "file"

    def _file_kind_from_suffix(self, suffix: str) -> str:
        return {
            ".md": "markdown",
            ".markdown": "markdown",
            ".json": "json",
            ".toml": "toml",
            ".yaml": "yaml",
            ".yml": "yaml",
            ".csv": "csv",
            ".py": "python",
            ".js": "javascript",
            ".ts": "typescript",
            ".log": "log",
        }.get(suffix, suffix.lstrip(".") or "text")

    def _clean_filename(self, value: Any) -> str:
        text = str(value or "").strip().replace("\\", "/").split("/")[-1].strip()
        return text[:160]

    def _safe_path_part(self, value: Any) -> str:
        text = str(value or "").strip() or "unknown"
        safe = "".join(char if char.isalnum() or char in {"-", "_", "."} else "_" for char in text)
        return safe[:80] or "unknown"

    def _safe_int(self, value: Any) -> int:
        try:
            return max(0, int(value or 0))
        except Exception:
            return 0
