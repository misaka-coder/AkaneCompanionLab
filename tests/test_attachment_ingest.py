from __future__ import annotations

import tempfile
import time
import unittest
import importlib.util
import types
from pathlib import Path
from typing import Any
from unittest.mock import patch
import sys

from companion_v01.attachment_inbox import AttachmentInboxService
from companion_v01.attachment_ingest import AttachmentIngestService, RemoteMediaDescriptor
from companion_v01.store import MemoryStore


class FakeVisionService:
    def __init__(self, store: MemoryStore) -> None:
        self.store = store
        self.scheduled: list[dict[str, Any]] = []

    def schedule_attachment_image_observation(
        self,
        *,
        attachment: dict[str, Any] | None,
        source_path: Path,
    ) -> dict[str, Any]:
        assert attachment is not None
        self.scheduled.append({"attachment": attachment, "source_path": source_path})
        self.store.update_attachment_inbox_item(
            profile_user_id=str(attachment.get("profile_user_id") or ""),
            session_id=str(attachment.get("session_id") or ""),
            attachment_id=str(attachment.get("attachment_id") or ""),
            status="ready",
            summary_title="窗边小猫",
            short_hint="一张白猫趴在窗边的照片。",
            detail={
                "type": "attachment_image_observation",
                "summary_title": "窗边小猫",
                "summary": "一张白猫趴在窗边的照片。",
                "entities": ["白猫", "窗边"],
                "mood_tags": ["安静"],
            },
            updated_at=200,
        )
        return {"status": "pending"}


class AttachmentIngestTests(unittest.TestCase):
    def test_local_text_file_is_registered_and_parsed(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source" / "计划.md"
            source.parent.mkdir(parents=True, exist_ok=True)
            source.write_text("# 今日计划\n\n- 写代码\n- 测试 QQ 附件\n", encoding="utf-8")

            store = MemoryStore(root / "db")
            inbox = AttachmentInboxService(store=store)
            service = AttachmentIngestService(
                base_dir=root / "attachments",
                store=store,
                attachment_service=inbox,
                vision_service=FakeVisionService(store),  # type: ignore[arg-type]
            )

            created = service.ingest_qq_attachments(
                profile_user_id="master",
                session_id="qq_pri_1",
                attachments=[
                    {
                        "kind": "document",
                        "origin_name": "计划.md",
                        "path": str(source),
                    }
                ],
                timestamp=100,
            )
            self.assertEqual(len(created), 1)

            item = self._wait_for_status(
                store,
                profile_user_id="master",
                session_id="qq_pri_1",
                status="ready",
            )

            self.assertEqual(item["attachment_handle"], "file_001")
            self.assertEqual(item["summary_title"], "计划.md")
            self.assertIn("文本文件", item["short_hint"])
            self.assertIn("今日计划", item["detail"].get("headings") or [])
            self.assertIn("测试 QQ 附件", item["detail"].get("text_preview") or "")
            self.assertTrue((root / "attachments" / item["storage_relpath"]).exists())

    @unittest.skipUnless(importlib.util.find_spec("docx") is not None, "python-docx is not installed")
    def test_local_docx_file_is_registered_and_parsed(self) -> None:
        from docx import Document  # type: ignore

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source" / "计划.docx"
            source.parent.mkdir(parents=True, exist_ok=True)
            document = Document()
            document.add_heading("学习目标", level=1)
            document.add_paragraph("每天复习数学和英语。")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "时间"
            table.cell(0, 1).text = "任务"
            table.cell(1, 0).text = "8:00"
            table.cell(1, 1).text = "背单词"
            document.save(str(source))

            store = MemoryStore(root / "db")
            inbox = AttachmentInboxService(store=store)
            service = AttachmentIngestService(
                base_dir=root / "attachments",
                store=store,
                attachment_service=inbox,
                vision_service=FakeVisionService(store),  # type: ignore[arg-type]
            )
            service.ingest_qq_attachments(
                profile_user_id="master",
                session_id="qq_pri_1",
                attachments=[{"kind": "document", "origin_name": "计划.docx", "path": str(source)}],
                timestamp=100,
            )

            item = self._wait_for_status(store, profile_user_id="master", session_id="qq_pri_1", status="ready")

            self.assertEqual(item["summary_title"], "计划.docx")
            self.assertEqual(item["detail"]["file_kind"], "docx")
            self.assertIn("学习目标", item["detail"].get("headings") or [])
            self.assertIn("每天复习数学和英语", item["detail"].get("text_preview") or "")
            self.assertEqual(item["detail"]["table_count"], 1)

    @unittest.skipUnless(importlib.util.find_spec("openpyxl") is not None, "openpyxl is not installed")
    def test_local_xlsx_file_is_registered_and_parsed(self) -> None:
        from openpyxl import Workbook  # type: ignore

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source" / "成绩.xlsx"
            source.parent.mkdir(parents=True, exist_ok=True)
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "成绩表"
            sheet.append(["姓名", "数学", "英语"])
            sheet.append(["Akane", 98, 95])
            workbook.save(str(source))

            store = MemoryStore(root / "db")
            inbox = AttachmentInboxService(store=store)
            service = AttachmentIngestService(
                base_dir=root / "attachments",
                store=store,
                attachment_service=inbox,
                vision_service=FakeVisionService(store),  # type: ignore[arg-type]
            )
            service.ingest_qq_attachments(
                profile_user_id="master",
                session_id="qq_pri_1",
                attachments=[{"kind": "document", "origin_name": "成绩.xlsx", "path": str(source)}],
                timestamp=100,
            )

            item = self._wait_for_status(store, profile_user_id="master", session_id="qq_pri_1", status="ready")

            self.assertEqual(item["summary_title"], "成绩.xlsx")
            self.assertEqual(item["detail"]["file_kind"], "xlsx")
            self.assertIn("成绩表", item["detail"].get("sheet_names") or [])
            self.assertIn("姓名", item["detail"].get("columns") or [])
            self.assertIn("Akane", item["detail"].get("text_preview") or "")

    def test_local_image_is_registered_and_delegated_to_vision_service(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source" / "cat.png"
            source.parent.mkdir(parents=True, exist_ok=True)
            source.write_bytes(b"\x89PNG\r\n\x1a\nstub")

            store = MemoryStore(root / "db")
            inbox = AttachmentInboxService(store=store)
            fake_vision = FakeVisionService(store)
            service = AttachmentIngestService(
                base_dir=root / "attachments",
                store=store,
                attachment_service=inbox,
                vision_service=fake_vision,  # type: ignore[arg-type]
            )

            created = service.ingest_qq_attachments(
                profile_user_id="master",
                session_id="qq_pri_1",
                attachments=[
                    {
                        "kind": "image",
                        "origin_name": "cat.png",
                        "path": str(source),
                    }
                ],
                timestamp=100,
            )
            self.assertEqual(len(created), 1)

            item = self._wait_for_status(
                store,
                profile_user_id="master",
                session_id="qq_pri_1",
                status="ready",
            )

            self.assertEqual(item["attachment_handle"], "img_001")
            self.assertEqual(item["summary_title"], "窗边小猫")
            self.assertEqual(item["detail"]["entities"], ["白猫", "窗边"])
            self.assertEqual(len(fake_vision.scheduled), 1)
            self.assertTrue(fake_vision.scheduled[0]["source_path"].exists())

    def test_image_uses_onebot_cache_before_direct_url_download(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            cached = root / "napcat-cache" / "cat.jpg"
            cached.parent.mkdir(parents=True, exist_ok=True)
            cached.write_bytes(b"cached image payload")

            store = MemoryStore(root / "db")
            inbox = AttachmentInboxService(store=store)
            fake_vision = FakeVisionService(store)
            service = AttachmentIngestService(
                base_dir=root / "attachments",
                store=store,
                attachment_service=inbox,
                vision_service=fake_vision,  # type: ignore[arg-type]
            )

            class FakeResponse:
                def raise_for_status(self) -> None:
                    return None

                def json(self) -> dict[str, Any]:
                    return {
                        "status": "ok",
                        "retcode": 0,
                        "data": {
                            "file": str(cached),
                            "url": "https://gchat.qpic.cn/download?bad=true",
                        },
                    }

            with patch("companion_v01.attachment_ingest.requests.post", return_value=FakeResponse()) as post_mock:
                with patch("companion_v01.attachment_ingest.requests.get") as get_mock:
                    created = service.ingest_qq_attachments(
                        profile_user_id="master",
                        session_id="qq_pri_1",
                        attachments=[
                            {
                                "kind": "image",
                                "file": "cat.jpg",
                                "origin_name": "cat.jpg",
                                "url": "https://gchat.qpic.cn/download?bad=true",
                            }
                        ],
                        timestamp=100,
                    )

                    self.assertEqual(len(created), 1)
                    item = self._wait_for_status(
                        store,
                        profile_user_id="master",
                        session_id="qq_pri_1",
                        status="ready",
                    )

            post_mock.assert_called()
            get_mock.assert_not_called()
            saved_path = root / "attachments" / item["storage_relpath"]
            self.assertEqual(saved_path.read_bytes(), b"cached image payload")
            self.assertEqual(item["summary_title"], "窗边小猫")

    def test_retry_failed_image_reuses_original_handle(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            cached = root / "napcat-cache" / "retry.jpg"
            cached.parent.mkdir(parents=True, exist_ok=True)
            cached.write_bytes(b"retry image payload")

            store = MemoryStore(root / "db")
            inbox = AttachmentInboxService(store=store)
            fake_vision = FakeVisionService(store)
            service = AttachmentIngestService(
                base_dir=root / "attachments",
                store=store,
                attachment_service=inbox,
                vision_service=fake_vision,  # type: ignore[arg-type]
            )
            failed = inbox.create_pending(
                profile_user_id="master",
                session_id="qq_pri_1",
                source="qq",
                kind="image",
                origin_name="retry.jpg",
                timestamp=100,
            )
            store.update_attachment_inbox_item(
                profile_user_id="master",
                session_id="qq_pri_1",
                attachment_id=failed["attachment_id"],
                status="failed",
                error_message="400 Client Error: Bad Request for url: https://gchat.qpic.cn/download?x=1",
                updated_at=110,
            )

            class FakeResponse:
                def raise_for_status(self) -> None:
                    return None

                def json(self) -> dict[str, Any]:
                    return {"status": "ok", "retcode": 0, "data": {"file": str(cached)}}

            with patch("companion_v01.attachment_ingest.requests.post", return_value=FakeResponse()):
                result = service.retry_attachment(
                    profile_user_id="master",
                    session_id="qq_pri_1",
                    target="img_001",
                    kind="image",
                    timestamp=200,
                )

                self.assertTrue(result["ok"])
                self.assertEqual(result["status"], "retry_started")
                item = self._wait_for_status(
                    store,
                    profile_user_id="master",
                    session_id="qq_pri_1",
                    status="ready",
                )

            self.assertEqual(item["attachment_handle"], "img_001")
            self.assertEqual(item["attachment_id"], failed["attachment_id"])
            self.assertEqual(item["summary_title"], "窗边小猫")
            self.assertTrue((root / "attachments" / item["storage_relpath"]).exists())

    def test_local_media_file_gets_lightweight_media_card(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source" / "clip.mp4"
            source.parent.mkdir(parents=True, exist_ok=True)
            source.write_bytes(b"fake video payload")

            store = MemoryStore(root / "db")
            inbox = AttachmentInboxService(store=store)
            service = AttachmentIngestService(
                base_dir=root / "attachments",
                store=store,
                attachment_service=inbox,
                vision_service=FakeVisionService(store),  # type: ignore[arg-type]
            )

            def fake_run(command, **kwargs):
                class Result:
                    returncode = 0
                    stderr = ""
                    stdout = """
                    {
                      "format": {
                        "format_name": "mov,mp4,m4a,3gp,3g2,mj2",
                        "duration": "65.5",
                        "size": "2048",
                        "bit_rate": "256000"
                      },
                      "streams": [
                        {"codec_type": "video", "codec_name": "h264", "width": 1280, "height": 720, "avg_frame_rate": "30/1"},
                        {"codec_type": "audio", "codec_name": "aac", "sample_rate": "44100", "channels": 2, "bit_rate": "128000"}
                      ]
                    }
                    """

                return Result()

            with patch("companion_v01.attachment_ingest.shutil.which", return_value="ffprobe"), patch(
                "companion_v01.attachment_ingest.subprocess.run",
                side_effect=fake_run,
            ):
                created = service.ingest_qq_attachments(
                    profile_user_id="master",
                    session_id="qq_pri_1",
                    attachments=[
                        {
                            "kind": "file",
                            "origin_name": "clip.mp4",
                            "path": str(source),
                        }
                    ],
                    timestamp=100,
                )
                self.assertEqual(len(created), 1)
                item = self._wait_for_status(
                    store,
                    profile_user_id="master",
                    session_id="qq_pri_1",
                    status="ready",
                )

            media_info = item["detail"].get("media_info") or {}
            self.assertEqual(item["attachment_handle"], "file_001")
            self.assertIn("时长 1:06", item["short_hint"])
            self.assertEqual(media_info["duration_seconds"], 65.5)
            self.assertEqual(media_info["audio"]["sample_rate"], 44100)
            self.assertEqual(media_info["video"]["width"], 1280)
            prompt = inbox.build_prompt_context(profile_user_id="master", session_id="qq_pri_1")
            self.assertIn("媒体信息", prompt)
            self.assertIn("1280x720", prompt)

    def test_fetch_media_from_url_registers_downloaded_media_into_workspace(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            store = MemoryStore(root / "db")
            inbox = AttachmentInboxService(store=store, base_dir=root / "attachments")
            service = AttachmentIngestService(
                base_dir=root / "attachments",
                store=store,
                attachment_service=inbox,
                vision_service=FakeVisionService(store),  # type: ignore[arg-type]
            )

            descriptor = RemoteMediaDescriptor(
                source_url="https://example.com/watch?v=1",
                webpage_url="https://example.com/watch?v=1",
                title="测试视频",
                ext="mp4",
                mime_type="video/mp4",
                kind="file",
                download_mode="yt_dlp",
                extractor="ExampleVideo",
                extractor_key="ExampleVideo",
                uploader="AkaneChannel",
                duration_seconds=12.5,
            )

            def fake_download(*, item, descriptor):
                target = root / "attachments" / "master" / "qq_pri_1" / "file_001.mp4"
                target.parent.mkdir(parents=True, exist_ok=True)
                target.write_bytes(b"remote video payload")
                return target

            def fake_build_card(*, source_path, item, mime_type):
                return {
                    "summary_title": "测试视频.mp4",
                    "short_hint": "媒体文件 测试视频.mp4，格式 mp4，约 20 bytes。",
                    "detail": {
                        "summary": "媒体文件 测试视频.mp4，格式 mp4，约 20 bytes。",
                        "file_kind": "mp4",
                        "mime_type": mime_type,
                        "file_size": 20,
                        "media_info": {
                            "format_name": "mp4",
                            "duration_seconds": 12.5,
                            "file_size": 20,
                            "audio": {"codec": "aac", "sample_rate": 44100, "channels": 2},
                            "video": {"codec": "h264", "width": 1280, "height": 720, "fps": 30.0},
                        },
                    },
                }

            with patch.object(service, "_fetch_remote_media_descriptor", return_value=descriptor):
                with patch.object(service, "_download_remote_media", side_effect=fake_download):
                    with patch.object(service, "_build_file_card", side_effect=fake_build_card):
                        result = service.fetch_media_from_urls(
                            profile_user_id="master",
                            session_id="qq_pri_1",
                            urls=["https://example.com/watch?v=1"],
                            timestamp=100,
                        )

            self.assertTrue(result["ok"])
            self.assertEqual(len(result["items"]), 1)
            item = result["items"][0]
            self.assertEqual(item["attachment_handle"], "file_001")
            self.assertEqual(item["status"], "ready")
            self.assertEqual(item["source"], "remote_url")
            self.assertEqual(item["detail"]["remote_source"]["platform"], "ExampleVideo")
            self.assertEqual(item["detail"]["remote_source"]["uploader"], "AkaneChannel")

            prompt = inbox.build_prompt_context(profile_user_id="master", session_id="qq_pri_1")
            self.assertIn("平台 ExampleVideo", prompt)
            self.assertIn("发布者 AkaneChannel", prompt)
            self.assertIn("媒体信息", prompt)

    def test_fetch_media_from_url_clears_stale_failed_entry_for_same_url(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            store = MemoryStore(root / "db")
            inbox = AttachmentInboxService(store=store, base_dir=root / "attachments")
            service = AttachmentIngestService(
                base_dir=root / "attachments",
                store=store,
                attachment_service=inbox,
                vision_service=FakeVisionService(store),  # type: ignore[arg-type]
            )
            url = "https://example.com/watch?v=stale"
            stale = store.add_attachment_inbox_item(
                profile_user_id="master",
                session_id="qq_pri_1",
                source="remote_url",
                kind="file",
                status="failed",
                origin_name="旧失败.mp4",
                source_event_id=url,
                error_message="Requested format is not available",
                timestamp=90,
            )
            descriptor = RemoteMediaDescriptor(
                source_url=url,
                webpage_url=url,
                title="测试视频",
                ext="mp4",
                mime_type="video/mp4",
                kind="file",
                download_mode="yt_dlp",
                extractor="ExampleVideo",
                extractor_key="ExampleVideo",
            )

            def fake_download(*, item, descriptor):
                target = root / "attachments" / "master" / "qq_pri_1" / "file_002.mp4"
                target.parent.mkdir(parents=True, exist_ok=True)
                target.write_bytes(b"remote video payload")
                return target

            def fake_build_card(*, source_path, item, mime_type):
                return {
                    "summary_title": "测试视频.mp4",
                    "short_hint": "媒体文件 测试视频.mp4。",
                    "detail": {"summary": "媒体文件 测试视频.mp4。"},
                }

            with patch.object(service, "_fetch_remote_media_descriptor", return_value=descriptor):
                with patch.object(service, "_download_remote_media", side_effect=fake_download):
                    with patch.object(service, "_build_file_card", side_effect=fake_build_card):
                        result = service.fetch_media_from_urls(
                            profile_user_id="master",
                            session_id="qq_pri_1",
                            urls=[url],
                            timestamp=100,
                        )

            self.assertTrue(result["ok"])
            stale_after = store.get_attachment_inbox_item(
                profile_user_id="master",
                session_id="qq_pri_1",
                attachment_id=stale["attachment_id"],
            )
            self.assertIsNotNone(stale_after)
            self.assertEqual(stale_after["status"], "cleared")
            prompt = inbox.build_prompt_context(profile_user_id="master", session_id="qq_pri_1")
            self.assertIn("测试视频", prompt)
            self.assertNotIn("Requested format", prompt)

    def test_remote_media_download_with_ytdlp_does_not_force_best_format(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            store = MemoryStore(root / "db")
            inbox = AttachmentInboxService(store=store, base_dir=root / "attachments")
            service = AttachmentIngestService(
                base_dir=root / "attachments",
                store=store,
                attachment_service=inbox,
                vision_service=FakeVisionService(store),  # type: ignore[arg-type]
            )

            captured_options: dict[str, Any] = {}
            downloaded_target = root / "attachments" / "master" / "qq_pri_1" / "file_001.mp4"
            downloaded_target.parent.mkdir(parents=True, exist_ok=True)
            downloaded_target.write_bytes(b"video")

            class FakeYoutubeDL:
                def __init__(self, options: dict[str, Any]) -> None:
                    captured_options.update(options)

                def __enter__(self):
                    return self

                def __exit__(self, exc_type, exc, tb) -> bool:
                    return False

                def download(self, urls: list[str]) -> None:
                    return None

            fake_module = types.SimpleNamespace(YoutubeDL=FakeYoutubeDL)
            cookiefile = root / "bilibili-cookies.txt"
            cookiefile.write_text("# Netscape HTTP Cookie File\n", encoding="utf-8")
            with patch.dict(sys.modules, {"yt_dlp": fake_module}):
                with patch("companion_v01.attachment_ingest.config.REMOTE_MEDIA_YTDLP_COOKIEFILE", str(cookiefile)):
                    with patch("companion_v01.attachment_ingest.config.REMOTE_MEDIA_YTDLP_REFERER", "https://www.bilibili.com/"):
                        with patch.object(service, "_locate_downloaded_remote_media_file", return_value=downloaded_target):
                            result = service._download_remote_media_with_yt_dlp(
                                descriptor=RemoteMediaDescriptor(
                                    source_url="https://b23.tv/demo",
                                    webpage_url="https://www.bilibili.com/video/BVdemo",
                                    title="测试视频",
                                    ext="mp4",
                                    mime_type="video/mp4",
                                    kind="file",
                                    download_mode="yt_dlp",
                                    extractor="BiliBili",
                                    extractor_key="BiliBili",
                                ),
                                target_dir=downloaded_target.parent,
                                handle="file_001",
                                timeout=30.0,
                                max_bytes=0,
                            )

            self.assertEqual(result, downloaded_target)
            self.assertNotIn("format", captured_options)
            self.assertEqual(captured_options["socket_timeout"], 30.0)
            self.assertEqual(captured_options["cookiefile"], str(cookiefile))
            headers = captured_options["http_headers"]
            self.assertIn("Mozilla/5.0", headers["User-Agent"])
            self.assertEqual(headers["Referer"], "https://www.bilibili.com/")

    def test_remote_media_412_error_suggests_cookiefile_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            store = MemoryStore(root / "db")
            inbox = AttachmentInboxService(store=store, base_dir=root / "attachments")
            service = AttachmentIngestService(
                base_dir=root / "attachments",
                store=store,
                attachment_service=inbox,
                vision_service=FakeVisionService(store),  # type: ignore[arg-type]
            )

            message = service._humanize_remote_fetch_error(
                "[BiliBili] 1ZJ6qBvEnZ: Unable to download JSON metadata: "
                "HTTP Error 412: Precondition Failed"
            )

            self.assertIn("平台风控", message)
            self.assertIn("REMOTE_MEDIA_YTDLP_COOKIEFILE", message)

    def test_ytdlp_common_options_support_browser_cookies(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            store = MemoryStore(root / "db")
            inbox = AttachmentInboxService(store=store, base_dir=root / "attachments")
            service = AttachmentIngestService(
                base_dir=root / "attachments",
                store=store,
                attachment_service=inbox,
                vision_service=FakeVisionService(store),  # type: ignore[arg-type]
            )

            with patch("companion_v01.attachment_ingest.config.REMOTE_MEDIA_YTDLP_COOKIEFILE", ""):
                with patch("companion_v01.attachment_ingest.config.REMOTE_MEDIA_YTDLP_COOKIES_FROM_BROWSER", "edge:Default"):
                    options = service._yt_dlp_common_options(timeout=12.0)

            self.assertEqual(options["socket_timeout"], 12.0)
            self.assertEqual(options["cookiesfrombrowser"], ("edge", "Default", None, None))
            self.assertNotIn("cookiefile", options)

    def test_remote_media_browser_cookie_copy_error_is_actionable(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            store = MemoryStore(root / "db")
            inbox = AttachmentInboxService(store=store, base_dir=root / "attachments")
            service = AttachmentIngestService(
                base_dir=root / "attachments",
                store=store,
                attachment_service=inbox,
                vision_service=FakeVisionService(store),  # type: ignore[arg-type]
            )

            message = service._humanize_remote_fetch_error(
                "ERROR: Could not copy Chrome cookie database. See "
                "https://github.com/yt-dlp/yt-dlp/issues/7271 for more info"
            )

            self.assertIn("浏览器仍在运行", message)
            self.assertIn("REMOTE_MEDIA_YTDLP_COOKIEFILE", message)

    def _wait_for_status(
        self,
        store: MemoryStore,
        *,
        profile_user_id: str,
        session_id: str,
        status: str,
    ) -> dict[str, Any]:
        deadline = time.time() + 3
        while time.time() < deadline:
            items = store.list_attachment_inbox_items(
                profile_user_id=profile_user_id,
                session_id=session_id,
                statuses=[status],
                limit=10,
            )
            if items:
                return items[0]
            time.sleep(0.05)
        self.fail(f"attachment did not reach status {status!r}")


if __name__ == "__main__":
    unittest.main()
